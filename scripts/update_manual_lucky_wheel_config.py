from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"
ASSET_DIR = ROOT / "custom" / "assets" / "manual-lucky-wheel-config"
LIST_IMAGE = ASSET_DIR / "page-overview.png"
MODAL_IMAGE = ASSET_DIR / "create-modal.png"

TITLE = "幸运转盘配置"
FONT_NAME = "微软雅黑"
BLUE = RGBColor(47, 85, 151)
CAPTION_GRAY = RGBColor(112, 126, 146)


def set_run(run, *, size: float = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def insert_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_body_after(anchor: Paragraph, text: str) -> Paragraph:
    paragraph = insert_after(anchor, style="Normal")
    run = paragraph.add_run(text)
    set_run(run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullet_after(anchor: Paragraph, label: str, text: str) -> Paragraph:
    paragraph = insert_after(anchor, style="List Bullet")
    label_run = paragraph.add_run(label)
    set_run(label_run, size=10, bold=True, color=BLUE)
    body_run = paragraph.add_run("：" + text)
    set_run(body_run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_image_after(anchor: Paragraph, image_path: Path, caption: str, width_inches: float = 6.35) -> Paragraph:
    image_para = insert_after(anchor, style="Normal")
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = image_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    image_para.paragraph_format.space_after = Pt(1)

    caption_para = insert_after(image_para, style="Caption")
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_para.add_run(caption)
    set_run(caption_run, size=9, color=CAPTION_GRAY)
    caption_para.paragraph_format.space_after = Pt(6)
    return caption_para


def reset_section(document: Document) -> Paragraph:
    start = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == TITLE:
            start = index
            break
    if start is None:
        raise RuntimeError(f"未找到《{TITLE}》章节。")

    end = len(document.paragraphs)
    for index in range(start + 1, len(document.paragraphs)):
        if document.paragraphs[index].style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            end = index
            break

    for paragraph in list(document.paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)

    anchor = document.paragraphs[start]
    for run in anchor.runs:
        set_run(run, size=10.5, bold=True, color=BLUE)
    return anchor


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    for path in (LIST_IMAGE, MODAL_IMAGE):
        if not path.exists():
            raise FileNotFoundError(path)

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = BACKUP_DIR / f"{DOCX_PATH.stem}.lucky-wheel-config-{stamp}.bak.docx"
    shutil.copy2(DOCX_PATH, backup_path)

    document = Document(DOCX_PATH)
    anchor = reset_section(document)

    last = add_body_after(
        anchor,
        "幸运转盘配置用于维护转盘活动的基础信息、奖项结构和页面启停状态。配置完成后会直接影响前台抽奖展示和奖励发放，因此在保存前应同时复核奖项数量、奖项类型、中奖概率和图片素材。",
    )
    last = add_image_after(
        last,
        LIST_IMAGE,
        "图：幸运转盘配置列表页，用于筛选、查看转盘配置状态并发起新增或维护操作。",
    )

    for label, text in [
        ("页面查询", "页面顶部支持按转盘名称和是否启用筛选，点击“搜索”可快速定位目标配置，点击“重置”可恢复全部数据列表。"),
        ("列表字段", "列表展示 ID、转盘名称、是否启用、奖项数量、描述和创建时间，方便确认活动标识、启用状态和配置完整度。"),
        ("启用状态", "“是否启用”用于控制该转盘配置是否可被前台活动引用。正式启用前应确认奖项、概率和图片均已配置完成。"),
        ("行内操作", "列表支持“详情、编辑、删除”等操作。编辑时可调整奖项信息和描述，删除前需确认该配置未被线上活动继续使用。"),
        ("新增入口", "点击“新增”打开新增幸运转盘配置弹窗，用于录入转盘名称、启用状态和具体奖项内容。"),
    ]:
        last = add_bullet_after(last, label, text)

    last = add_image_after(
        last,
        MODAL_IMAGE,
        "图：新增幸运转盘配置弹窗，用于维护转盘基础信息、奖项配置和奖项图片。",
        width_inches=6.2,
    )

    for label, text in [
        ("转盘名称", "转盘名称用于区分不同抽奖活动，建议直接使用活动主题或奖池名称，便于后续筛选和关联活动。"),
        ("是否启用", "弹窗中的启用开关用于设置该配置保存后是否立即生效。若奖项尚未全部配置完成，建议先关闭启用状态再保存。"),
        ("奖项配置", "奖项区域按行维护每个奖项，支持填写奖项名称、奖项标识、奖项类型、奖项数值和中奖概率。每一行都代表转盘上的一个独立奖项。"),
        ("奖项标识", "奖项标识用于系统识别具体奖项，建议保持唯一且与活动配置口径一致，避免后续派奖或统计时出现混淆。"),
        ("奖项类型与数值", "奖项类型决定奖励发放方式，奖项数值用于填写对应的奖励额度、数量或关联值。录入时应与实际奖励规则保持一致。"),
        ("中奖概率", "中奖概率用于控制每个奖项的抽中权重。修改概率后应整体复核各奖项分布，避免出现总概率异常或高价值奖项权重设置错误。"),
        ("奖项图标", "每个奖项都可上传图标，页面提示建议使用 80×80 或接近比例的 PNG、JPG、GIF 图片。上传前应确认图片清晰、内容与奖项名称一致。"),
        ("添加奖项", "点击“添加奖项”可继续增加奖项行。新增后应同步补全名称、类型、概率和图标，避免保存后出现空奖项。"),
        ("描述", "描述字段用于补充转盘活动说明、使用场景或备注信息，便于运营人员后续识别该配置用途。"),
        ("保存检查", "点击“确定”前，重点检查奖项数量是否完整、概率是否合理、图标是否上传成功，以及启用状态是否符合上线计划。"),
        ("取消与关闭", "点击“取消”或右上角关闭按钮会放弃当前未保存内容，适用于中止新增或重新整理奖项配置。"),
    ]:
        last = add_bullet_after(last, label, text)

    document.save(DOCX_PATH)
    print(f"updated lucky wheel config; backup={backup_path}")


if __name__ == "__main__":
    main()
