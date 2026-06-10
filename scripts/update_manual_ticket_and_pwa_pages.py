from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"

SITE_ASSET_DIR = ROOT / "custom" / "assets" / "manual-site-app-config"
TICKET_ASSET_DIR = ROOT / "custom" / "assets" / "manual-ticket-center"
GLOBAL_ASSET_DIR = ROOT / "custom" / "assets" / "manual-global-config"

PWA_IMAGE = SITE_ASSET_DIR / "pwa-config.png"
TICKET_LIST_IMAGE = TICKET_ASSET_DIR / "page-overview.png"
TICKET_MODAL_IMAGE = TICKET_ASSET_DIR / "add-config-modal.png"
TICKET_USAGE_IMAGE = TICKET_ASSET_DIR / "usage-records.png"
GLOBAL_IMAGE = GLOBAL_ASSET_DIR / "free-reward-limit.png"

FONT_NAME = "微软雅黑"
BLUE = RGBColor(47, 85, 151)
CAPTION_GRAY = RGBColor(112, 126, 146)


def set_run(run, *, size: float = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def insert_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_body_after(anchor: Paragraph, text: str) -> Paragraph:
    paragraph = insert_after(anchor, style="Normal")
    run = paragraph.add_run(text)
    set_run(run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullet_after(anchor: Paragraph, label: str, text: str) -> Paragraph:
    paragraph = insert_after(anchor, style="List Bullet")
    label_run = paragraph.add_run(label)
    set_run(label_run, size=10, bold=True, color=BLUE)
    body_run = paragraph.add_run("：" + text)
    set_run(body_run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_picture_after(anchor: Paragraph, image_path: Path, width: float = 6.45) -> Paragraph:
    paragraph = insert_after(anchor, style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def add_caption_after(anchor: Paragraph, text: str) -> Paragraph:
    paragraph = insert_after(anchor, style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run(run, size=9, color=CAPTION_GRAY)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_heading_after(anchor: Paragraph, text: str, level_style: str = "Heading 3", size: float = 10.5) -> Paragraph:
    paragraph = insert_after(anchor, style=level_style)
    run = paragraph.add_run(text)
    set_run(run, size=size, bold=True, color=BLUE)
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def find_heading(document: Document, title: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == title:
            return paragraph
    raise RuntimeError(f"未找到章节：{title}")


def clear_section(document: Document, title: str) -> Paragraph:
    paragraphs = document.paragraphs
    start = None
    for index, paragraph in enumerate(paragraphs):
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == title:
            start = index
            break
    if start is None:
        raise RuntimeError(f"未找到“{title}”章节。")

    end = len(paragraphs)
    for index in range(start + 1, len(paragraphs)):
        if paragraphs[index].style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            end = index
            break

    for paragraph in list(paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)

    return document.paragraphs[start]


def replace_pwa_subsection(document: Document) -> None:
    paragraphs = document.paragraphs
    start = None
    end = None
    for index, paragraph in enumerate(paragraphs):
        if "PWA配置" in paragraph.text.strip():
            start = index
            break
    if start is None:
        raise RuntimeError("未找到 PWA配置 位置。")

    for index in range(start + 1, len(paragraphs)):
        if paragraphs[index].style.name == "Heading 4":
            end = index
            break
    if end is None:
        raise RuntimeError("未找到 PWA配置 后续小节。")

    anchor = paragraphs[start - 1]
    for paragraph in list(paragraphs[start:end]):
        delete_paragraph(paragraph)

    last = add_heading_after(anchor, "PWA配置", "Heading 4", 10.5)
    last = add_picture_after(last, PWA_IMAGE, 6.45)
    last = add_caption_after(last, "图8：PWA 配置页面，用于维护应用名称、显示模式、屏幕方向和桌面图标。")

    bullets = [
        ("应用名称", "用于配置应用在安装页或浏览器提示中的名称，建议与站点品牌名保持一致，避免玩家安装后无法识别。"),
        ("桌面显示名称", "用于控制用户添加到桌面后的图标名称，应保持简短清晰，避免在移动端桌面被截断。"),
        ("应用描述", "用于补充应用用途或品牌说明，建议与站点定位、推广文案保持一致。"),
        ("显示模式", "当前页面支持配置显示模式，例如 minimal-ui，用于控制 PWA 打开后的浏览器 UI 保留方式。调整时应结合移动端实际打开效果检查导航栏和返回行为。"),
        ("屏幕方向", "用于限制应用安装后的展示方向，例如 portrait 竖屏。保存前需确认与前台主要页面的适配方向一致。"),
        ("启用状态", "通过启用状态控制当前 PWA 配置是否生效。关闭后应同步评估是否影响桌面快捷入口和安装引导。"),
        ("应用图标", "页面底部需要上传桌面图标和 512×512 启动图。上传时要按页面提示控制尺寸、格式和文件大小，避免安装时图标模糊或启动页拉伸。"),
        ("操作方式", "点击右上角编辑进入编辑态，更新字段或图片后保存；上线前建议实际安装一次，确认名称、图标、启动图和打开体验都符合预期。"),
    ]
    for label, text in bullets:
        last = add_bullet_after(last, label, text)


def update_ticket_center(document: Document) -> None:
    anchor = clear_section(document, "票券中心")
    last = add_body_after(
        anchor,
        "票券中心用于维护票券配置、发放状态和基础使用规则，适合运营人员统一管理金蛋、随机红包、优惠百分比等票券模板，并按业务需要随时新增、停发或恢复。",
    )
    last = add_picture_after(last, TICKET_LIST_IMAGE, 6.45)
    last = add_caption_after(last, "图：票券中心列表页，用于查询票券配置、查看派发与使用情况，并执行停发、恢复正常、修改等操作。")

    bullets = [
        ("搜索区", "页面顶部支持按票券名称搜索，并提供重置与新增配置入口，适合在票券数量较多时快速定位目标配置。"),
        ("列表字段", "列表集中展示 ID、票券类型、票券名称、有效期、已派发数量、已使用数量、创建时间和状态，便于同步核对配置内容与实际使用情况。"),
        ("状态管理", "已启用票券可执行停发，停发后可恢复正常；操作前应确认是否还有前台活动、任务或派奖流程继续引用该票券。"),
        ("派发与使用数量", "已派发数量用于观察当前票券投放规模，已使用数量用于判断核销消耗情况。活动复盘时可与票券使用记录联动查看。"),
        ("修改", "点击修改进入编辑态，可调整票券名称、使用须知、金额流水规则、有效期和启用状态。修改前需确认不会影响正在发放中的奖励口径。"),
        ("删除", "只有停发状态的票券才提供删除操作。删除前应确认该票券不再被活动、任务或历史派奖流程引用，避免后续查询缺少依据。"),
    ]
    for label, text in bullets:
        last = add_bullet_after(last, label, text)

    last = add_picture_after(last, TICKET_MODAL_IMAGE, 5.75)
    last = add_caption_after(last, "图：新增票券配置弹窗，用于填写票券类型、票券信息、金额流水规则、有效期和启用状态。")

    modal_bullets = [
        ("票券类型", "支持优惠百分比、随机红包和金蛋三种类型。应先选择类型，再继续填写对应金额规则，避免后续字段口径不一致。"),
        ("票券名称", "用于后台识别和前台展示，建议命名清晰并带上活动主题或适用场景，方便后续搜索与复盘。"),
        ("使用须知", "用于填写票券使用规则、领取条件、有效期说明和失效条件。文案应尽量覆盖玩家最关心的使用限制。"),
        ("票券图片", "上传票券图标时需按页面要求控制在 64×64 或等比例，并校验格式和大小，避免前台卡片展示失真。"),
        ("设置金额及流水", "不同票券类型对应不同配置区域。优惠百分比需要填写优惠值和打码量倍数；随机红包和金蛋需要配置稽核倍数及概率规则。"),
        ("概率规则", "随机红包和金蛋配置中需维护金额区间或奖励金额与中奖概率。提交前应检查单项概率与总概率口径，避免派奖异常。"),
        ("有效期", "支持长期有效和领取票券后生效两类方式。选择领取票券后时，应同步明确生效时点和失效时长。"),
        ("充值限制", "可限制领取或使用票券时的充值金额门槛，适合用于充值类优惠券或高门槛奖励。"),
        ("是否启用", "新增时可直接设置启用状态。若配置尚未核对完成，建议先关闭，待活动排期确认后再启用。"),
        ("保存与取消", "点击确定保存当前票券配置；点击取消或关闭弹窗会放弃未保存内容。正式提交前建议复核名称、规则、概率和限制条件。"),
    ]
    for label, text in modal_bullets:
        last = add_bullet_after(last, label, text)


def update_ticket_usage(document: Document) -> None:
    anchor = clear_section(document, "票券使用记录")
    last = add_body_after(
        anchor,
        "票券使用记录用于查询玩家票券从派发到核销的完整使用明细，适合运营、客服和风控人员按用户、票券类型和状态追踪奖励去向。",
    )
    last = add_picture_after(last, TICKET_USAGE_IMAGE, 6.45)
    last = add_caption_after(last, "图：票券使用记录页面，用于筛选票券明细并核对派发时间、使用时间、票券类型和当前状态。")

    bullets = [
        ("筛选条件", "页面顶部支持按用户 ID、票券 ID、票券类型和状态组合筛选，适合快速定位某一张票券或某一位用户的奖励记录。"),
        ("时间字段", "派发时间用于确认票券发放时间点，使用时间用于确认实际核销或使用时间。处理客诉时建议两个时间一起核对。"),
        ("用户信息", "用户 ID 与用户名用于定位具体玩家，便于继续联查会员资料、账变记录、活动记录或客服处理记录。"),
        ("票券名称与类型", "票券名称用于识别具体活动或奖励来源，票券类型用于区分金蛋、随机红包、充值优惠券等配置类别。"),
        ("状态", "当前页面主要展示待核销与已核销状态。待核销表示票券已发放但尚未完成使用，已核销表示票券已被消耗或完成兑换。"),
        ("分页查看", "底部支持页码跳转和每页条数切换，记录较多时可分批查看，避免遗漏早期派发数据。"),
        ("使用要点", "排查奖励未到账、重复领取或票券失效问题时，建议结合票券中心配置、活动规则和账变记录交叉核对。"),
    ]
    for label, text in bullets:
        last = add_bullet_after(last, label, text)


def update_global_config(document: Document) -> None:
    anchor = clear_section(document, "全局配置")
    last = add_body_after(
        anchor,
        "全局配置用于维护优惠活动的统一限制规则。当前页面只有“免费奖励领取限制”卡片，适合在活动派奖前统一控制免费奖励的领取门槛。",
    )
    last = add_picture_after(last, GLOBAL_IMAGE, 6.2)
    last = add_caption_after(last, "图：全局配置页面，当前仅包含“免费奖励领取限制”卡片，用于设置用户余额门槛和启用状态。")

    bullets = [
        ("配置状态", "通过开关控制免费奖励领取限制是否生效。启用后，系统会按当前限制条件统一拦截符合条件的奖励领取。"),
        ("领取限制", "页面支持设置用户余额阈值。当用户余额高于设定值时，不允许领取免费奖励。调整该数值前应先确认活动策略和用户分层方案。"),
        ("影响范围", "该限制会影响除充值直送活动以外的系统赠送类活动、任务奖励和代理佣金领取，因此保存前需同步评估运营侧与代理侧影响。"),
        ("保存操作", "修改开关或限制值后点击保存生效。上线前建议先确认公告口径、活动说明和客服话术，避免玩家因规则变化产生疑问。"),
        ("使用要点", "当站点需要限制高余额用户继续领取免费福利时，可优先使用该卡片统一控制；调整后应关注相关活动领取数据是否出现明显波动。"),
    ]
    for label, text in bullets:
        last = add_bullet_after(last, label, text)


def copy_image_from_docx(document: Document, paragraph_index: int, image_index: int, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    blips = document.paragraphs[paragraph_index]._element.xpath('.//*[local-name()="blip"]')
    if image_index >= len(blips):
        raise RuntimeError(f"段落 {paragraph_index} 中不存在第 {image_index + 1} 张图片。")
    rid = blips[image_index].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
    part = document.part.related_parts[rid]
    output_path.write_bytes(part.blob)


def extract_assets(document: Document) -> None:
    copy_image_from_docx(document, 151, 0, PWA_IMAGE)
    copy_image_from_docx(document, 582, 0, TICKET_LIST_IMAGE)
    copy_image_from_docx(document, 582, 1, TICKET_MODAL_IMAGE)
    copy_image_from_docx(document, 587, 0, TICKET_USAGE_IMAGE)
    copy_image_from_docx(document, 592, 0, GLOBAL_IMAGE)


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{DOCX.stem}.ticket-pwa-{stamp}.bak.docx"
    shutil.copy2(DOCX, backup)

    document = Document(DOCX)
    extract_assets(document)
    replace_pwa_subsection(document)
    update_ticket_center(document)
    update_ticket_usage(document)
    update_global_config(document)
    document.save(DOCX)

    print(
        "updated manual sections;"
        f" backup={backup};"
        f" pwa={PWA_IMAGE};"
        f" ticket_list={TICKET_LIST_IMAGE};"
        f" ticket_modal={TICKET_MODAL_IMAGE};"
        f" ticket_usage={TICKET_USAGE_IMAGE};"
        f" global={GLOBAL_IMAGE}"
    )


if __name__ == "__main__":
    main()
