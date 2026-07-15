from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
BACKUP_DIR = ROOT / "backups"
VIP_SETTINGS = "VIP设置"
VIP_DOWNGRADE = "VIP降级配置"
OLD_SUMMARY = (
    "VIP设置用于维护会员等级规则、保级与降级条件、各档位俸禄以及大 R 分层标准，"
    "是统一管理会员成长体系和高价值用户分层策略的核心页面。"
)
NEW_SUMMARY = (
    "VIP设置用于维护会员等级规则、各档位俸禄以及大 R 分层标准，"
    "是统一管理会员成长体系和高价值用户分层策略的核心页面。"
)
EXPECTED_REMAINING_HEADINGS = [
    "VIP晋级条件",
    "日俸禄",
    "周俸禄",
    "月俸禄",
    "大R等级定义",
]


def find_docx() -> Path:
    candidates = sorted(
        path for path in ROOT.glob("*.docx") if not path.name.startswith("~$")
    )
    if len(candidates) != 1:
        raise RuntimeError(f"expected one editable docx, found {len(candidates)}")
    return candidates[0]


def ensure_writable(path: Path) -> None:
    try:
        with path.open("r+b"):
            pass
    except PermissionError as exc:
        raise RuntimeError(f"DOCX_LOCKED: close the document and rerun: {path}") from exc


def backup_docx(path: Path) -> Path:
    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{path.stem}.remove-vip-downgrade-{stamp}.bak.docx"
    shutil.copy2(path, backup)
    return backup


def drawing_count(paragraphs) -> int:
    return sum(
        len(paragraph._element.xpath(".//w:drawing")) for paragraph in paragraphs
    )


def find_unique(paragraphs, *, text: str, style: str, start: int, end: int) -> int:
    matches = [
        index
        for index in range(start, end)
        if paragraphs[index].style.name == style
        and paragraphs[index].text.strip() == text
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"expected one {style} heading {text!r}, found {len(matches)}"
        )
    return matches[0]


def vip_section_bounds(paragraphs) -> tuple[int, int]:
    start = find_unique(
        paragraphs,
        text=VIP_SETTINGS,
        style="Heading 2",
        start=0,
        end=len(paragraphs),
    )
    end = next(
        (
            index
            for index in range(start + 1, len(paragraphs))
            if paragraphs[index].style.name in ("Heading 1", "Heading 2")
        ),
        len(paragraphs),
    )
    return start, end


def validate_saved(path: Path, expected_total_images: int) -> None:
    document = Document(path)
    paragraphs = document.paragraphs
    full_text = "\n".join(paragraph.text for paragraph in paragraphs)
    if VIP_DOWNGRADE in full_text or "保级与降级条件" in full_text:
        raise RuntimeError("VIP downgrade text remains after save")

    start, end = vip_section_bounds(paragraphs)
    headings = [
        paragraph.text.strip()
        for paragraph in paragraphs[start:end]
        if paragraph.style.name == "Heading 3"
    ]
    if headings != EXPECTED_REMAINING_HEADINGS:
        raise RuntimeError(f"unexpected remaining VIP headings: {headings}")
    if drawing_count(paragraphs) != expected_total_images:
        raise RuntimeError("document image count changed by more than one")
    if drawing_count(paragraphs[start:end]) != 5:
        raise RuntimeError("expected five remaining images in VIP settings")


def main() -> None:
    path = find_docx()
    ensure_writable(path)
    document = Document(path)
    paragraphs = document.paragraphs
    total_images_before = drawing_count(paragraphs)

    vip_start, vip_end = vip_section_bounds(paragraphs)
    summary_matches = [
        paragraph
        for paragraph in paragraphs[vip_start + 1 : vip_end]
        if paragraph.text.strip() == OLD_SUMMARY
    ]
    if len(summary_matches) != 1:
        raise RuntimeError(f"expected one VIP summary, found {len(summary_matches)}")

    downgrade_start = find_unique(
        paragraphs,
        text=VIP_DOWNGRADE,
        style="Heading 3",
        start=vip_start + 1,
        end=vip_end,
    )
    downgrade_end = next(
        index
        for index in range(downgrade_start + 1, vip_end)
        if paragraphs[index].style.name in ("Heading 1", "Heading 2", "Heading 3")
    )
    removal = paragraphs[downgrade_start:downgrade_end]
    if drawing_count(removal) != 1:
        raise RuntimeError("expected exactly one image in VIP downgrade section")

    backup = backup_docx(path)
    summary_matches[0].text = NEW_SUMMARY
    for paragraph in reversed(removal):
        element = paragraph._element
        element.getparent().remove(element)

    document.save(path)
    validate_saved(path, expected_total_images=total_images_before - 1)
    print(f"updated={path}")
    print(f"backup={backup}")


if __name__ == "__main__":
    main()
