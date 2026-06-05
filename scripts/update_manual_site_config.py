from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

base = Path('.')
docx_path = base / 'B端后台操作手册.docx'
img_dir = base / 'custom' / 'assets' / 'manual-site-config'
img_dir.mkdir(parents=True, exist_ok=True)

font_paths = [
    'C:/Windows/Fonts/msyh.ttc',
    'C:/Windows/Fonts/simhei.ttf',
    'C:/Windows/Fonts/simsun.ttc',
]
def font(size, bold=False):
    for p in font_paths:
        if Path(p).exists():
            return ImageFont.truetype(p, size)
    return ImageFont.load_default()

f_title = font(26, True)
f_text = font(18)
f_small = font(15)
f_badge = font(13)
blue = (47, 141, 255)
deep = (37, 64, 97)
line = (225, 230, 238)
muted = (112, 126, 146)
text = (16, 24, 40)
red = (239, 68, 68)


def rounded(draw, box, r, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=r, fill=fill, outline=outline, width=width)

def header(draw, w, title, action='编辑'):
    rounded(draw, (8, 8, w-8, 92), 10, (255,255,255), line)
    draw.rectangle((8, 72, w-8, 73), fill=line)
    draw.rectangle((28, 31, 33, 55), fill=blue)
    draw.text((43, 30), title, font=f_title, fill=text)
    bw = 96 if action == '新增' else 82
    rounded(draw, (w-120, 28, w-28, 60), 6, blue)
    draw.text((w-97, 35), '+ 新增' if action == '新增' else '编辑', font=f_small, fill=(255,255,255))

def save_card(name, title, draw_body, h=430, w=1280, action='编辑'):
    im = Image.new('RGB', (w,h), (248,250,253))
    d = ImageDraw.Draw(im)
    rounded(d, (4,4,w-4,h-4), 10, (255,255,255), line)
    header(d,w,title,action)
    draw_body(d,w,h)
    path = img_dir / f'{name}.png'
    im.save(path)
    return path

def body_icons(d,w,h):
    d.text((270,128),'Logo',font=f_text,fill=text); d.text((320,128),'*',font=f_text,fill=red)
    rounded(d,(270,158,430,238),8,(250,251,253),(207,213,222))
    d.text((292,184),'WTGAME',font=font(30,True),fill=(42,88,140))
    d.text((270,258),'要求 200*60规格或等比例，PNG、JPG、GIF图片，500KB以内',font=f_small,fill=muted)
    d.text((760,128),'Favicon',font=f_text,fill=text); d.text((830,128),'*',font=f_text,fill=red)
    rounded(d,(760,158,840,238),8,(18,24,32),(207,213,222))
    d.text((779,183),'ICON',font=font(19,True),fill=(255,176,44))
    d.text((760,258),'要求 32*32规格或等比例，PNG、JPG、GIF图片，100KB以内',font=f_small,fill=muted)

def body_currency(d,w,h):
    d.text((360,128),'货币图标',font=f_text,fill=text); d.text((440,128),'*',font=f_text,fill=red)
    rounded(d,(360,160,440,240),8,(255,248,220),(207,213,222))
    d.ellipse((370,168,430,228),fill=(245,183,0),outline=(224,150,0),width=6)
    d.text((391,181),'$',font=font(42,True),fill=(255,236,115))
    d.text((780,128),'货币符号',font=f_text,fill=text); d.text((860,128),'*',font=f_text,fill=red)
    d.text((780,164),'$',font=font(28,True),fill=text)

def body_share(d,w,h):
    labels=['og:title  *  -','og:description  *','og:image  *','og:url  *  -']
    y=116
    for lab in labels:
        d.text((62,y),lab,font=f_small,fill=text)
        if 'image' in lab:
            rounded(d,(180,y-8,360,y+82),6,(246,248,252),(218,224,232)); d.text((245,y+25),'未上传',font=f_text,fill=muted)
            y+=96
        else:
            y+=42
    d.text((990,112),'预览',font=f_text,fill=deep)
    rounded(d,(990,145,1210,335),8,(255,255,255),line)
    d.rectangle((990,145,1210,245),fill=(44,177,200))
    d.text((1080,184),'OG',font=font(28,True),fill=(255,255,255))
    d.text((1010,262),'标题',font=font(17,True),fill=text)
    d.text((1010,290),'描述',font=f_small,fill=text)
    d.text((1010,316),'https://example.com',font=f_small,fill=(24,91,255))

def mini_card(d,x,y,title,tag,on=True,icon='coin'):
    rounded(d,(x,y,x+235,y+170),6,(255,255,255),line)
    rounded(d,(x,y,x+235,y+42),6,(248,250,253),None)
    rounded(d,(x+12,y+12,x+64,y+26),8,(237,245,255))
    d.text((x+17,y+10),tag,font=f_badge,fill=(37,99,235))
    d.text((x+75,y+10),title,font=f_small,fill=text)
    rounded(d,(x+190,y+12,x+220,y+26),9,blue if on else (190,194,201))
    d.ellipse((x+205 if on else x+192,y+13,x+218 if on else x+205,y+26),fill=(255,255,255))
    d.text((x+20,y+58),'路由    /promotion?tab='+tag,font=f_small,fill=muted)
    d.text((x+20,y+82),'排序    1',font=f_small,fill=muted)
    d.line((x+20,y+106,x+215,y+106),fill=line)
    d.text((x+108,y+118),'图标',font=f_small,fill=muted,anchor='ma')
    rounded(d,(x+105,y+134,x+145,y+164),5,(255,239,199),line)

def body_activity(d,w,h):
    positions=[(24,112,'活动','activity',True),(276,112,'任务','task',True),(528,112,'返水','rebate',True),(780,112,'利息宝','interest',False),(1032,112,'VIP','vip',True),(24,300,'待领取','unclaimed',False),(276,300,'领取记录','claim_records',True)]
    for x,y,t,tag,on in positions: mini_card(d,x,y,t,tag,on)

def body_quick(d,w,h):
    positions=[(24,112,'分享赚钱','share_earn'),(276,112,'待领取','unclaimed'),(528,112,'利息宝','interest'),(780,112,'实时返水','rebate'),(1032,112,'任务中心','task_center'),(24,300,'VIP','vip')]
    for x,y,t,tag in positions: mini_card(d,x,y,t,tag,True)

def body_float(d,w,h):
    def fcard(x,y,title,items):
        rounded(d,(x,y,x+260,y+250),6,(255,255,255),line)
        d.text((x+18,y+24),title,font=f_text,fill=text)
        rounded(d,(x+202,y+22,x+238,y+40),9,blue); d.ellipse((x+221,y+23,x+238,y+40),fill=(255,255,255))
        d.text((x+20,y+62),'位置    左',font=f_small,fill=muted)
        d.text((x+20,y+88),'排序    1',font=f_small,fill=muted)
        d.text((x+20,y+114),'按钮数  '+str(len(items)),font=f_small,fill=muted)
        yy=y+148
        for it in items:
            rounded(d,(x+22,yy,x+48,yy+22),4,(255,242,204),line)
            d.text((x+58,yy),it,font=f_small,fill=text)
            rounded(d,(x+202,yy+2,x+238,yy+20),9,blue); d.ellipse((x+221,yy+3,x+238,yy+20),fill=(255,255,255))
            yy+=34
        d.line((x+14,y+218,x+246,y+218),fill=line)
        d.text((x+172,y+228),'编辑',font=f_small,fill=blue); d.text((x+214,y+228),'删除',font=f_small,fill=red)
    fcard(24,116,'左侧悬浮框1',['抽奖助力','邀请','任务'])
    fcard(314,116,'右侧悬浮框1',['转盘抽奖'])
    fcard(604,116,'右侧悬浮窗2',['vip','推广'])

cards = [
 ('site-icons','网站图标',body_icons,360,'编辑'),('currency','币种管理',body_currency,330,'编辑'),('share-card','分享卡片配置（Social Share / Open Graph）',body_share,380,'编辑'),('slide-promo','侧滑优惠中心',body_activity,520,'新增'),('quick-actions','首页快捷操作配置',body_quick,520,'新增'),('float-buttons','首页悬浮按钮配置',body_float,420,'新增')]
paths=[]
for name,title,fn,h,act in cards:
    paths.append(save_card(name,title,fn,h=h,action=act))


def set_run(run, size=10, bold=False, color=None):
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_para_spacing(p):
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15

def clear_para(p):
    for r in list(p.runs):
        p._p.remove(r._r)

def add_after(paragraph, text='', style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p = new_p
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p

def delete_paragraph(p):
    el = p._element
    el.getparent().remove(el)
    p._p = p._element = None

def set_heading(p, text, size, color=deep):
    clear_para(p)
    p.style = 'Heading 3'
    r=p.add_run(text)
    set_run(r,size=size,bold=True,color=color)
    set_para_spacing(p)

def add_bullet_after(anchor, prefix, rest):
    p=add_after(anchor, style='List Bullet')
    r=p.add_run(prefix)
    set_run(r,10,True,deep)
    r=p.add_run('：'+rest)
    set_run(r,10,False,None)
    set_para_spacing(p)
    return p

def add_normal_after(anchor, text, size=10, italic=False):
    p=add_after(anchor, style='Normal')
    r=p.add_run(text)
    set_run(r,size,False,muted if italic else None)
    r.italic = italic
    set_para_spacing(p)
    return p

def add_picture_after(anchor, path):
    p=add_after(anchor, style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.45))
    p.paragraph_format.space_after = Pt(2)
    return p

doc=Document(str(docx_path))
start = next(i for i,p in enumerate(doc.paragraphs) if p.text.strip()=='站点配置' and i>60)
end = next(i for i,p in enumerate(doc.paragraphs[start+1:], start+1) if p.text.strip()=='运营中心')
anchor = doc.paragraphs[start]
# delete old content between station config heading and next section
for p in list(doc.paragraphs[start+1:end]):
    delete_paragraph(p)
set_heading(anchor,'站点配置',12)

intro = add_after(anchor, style='List Bullet')
r=intro.add_run('站点配置用于维护站点基础展示、首页入口、分享卡片和活动入口等前台基础配置。运营人员可在该页面按配置卡片分别编辑网站图标、币种信息、分享卡片、侧滑优惠中心、首页快捷操作和首页悬浮按钮。配置保存后会影响前台品牌露出、用户访问入口、社交分享效果和首页运营位展示，上线前需重点检查图片尺寸、路由地址、排序、启用状态和移动端展示效果。')
set_run(r,10,False,None); set_para_spacing(intro)
last=intro
sections = [
 ('网站图标','图1：网站图标配置卡片，用于维护站点 Logo 与 Favicon。', paths[0], [
  ('Logo','上传或替换站点主 Logo，通常展示在前台页头、登录页或品牌露出位置；需按页面提示控制图片尺寸、格式和文件大小。'),
  ('Favicon','上传浏览器标签页图标，建议使用 32*32 或等比例图片，保证在浏览器标签、收藏夹等位置清晰展示。'),
  ('操作方式','点击右上角编辑进入编辑态，替换图片后保存；保存前检查图片是否清晰、是否留白过多、是否符合品牌规范。')]),
 ('币种管理','图2：币种管理配置卡片，用于维护站点货币图标和货币符号。', paths[1], [
  ('货币图标','上传前台展示用的币种图标，常用于钱包、金额、充值或活动奖励等场景。'),
  ('货币符号','配置金额前后展示的币种符号，例如 $；需与站点实际结算币种保持一致。'),
  ('操作方式','点击编辑后调整图标或符号，保存前确认币种展示不会与活动、充值、提现页面口径冲突。')]),
 ('分享卡片配置','图3：分享卡片配置卡片，用于维护 Social Share / Open Graph 分享信息。', paths[2], [
  ('og:title','设置外部渠道分享时展示的标题，建议简洁表达品牌或活动主题。'),
  ('og:description','设置分享描述文案，用于补充说明站点、活动或推广卖点。'),
  ('og:image','上传分享预览图，需检查图片比例、清晰度和裁切效果；未上传时预览区会显示默认占位。'),
  ('og:url','设置分享落地地址，需确认链接可访问且与当前活动或站点入口一致。'),
  ('预览','右侧预览区用于检查标题、描述、图片和链接的最终展示效果，保存前应整体核对。')]),
 ('侧滑优惠中心','图4：侧滑优惠中心配置卡片，用于维护前台侧滑活动入口。', paths[3], [
  ('活动卡片','每个卡片对应一个侧滑入口，包含标识、名称、路由、排序、图标和启用状态。'),
  ('新增','点击右上角新增创建新的侧滑入口，填写名称、路由、排序并上传图标。'),
  ('编辑/删除','点击卡片底部编辑可调整已有入口；删除前需确认该入口不再被活动或运营页面使用。'),
  ('启用状态','开关开启后前台展示，关闭后隐藏；排序数字越小通常展示越靠前。')]),
 ('首页快捷操作配置','图5：首页快捷操作配置卡片，用于维护首页快捷入口。', paths[4], [
  ('快捷入口','配置首页常用功能入口，如分享赚钱、待领取、利息宝、实时返水、任务中心、VIP 等。'),
  ('展示信息','每个入口包含名称、路由、排序、颜色、图标和启用状态，颜色用于前台入口视觉区分。'),
  ('操作方式','通过新增、编辑、删除维护入口列表；保存前检查路由是否正确、图标是否匹配业务含义、排序是否符合首页展示优先级。')]),
 ('首页悬浮按钮配置','图6：首页悬浮按钮配置卡片，用于维护首页左右侧悬浮按钮组。', paths[5], [
  ('悬浮框','按左侧或右侧配置悬浮按钮组，每组包含位置、排序、按钮数和启用状态。'),
  ('按钮项','每个按钮项可配置图标、名称、跳转目标和开关状态，用于抽奖助力、邀请、任务、VIP、推广等快速入口。'),
  ('操作方式','点击新增创建悬浮框，点击编辑维护按钮项；保存前检查左右位置是否冲突、按钮数量是否过多、移动端是否遮挡主要内容。')])
]
for title, caption, path, bullets in sections:
    p=add_after(last, style='Heading 3'); set_heading(p,title,11); last=p
    p=add_picture_after(last,path); last=p
    p=add_normal_after(last,caption,9,True); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; last=p
    for pref,rest in bullets:
        last=add_bullet_after(last,pref,rest)

# Apply requested formatting to the newly-added site config block only
# (above functions already set it)
doc.save(str(docx_path))
print('saved', docx_path)
print('images', len(paths), img_dir)
