from __future__ import annotations

import base64
import tempfile
import unittest
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

from docx import Document

import update_manual_sections as updater


PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


class UpdateManualSectionsTests(unittest.TestCase):
    def test_find_default_docx_ignores_word_lock_file(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            target = root / "manual.docx"
            target.write_bytes(b"docx")
            (root / "~$nual.docx").write_bytes(b"lock")

            self.assertEqual(updater.find_default_docx(root), target)

    def test_find_default_docx_requires_an_unambiguous_target(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            (root / "one.docx").write_bytes(b"one")
            (root / "two.docx").write_bytes(b"two")

            with self.assertRaisesRegex(RuntimeError, "multiple docx files"):
                updater.find_default_docx(root)

    def test_heading_two_range_includes_heading_three_children(self) -> None:
        document = Document()
        document.add_heading("Chapter", level=1)
        document.add_heading("Target", level=2)
        document.add_paragraph("Introduction")
        document.add_heading("Child", level=3)
        document.add_paragraph("Child body")
        document.add_heading("Next", level=2)

        location = updater.find_section_range(
            document, title="Target", heading_style="Heading 2"
        )

        self.assertEqual(location.start, 1)
        self.assertEqual(location.end, 5)

    def test_duplicate_heading_requires_parent_or_match_index(self) -> None:
        document = Document()
        document.add_heading("First", level=2)
        document.add_heading("Overview", level=3)
        document.add_heading("Second", level=2)
        document.add_heading("Overview", level=3)

        with self.assertRaisesRegex(RuntimeError, "ambiguous"):
            updater.find_section_range(
                document, title="Overview", heading_style="Heading 3"
            )

        location = updater.find_section_range(
            document,
            title="Overview",
            heading_style="Heading 3",
            parent_title="Second",
        )
        self.assertEqual(location.start, 3)

    def test_keep_existing_media_defaults_to_false(self) -> None:
        self.assertFalse(updater.resolve_keep_existing_media({}, {}))

    def test_chat_images_are_loaded_without_creating_output_directory(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            config_dir = Path(temp_dir)
            output_dir = config_dir / "manual-images"
            task = {
                "thread_read_file": "thread-read.json",
                "output_dir": str(output_dir),
                "expected_count": 1,
                "names": ["overview.png"],
            }
            mocked_result = ({}, {"id": "message-1"}, [("png", PNG_BYTES)])

            with patch.object(
                updater.chat_images,
                "extract_from_thread_read_file",
                return_value=mocked_result,
            ):
                loaded = updater.load_chat_images(task, config_dir)

            self.assertIn("overview.png", loaded)
            self.assertEqual(loaded["overview.png"].data, PNG_BYTES)
            self.assertFalse(output_dir.exists())

    def test_missing_path_falls_back_to_embedded_image(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            context = updater.BlockContext(
                config_dir=Path(temp_dir),
                embedded_images=[updater.EmbeddedImage(PNG_BYTES)],
                chat_images={},
            )

            source = updater.resolve_image_source(
                {"type": "image", "path": "missing.png"}, context
            )

            self.assertEqual(source, PNG_BYTES)

    def test_collects_embedded_images_in_memory(self) -> None:
        document = Document()
        document.add_heading("Target", level=2)
        paragraph = document.add_paragraph()
        paragraph.add_run().add_picture(BytesIO(PNG_BYTES))
        document.add_paragraph("Figure caption", style="Caption")
        document.add_heading("Next", level=2)
        location = updater.find_section_range(
            document, title="Target", heading_style="Heading 2"
        )

        images = updater.collect_section_images(document, location)

        self.assertEqual(len(images), 1)
        self.assertEqual(images[0].caption, "Figure caption")
        self.assertTrue(images[0].data.startswith(b"\x89PNG"))


if __name__ == "__main__":
    unittest.main()
