from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"
ASSET_DIR = ROOT / "custom" / "assets" / "manual-message-push"
OVERVIEW_IMAGE = ASSET_DIR / "message-push-list.png"

TITLE = "消息推送"
FONT_NAME = "微软雅黑"
BLUE = RGBColor(47, 85, 151)
CAPTION_GRAY = RGBColor(112, 126, 146)


ROWS = [
    ("9e054f55-0c0c...", "端午送粽子达...", "公告", "VIP玩家", "2026-06-08 16:30", "2026-06-08 16:35", "已结束", "rebeca", "2026-06-08T16:26:57.60...", "详情  撤回"),
    ("32334070-748...", "帮派专属", "跑马灯", "全体玩家", "2026-06-08 16:05", "2026-06-17 00:00", "已发送", "rebeca", "2026-06-08T16:05:33.36...", "详情  撤回"),
    ("1961d74c-aff9...", "帮派专属", "跑马灯", "VIP玩家", "2026-06-08 16:03", "2026-06-17 00:00", "已发送", "rebeca", "2026-06-08T16:03:59.78...", "详情  撤回"),
    ("7ca99faf-2d8f...", "游戏是我们的...", "跑马灯", "全体玩家", "2026-06-08 16:02", "2026-06-27 00:00", "已发送", "rebeca", "2026-06-08T16:02:18.19...", "详情  撤回"),
    ("7f978150-f93b...", "全服大型更新...", "消息", "全体玩家", "2026-06-08 15:56", "2026-06-18 00:00", "已发送", "rebeca", "2026-06-08T15:56:26.81...", "详情  撤回"),
    ("f3794ee5-892...", "游戏维护更新...", "公告", "全体玩家", "2026-06-08 15:46", "2026-06-30 00:00", "已发送", "rebeca", "2026-06-08T15:46:24.92...", "详情  撤回"),
    ("0203437e-d5c...", "端午假期福利", "公告", "全体玩家", "2026-06-06 17:27", "2026-06-19 00:00", "已发送", "rebeca", "2026-06-06T17:27:48.28...", "详情  撤回"),
    ("049c73df-18d...", "cs", "消息", "自定义玩家", "2026-05-27 17:56", "2026-05-29 00:00", "已结束", "qitian", "2026-05-27T17:56:35.16...", "详情  撤回"),
    ("1dc53aee-4f5a...", "css", "消息", "自定义玩家", "2026-05-27 17:55", "2026-05-29 00:00", "已结束", "qitian", "2026-05-27T17:55:31.21...", "详情  撤回"),
    ("d2798f2d-1e0...", "cs", "消息", "自定义玩家", "2026-05-27 17:49", "2026-05-29 00:00", "已结束", "qitian", "2026-05-27T17:49:29.08...", "详情  撤回"),
]


SECTIONS = [
    (
        "页面说明",
        [
            ("页面用途", "消息推送用于创建和管理面向玩家的站内消息、公告和跑马灯内容。运营人员可按消息类型、状态和标题查询推送记录，查看详情，或对仍需下线的消息执行撤回。"),
            ("适用场景", "适用于活动通知、维护公告、版本更新、节假日福利、定向 VIP 通知和指定玩家消息触达。推送前应确认标题、收件人范围、发送时间、结束时间和显示内容。"),
        ],
    ),
    (
        "筛选条件",
        [
            ("消息类型", "按消息、公告、跑马灯筛选列表。排查某类触达内容时先选择消息类型，再结合状态或标题缩小范围。"),
            ("状态", "按未开始、已发送、已结束、已撤回筛选。上线中的内容通常查看已发送；复核历史内容时查看已结束或已撤回。"),
            ("标题", "输入标题关键字查询，例如端午、维护、帮派专属等。标题不需要输入完整内容，建议使用较短关键字。"),
            ("搜索", "按当前筛选条件刷新列表。筛选无结果时，先检查类型和状态是否过窄，再清空标题重新查询。"),
            ("重置", "清空消息类型、状态和标题，恢复默认列表。"),
            ("新增", "打开新增弹窗，用于配置新的消息、公告或跑马灯推送。"),
        ],
    ),
    (
        "列表字段",
        [
            ("消息编号", "系统生成的唯一编号，用于定位单条推送记录。列表中可能截断显示，查看详情可确认完整编号。"),
            ("标题", "展示推送标题。标题应能让运营人员快速判断内容主题，避免使用无法识别的测试标题上线。"),
            ("消息类型", "显示该条记录属于消息、公告或跑马灯。不同类型影响玩家端展示位置和触达方式。"),
            ("收件人", "显示推送范围，包括全体玩家、VIP玩家、R玩家或自定义玩家。发送前必须确认范围，避免误发。"),
            ("发送时间", "显示实际或计划发送时间。立即发送的记录通常以创建后的发送时间为准，定时发送按配置时间触发。"),
            ("结束时间", "显示内容停止展示或失效的时间。活动类消息应与活动结束时间保持一致或略晚，避免提前下线。"),
            ("状态", "未开始表示等待发送，已发送表示当前已触达或展示中，已结束表示超过结束时间，已撤回表示人工下线。"),
            ("操作人", "记录最后创建或操作该消息的后台账号，用于问题追溯。"),
            ("操作时间", "记录最后一次操作时间。排查误发、撤回或内容变更时，结合操作人一起核对。"),
            ("操作", "详情用于查看完整内容和配置；撤回用于将消息下线。已撤回记录可能展示删除入口，应按权限和业务要求谨慎处理。"),
        ],
    ),
    (
        "新增消息",
        [
            ("选择消息类型", "点击“新增”后先选择消息、公告或跑马灯。选择前应确认玩家端展示位置和内容形式。"),
            ("设置弹窗频次", "可选择不弹窗、登录后弹一次、每次进入首页弹窗、每日一次或每小时一次。高频弹窗会影响玩家体验，只适合重要通知。"),
            ("选择收件人", "可选择全体玩家、VIP玩家、R玩家或自定义玩家。选择 VIP玩家时需要勾选 VIP 等级；选择 R玩家时需要勾选玩家类型；选择自定义玩家时需录入收件人 ID，多个 ID 用逗号分隔。"),
            ("填写标题和显示内容", "标题应简明描述推送主题；显示内容为玩家看到的正文，可通过编辑器设置基础格式。保存前检查错别字、活动时间、奖励条件和链接跳转。"),
            ("设置发送时间", "选择立即发送时，确认后直接触发；选择定时发送时，需要填写计划发送时间。定时内容应提前完成复核，避免临近发送才修改。"),
            ("设置结束时间", "结束时间为必填项。应晚于发送时间，并与公告、活动或维护窗口的有效期一致。"),
            ("提交保存", "点击确定后完成创建。若显示内容为空或必填项未填写，需补齐后再提交；点击取消或关闭会放弃本次新增。"),
        ],
    ),
    (
        "详情与撤回",
        [
            ("查看详情", "在目标记录行点击“详情”，查看消息编号、状态、消息类型、弹窗频次、收件人、发送时间、结束时间、标题和显示内容。用于复核已发送内容或排查玩家反馈。"),
            ("执行撤回", "对需要提前下线的消息点击“撤回”，在确认弹窗中核对标题后点击确认。撤回后状态更新为已撤回，玩家端不应继续展示该内容。"),
            ("撤回判断", "维护取消、活动规则变更、误选收件人、标题或正文有误时应及时撤回。已结束记录一般无需撤回，除非仍在玩家端异常展示。"),
        ],
    ),
    (
        "注意事项",
        [
            ("推送范围", "全体玩家和 VIP/R 玩家范围影响较大，提交前必须二次复核。自定义玩家 ID 要检查格式、分隔符和名单来源。"),
            ("时间配置", "发送时间不能晚于结束时间。定时发送内容应预留审核时间，避免刚创建就进入发送窗口。"),
            ("内容合规", "标题和正文不得包含错误活动规则、过期链接、未确认奖励承诺或内部后台信息。涉及充值、提现、维护和风控的通知要与对应业务负责人确认。"),
            ("弹窗频次", "频次越高对玩家打扰越明显。普通活动通知建议控制弹窗频次，重要维护或风险提示可适当提高。"),
            ("撤回影响", "撤回只处理展示和触达状态，不会自动补发新消息。若需要更正内容，应撤回旧消息后重新新增正确内容。"),
        ],
    ),
    (
        "常见问题",
        [
            ("为什么列表看不到刚创建的消息", "先点击重置清空筛选，再检查状态是否为未开始或已发送；如果选择了定时发送，发送时间未到时通常不会按已发送展示。"),
            ("为什么玩家没有收到消息", "检查收件人范围是否正确、玩家是否属于所选 VIP/R 类型或自定义 ID 名单、发送时间是否已到，以及结束时间是否已经过期。"),
            ("误发给全体玩家怎么办", "立即在列表中定位该标题并执行撤回；如果需要发送更正说明，再新增一条正确范围或正确内容的消息。"),
            ("已结束和已撤回有什么区别", "已结束通常是到达结束时间后自然失效；已撤回是运营人员手动提前下线。排查展示异常时，两种状态都要结合操作时间和结束时间判断。"),
        ],
    ),
]


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    for name in candidates:
        path = Path(name)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def text_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill: str) -> None:
    x1, y1, x2, y2 = box
    bbox = draw.textbbox((0, 0), text, font=font)
    x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
    y = y1 + (y2 - y1 - (bbox[3] - bbox[1])) / 2 - 1
    draw.text((x, y), text, font=font, fill=fill)


def draw_button(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, primary: bool, font) -> None:
    fill = "#348fff" if primary else "#ffffff"
    outline = "#348fff" if primary else "#d8dee8"
    text_fill = "#ffffff" if primary else "#344154"
    draw.rounded_rectangle(box, radius=4, fill=fill, outline=outline, width=1)
    text_center(draw, box, text, font, text_fill)


def draw_status(draw: ImageDraw.ImageDraw, x: int, y: int, status: str, font) -> None:
    color = "#43b200" if status == "已发送" else "#ff2c2c"
    bg = "#ecffd9" if status == "已发送" else "#fff0f0"
    draw.rounded_rectangle((x - 28, y - 9, x + 28, y + 9), radius=9, fill=bg)
    text_center(draw, (x - 28, y - 9, x + 28, y + 9), status, font, color)


def create_overview_image() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    image = Image.new("RGB", (1666, 705), "#ffffff")
    draw = ImageDraw.Draw(image)
    f12 = get_font(20)
    f13 = get_font(22)
    f13b = get_font(22, True)

    draw.rectangle((0, 0, 1666, 705), fill="#ffffff")

    y = 17
    draw.text((12, y + 6), "消息类型:", fill="#344154", font=f12)
    draw.rounded_rectangle((78, y, 180, y + 32), radius=3, fill="#ffffff", outline="#d8dee8")
    draw.text((90, y + 7), "请选择类型", fill="#aeb8c6", font=f12)
    draw.text((163, y + 7), "▼", fill="#b7c0ce", font=get_font(15))

    draw.text((196, y + 6), "状态:", fill="#344154", font=f12)
    draw.rounded_rectangle((232, y, 338, y + 32), radius=3, fill="#ffffff", outline="#d8dee8")
    draw.text((242, y + 7), "请选择状态", fill="#aeb8c6", font=f12)
    draw.text((321, y + 7), "▼", fill="#b7c0ce", font=get_font(15))

    draw.text((354, y + 6), "标题:", fill="#344154", font=f12)
    draw.rounded_rectangle((390, y, 532, y + 32), radius=3, fill="#ffffff", outline="#d8dee8")
    draw.text((401, y + 7), "请输入标题", fill="#aeb8c6", font=f12)

    draw_button(draw, (546, y, 600, y + 32), "搜索", True, f12)
    draw_button(draw, (620, y, 674, y + 32), "重置", False, f12)
    draw_button(draw, (688, y, 742, y + 32), "新增", True, f12)

    table_x, table_y = 12, 65
    widths = [108, 108, 162, 162, 220, 218, 146, 220, 162, 126]
    headers = ["消息编号", "标题", "消息类型", "收件人", "发送时间", "结束时间", "状态", "操作人", "操作时间", "操作"]
    row_h = 43
    header_h = 42
    table_w = sum(widths)

    draw.rectangle((table_x, table_y, table_x + table_w, table_y + header_h + row_h * len(ROWS)), outline="#e6eaf0", width=1)
    draw.rectangle((table_x, table_y, table_x + table_w, table_y + header_h), fill="#fafbfc", outline="#e6eaf0")

    x = table_x
    for width, header in zip(widths, headers):
        draw.line((x, table_y, x, table_y + header_h + row_h * len(ROWS)), fill="#e6eaf0", width=1)
        text_center(draw, (x, table_y, x + width, table_y + header_h), header, f13b, "#26364f")
        x += width
    draw.line((x, table_y, x, table_y + header_h + row_h * len(ROWS)), fill="#e6eaf0", width=1)

    for index, row in enumerate(ROWS):
        y1 = table_y + header_h + row_h * index
        y2 = y1 + row_h
        draw.rectangle((table_x, y1, table_x + table_w, y2), fill="#ffffff", outline="#e6eaf0")
        x = table_x
        for col, (width, value) in enumerate(zip(widths, row)):
            draw.line((x, y1, x, y2), fill="#e6eaf0", width=1)
            if col == 6:
                draw_status(draw, x + width // 2, y1 + row_h // 2, value, get_font(18))
            elif col == 9:
                draw.text((x + 30, y1 + 12), "详情", fill="#2386ff", font=f12)
                draw.text((x + 68, y1 + 12), "撤回", fill="#ff4d4f", font=f12)
            else:
                text_center(draw, (x, y1, x + width, y2), value, f12, "#2b3b52")
            x += width
        draw.line((x, y1, x, y2), fill="#e6eaf0", width=1)

    pager_y = 698 - 36
    pager_x = 1168
    for label in ["|‹", "‹‹", "‹", "1"]:
        w = 24 if label != "1" else 52
        draw.rounded_rectangle((pager_x, pager_y, pager_x + w, pager_y + 30), radius=3, fill="#ffffff", outline="#d8dee8")
        text_center(draw, (pager_x, pager_y, pager_x + w, pager_y + 30), label, f12, "#7b8798")
        pager_x += w + 10
    draw.text((pager_x, pager_y + 6), "/ 2", fill="#344154", font=f12)
    pager_x += 40
    for label in ["›", "››", "›|"]:
        draw.text((pager_x, pager_y + 6), label, fill="#344154", font=f13)
        pager_x += 36
    draw.rounded_rectangle((pager_x, pager_y, pager_x + 112, pager_y + 30), radius=3, fill="#ffffff", outline="#d8dee8")
    draw.text((pager_x + 16, pager_y + 7), "20条/页", fill="#344154", font=f12)
    draw.text((pager_x + 90, pager_y + 7), "▼", fill="#b7c0ce", font=get_font(15))
    draw.text((pager_x + 124, pager_y + 7), "共 39 条记录", fill="#344154", font=f12)

    image.save(OVERVIEW_IMAGE)


def set_run(run, *, size: float = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def insert_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_heading_after(anchor, text: str):
    paragraph = insert_after(anchor, style="Heading 4")
    run = paragraph.add_run(text)
    set_run(run, size=11, bold=True, color=BLUE)
    paragraph.paragraph_format.space_after = Pt(5)
    return paragraph


def add_body_after(anchor, text: str):
    paragraph = insert_after(anchor, style="Normal")
    run = paragraph.add_run(text)
    set_run(run)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullet_after(anchor, label: str, text: str):
    paragraph = insert_after(anchor, style="List Bullet")
    label_run = paragraph.add_run(label)
    set_run(label_run, bold=True, color=BLUE)
    body_run = paragraph.add_run("：" + text)
    set_run(body_run)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_picture_after(anchor, image_path: Path, width: float = 6.45):
    paragraph = insert_after(anchor, style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def add_caption_after(anchor, text: str):
    paragraph = insert_after(anchor, style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run(run, size=9, color=CAPTION_GRAY)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def reset_section(document: Document):
    start = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == TITLE:
            start = index
            break
    if start is None:
        raise RuntimeError("未找到“消息推送”章节。")

    end = len(document.paragraphs)
    for index in range(start + 1, len(document.paragraphs)):
        if document.paragraphs[index].style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            end = index
            break

    for paragraph in list(document.paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)

    anchor = document.paragraphs[start]
    for run in anchor.runs:
        set_run(run, size=12.5, bold=True, color=BLUE)
    return anchor


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    create_overview_image()

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{DOCX.stem}.message-push-{stamp}.bak.docx"
    shutil.copy2(DOCX, backup)

    document = Document(DOCX)
    anchor = reset_section(document)
    last = add_body_after(
        anchor,
        "消息推送用于统一管理玩家侧通知触达，包括消息、公告和跑马灯。页面支持按消息类型、状态和标题查询记录，并提供新增、详情查看、撤回等操作，适合运营人员发布活动通知、维护公告、版本更新和定向玩家提醒。",
    )

    last = add_heading_after(last, "页面说明")
    last = add_picture_after(last, OVERVIEW_IMAGE, 6.5)
    last = add_caption_after(last, "图：消息推送列表，用于查询消息、公告和跑马灯推送记录。")

    for section_title, items in SECTIONS:
        if section_title != "页面说明":
            last = add_heading_after(last, section_title)
        for label, text in items:
            last = add_bullet_after(last, label, text)

    try:
        document.save(DOCX)
    except PermissionError:
        print(f"PERMISSION_ERROR backup={backup}")
        raise

    print(f"updated message push; backup={backup}; image={OVERVIEW_IMAGE}")


if __name__ == "__main__":
    main()
