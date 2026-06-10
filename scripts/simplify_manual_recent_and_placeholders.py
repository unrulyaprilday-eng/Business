from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"
PLATFORM_IMAGE = ROOT / "custom" / "assets" / "manual" / "platform-notice" / "platform-notice-list.png"
PUSH_IMAGE = ROOT / "custom" / "assets" / "manual-message-push" / "message-push-list.png"

FONT_NAME = "微软雅黑"
BLUE = RGBColor(47, 85, 151)
CAPTION_GRAY = RGBColor(112, 126, 146)


PAGE_SUMMARIES = {
    "站点维护管理": [
        ("页面功能", "用于设置站点维护状态、维护时间和前台提示内容。"),
        ("主要操作", "查看维护配置，新增或编辑维护计划，启用、停用维护状态。"),
        ("使用说明", "保存前确认维护时间、影响站点和提示文案，避免误影响正常访问。"),
    ],
    "给平台提建议": [
        ("页面功能", "用于向平台提交产品、运营或系统使用建议。"),
        ("主要操作", "填写建议标题、建议内容和联系方式，提交后在列表查看处理状态。"),
        ("使用说明", "建议内容应描述具体问题、使用场景和期望结果，便于平台跟进。"),
    ],
    "会员投诉列表": [
        ("页面功能", "用于查看玩家提交的投诉、意见和处理状态。"),
        ("主要操作", "按会员、状态或时间筛选投诉，查看详情并记录处理结果。"),
        ("使用说明", "处理投诉前先核对会员账号、问题类型和相关订单或活动记录。"),
    ],
    "分享配置": [
        ("页面功能", "用于维护分享标题、描述、图片和渠道参数。"),
        ("主要操作", "编辑分享文案、配置分享图片，并设置渠道或落地页参数。"),
        ("使用说明", "发布前确认分享配置与活动、品牌和投放页面保持一致。"),
    ],
    "宣传管理": [
        ("页面功能", "用于维护站点宣传内容、推广素材和展示状态。"),
        ("主要操作", "新增、编辑、上架、下架宣传内容，检查标题、图片和跳转地址。"),
        ("使用说明", "发布前确认素材清晰、链接有效，内容与当前活动或站点配置一致。"),
    ],
    "任务中心": [
        ("页面功能", "用于配置玩家任务、查看任务状态和管理任务奖励。"),
        ("主要操作", "新增或编辑任务，设置任务条件、奖励、时间范围和启停状态。"),
        ("使用说明", "上线前确认任务条件、奖励发放规则和活动时间，避免重复奖励或无法完成。"),
    ],
    "落地页管理": [
        ("页面功能", "用于维护推广落地页、页面模板和投放链接。"),
        ("主要操作", "新增、编辑、预览、启用或停用落地页，配置标题、素材和跳转地址。"),
        ("使用说明", "投放前检查页面内容、渠道归属和跳转链路，确保推广链接可访问。"),
    ],
    "活动排行统计": [
        ("页面功能", "用于查看活动参与、排名和奖励相关统计。"),
        ("主要操作", "按活动、时间或玩家筛选数据，查看排行、参与人数和奖励发放情况。"),
        ("使用说明", "统计结果用于活动复盘，需结合活动规则和奖励记录一起核对。"),
    ],
    "优惠活动列表": [
        ("页面功能", "用于管理站点优惠活动和活动状态。"),
        ("主要操作", "查询活动，新增、编辑、复制、启用、停用或查看活动数据。"),
        ("使用说明", "上线前确认活动时间、参与条件、奖励规则和展示入口。"),
    ],
    "幸运转盘配置": [
        ("页面功能", "用于配置幸运转盘活动、奖品和抽奖规则。"),
        ("主要操作", "设置奖品、中奖概率、参与条件、活动时间和启停状态。"),
        ("使用说明", "发布前核对奖品库存、概率配置和奖励发放方式。"),
    ],
    "新增活动": [
        ("页面功能", "用于创建新的优惠活动。"),
        ("主要操作", "选择活动类型，填写活动信息、规则、奖励、展示素材和时间范围。"),
        ("使用说明", "提交前逐项检查必填字段、奖励条件和前台展示文案。"),
    ],
    "活动分析": [
        ("页面功能", "用于查看活动效果和玩家参与数据。"),
        ("主要操作", "按活动和时间筛选，查看参与人数、领取金额、转化效果等指标。"),
        ("使用说明", "分析时结合活动成本、充值投注表现和奖励发放记录判断效果。"),
    ],
    "优惠数据报表": [
        ("页面功能", "用于统计优惠活动相关数据。"),
        ("主要操作", "按活动、时间、渠道或玩家筛选，查看领取、消耗和转化数据。"),
        ("使用说明", "报表数据用于活动复盘和成本控制，口径以页面筛选条件为准。"),
    ],
    "票券中心": [
        ("页面功能", "用于管理优惠券、票券模板和发放规则。"),
        ("主要操作", "新增、编辑、启用、停用票券，配置面额、使用门槛和有效期。"),
        ("使用说明", "发券前确认适用范围、有效期和使用限制，避免与活动规则冲突。"),
    ],
    "票券使用记录": [
        ("页面功能", "用于查询玩家票券领取和使用记录。"),
        ("主要操作", "按玩家、票券、状态或时间筛选，查看使用时间和关联活动。"),
        ("使用说明", "排查奖励问题时可结合票券状态、账变记录和活动记录核对。"),
    ],
    "全局配置": [
        ("页面功能", "用于维护优惠活动相关的全局规则和默认参数。"),
        ("主要操作", "查看并编辑活动通用开关、限制条件和基础配置。"),
        ("使用说明", "全局配置会影响多个活动，保存前需确认影响范围。"),
    ],
    "会员提现设置": [
        ("页面功能", "用于配置会员提现规则和限制。"),
        ("主要操作", "设置提现额度、次数、手续费、审核规则和启停状态。"),
        ("使用说明", "调整后会影响玩家提现体验，需与风控和财务口径一致。"),
    ],
    "会员充值配置": [
        ("页面功能", "用于配置会员充值入口、限额和展示规则。"),
        ("主要操作", "维护充值金额、通道展示、优惠提示和启停状态。"),
        ("使用说明", "保存前确认充值通道可用、金额范围合理、前台展示正确。"),
    ],
    "充值订单列表": [
        ("页面功能", "用于查询会员充值订单和订单状态。"),
        ("主要操作", "按会员、订单号、通道、状态或时间筛选，查看订单详情。"),
        ("使用说明", "处理充值问题时应核对支付状态、到账状态和账变记录。"),
    ],
    "提现订单列表": [
        ("页面功能", "用于查询会员提现订单和处理进度。"),
        ("主要操作", "按会员、订单号、状态或时间筛选，查看审核和出款信息。"),
        ("使用说明", "处理前核对账户信息、风控结果和资金状态。"),
    ],
    "提现风控审核": [
        ("页面功能", "用于处理触发风控的提现订单。"),
        ("主要操作", "查看风控原因、会员信息和订单详情，执行通过、拒绝或补充审核。"),
        ("使用说明", "审核结果会影响资金出款，需按风控规则和财务要求处理。"),
    ],
    "支付通道数据": [
        ("页面功能", "用于查看支付通道的交易和成功率数据。"),
        ("主要操作", "按通道、时间或状态筛选，查看充值笔数、金额和成功率。"),
        ("使用说明", "用于判断通道质量，异常时联系通道配置或财务人员处理。"),
    ],
    "人工加扣款": [
        ("页面功能", "用于对会员账户执行人工加款或扣款。"),
        ("主要操作", "填写会员、金额、类型和备注，提交后查看处理结果。"),
        ("使用说明", "此操作直接影响会员余额，提交前必须核对会员和金额。"),
    ],
    "账变记录": [
        ("页面功能", "用于查询会员资金变动明细。"),
        ("主要操作", "按会员、账变类型、订单号或时间筛选，查看余额变化。"),
        ("使用说明", "排查资金问题时优先结合订单、活动和人工加扣款记录核对。"),
    ],
    "资金账变类型": [
        ("页面功能", "用于维护资金账变类型和展示名称。"),
        ("主要操作", "查看、启用、停用或调整账变类型配置。"),
        ("使用说明", "类型配置会影响账变筛选和前后台展示，调整前需确认口径。"),
    ],
    "打码管理": [
        ("页面功能", "用于查看和管理会员打码量要求。"),
        ("主要操作", "查询会员打码进度，核对活动、充值或人工调整产生的打码要求。"),
        ("使用说明", "处理提现或活动问题时，应结合打码完成情况判断。"),
    ],
    "综合数据报表": [
        ("页面功能", "用于查看站点综合运营数据。"),
        ("主要操作", "按时间筛选充值、提现、投注、输赢、会员等关键指标。"),
        ("使用说明", "用于日常经营看板和趋势分析，需注意筛选时间口径。"),
    ],
    "游戏人数据表": [
        ("页面功能", "用于查看游戏维度的玩家人数数据。"),
        ("主要操作", "按时间、游戏或渠道筛选，查看活跃、投注和参与人数。"),
        ("使用说明", "用于判断游戏热度和玩家分布，分析时结合投注报表。"),
    ],
    "充值留存报表": [
        ("页面功能", "用于查看充值用户后续留存情况。"),
        ("主要操作", "按首充日期、渠道或站点筛选，查看次日、多日留存。"),
        ("使用说明", "用于评估充值用户质量和运营活动效果。"),
    ],
    "代理结算明细": [
        ("页面功能", "用于查看代理结算明细和结算结果。"),
        ("主要操作", "按代理、结算周期或状态筛选，查看佣金、返水和调整项。"),
        ("使用说明", "结算前需核对代理配置、会员数据和异常调整。"),
    ],
    "VIP数据统计": [
        ("页面功能", "用于统计各 VIP 等级会员数据。"),
        ("主要操作", "按时间或等级查看会员数量、充值、投注和升降级情况。"),
        ("使用说明", "用于评估 VIP 结构和高价值会员运营效果。"),
    ],
    "数据排行": [
        ("页面功能", "用于查看会员、游戏或渠道相关排行。"),
        ("主要操作", "按指标和时间筛选，查看充值、投注、输赢等排名。"),
        ("使用说明", "排行用于发现重点会员和异常数据，需结合明细记录核对。"),
    ],
    "游戏记录": [
        ("页面功能", "用于查询玩家游戏投注记录。"),
        ("主要操作", "按玩家、游戏、订单号、状态或时间筛选，查看投注明细。"),
        ("使用说明", "排查输赢争议时应结合游戏记录、账变和订单状态。"),
    ],
    "用户留存": [
        ("页面功能", "用于查看注册或活跃用户的留存情况。"),
        ("主要操作", "按日期、渠道或用户分组筛选，查看留存比例。"),
        ("使用说明", "留存分析用于评估渠道质量和运营活动效果。"),
    ],
    "用户分析": [
        ("页面功能", "用于分析用户结构、活跃和行为数据。"),
        ("主要操作", "按时间、渠道或用户类型筛选，查看新增、活跃、充值等指标。"),
        ("使用说明", "用于用户运营决策，需结合留存、充值和活动数据判断。"),
    ],
    "LTV": [
        ("页面功能", "用于查看用户生命周期价值数据。"),
        ("主要操作", "按注册日期、渠道或用户分组筛选，查看不同周期价值表现。"),
        ("使用说明", "用于评估投放回收和用户质量，需与渠道成本结合分析。"),
    ],
    "活动统计报表": [
        ("页面功能", "用于统计活动参与和奖励数据。"),
        ("主要操作", "按活动和时间筛选，查看参与人数、奖励金额和转化情况。"),
        ("使用说明", "用于活动效果复盘，需结合优惠数据和账变记录核对。"),
    ],
    "充值分布表": [
        ("页面功能", "用于查看充值金额、次数和用户分布。"),
        ("主要操作", "按时间、金额区间或渠道筛选，查看充值结构。"),
        ("使用说明", "用于分析充值习惯和优化充值配置。"),
    ],
    "任务统计报表": [
        ("页面功能", "用于查看任务参与和完成数据。"),
        ("主要操作", "按任务和时间筛选，查看参与人数、完成数和奖励发放。"),
        ("使用说明", "用于评估任务效果，需结合任务配置和奖励记录。"),
    ],
    "用户分群查询": [
        ("页面功能", "用于查询用户分群结果和分群成员。"),
        ("主要操作", "按分群条件、标签或会员信息筛选，查看命中用户。"),
        ("使用说明", "分群结果可用于活动、消息推送和精细化运营。"),
    ],
    "推荐频道管理": [
        ("页面功能", "用于维护前台推荐频道和展示内容。"),
        ("主要操作", "新增、编辑、排序、启用或停用推荐频道。"),
        ("使用说明", "调整前确认频道内容、排序和前台展示效果。"),
    ],
    "游戏类型设置": [
        ("页面功能", "用于维护游戏类型分类。"),
        ("主要操作", "新增、编辑、排序、启用或停用游戏类型。"),
        ("使用说明", "类型配置会影响前台游戏筛选和后台统计归类。"),
    ],
    "游戏频道管理": [
        ("页面功能", "用于维护游戏频道和频道内游戏展示。"),
        ("主要操作", "配置频道名称、排序、展示游戏和启停状态。"),
        ("使用说明", "保存后需检查前台频道展示和游戏入口是否正确。"),
    ],
    "子游戏列表": [
        ("页面功能", "用于查看和管理子游戏信息。"),
        ("主要操作", "按厂商、类型或状态筛选，启用、停用或调整游戏展示。"),
        ("使用说明", "调整游戏状态会影响玩家进入游戏，操作前需确认影响范围。"),
    ],
    "游戏统计": [
        ("页面功能", "用于查看游戏投注和输赢统计。"),
        ("主要操作", "按游戏、厂商或时间筛选，查看投注人数、投注额和输赢。"),
        ("使用说明", "用于分析游戏表现，异常时结合游戏记录核对。"),
    ],
    "游戏厂商统计": [
        ("页面功能", "用于查看各游戏厂商数据表现。"),
        ("主要操作", "按厂商和时间筛选，查看投注、输赢和活跃数据。"),
        ("使用说明", "用于评估厂商质量和游戏引入效果。"),
    ],
    "在线玩家列表": [
        ("页面功能", "用于查看当前在线玩家。"),
        ("主要操作", "按会员、渠道或状态筛选，查看在线时间、终端和会员信息。"),
        ("使用说明", "用于实时运营观察和异常在线排查。"),
    ],
    "所有会员": [
        ("页面功能", "用于查询和管理会员基础信息。"),
        ("主要操作", "按会员账号、ID、状态或注册时间筛选，查看资料、余额和行为信息。"),
        ("使用说明", "处理会员问题时先确认账号身份，再进入资金、订单或活动明细核对。"),
    ],
    "VIP设置": [
        ("页面功能", "用于配置 VIP 等级、权益和升级条件。"),
        ("主要操作", "设置等级名称、升级条件、保级条件、奖励和启停状态。"),
        ("使用说明", "调整后会影响会员等级和权益展示，需提前确认规则。"),
    ],
    "有奖反馈": [
        ("页面功能", "用于查看和处理会员反馈奖励。"),
        ("主要操作", "筛选反馈记录，查看内容、处理状态并发放或拒绝奖励。"),
        ("使用说明", "处理前核对反馈内容、奖励规则和会员信息。"),
    ],
    "代理列表": [
        ("页面功能", "用于查询和管理代理账号。"),
        ("主要操作", "按代理账号、状态或时间筛选，查看代理信息和下级数据。"),
        ("使用说明", "调整代理状态前需确认结算、会员归属和业务影响。"),
    ],
    "代理配置": [
        ("页面功能", "用于维护代理返佣、返水和结算规则。"),
        ("主要操作", "设置代理等级、比例、结算周期和启停状态。"),
        ("使用说明", "配置变更会影响代理收益，保存前需复核规则。"),
    ],
    "代理领取记录": [
        ("页面功能", "用于查看代理奖励或佣金领取记录。"),
        ("主要操作", "按代理、状态或时间筛选，查看领取金额和处理结果。"),
        ("使用说明", "排查收益问题时结合结算明细和账变记录核对。"),
    ],
    "代理数据查询": [
        ("页面功能", "用于查询代理维度的运营数据。"),
        ("主要操作", "按代理和时间筛选，查看下级会员、充值、投注和佣金数据。"),
        ("使用说明", "用于代理运营分析和结算核对。"),
    ],
    "黑名单": [
        ("页面功能", "用于管理受限制的玩家或风险对象。"),
        ("主要操作", "新增、查询、解除黑名单，查看限制原因和操作记录。"),
        ("使用说明", "加入或解除黑名单前需确认风险原因和处理依据。"),
    ],
    "刷子监控": [
        ("页面功能", "用于监控疑似异常套利或刷量玩家。"),
        ("主要操作", "按玩家、风险类型或时间筛选，查看异常指标和处理记录。"),
        ("使用说明", "处理前结合投注、活动、充值和设备信息综合判断。"),
    ],
    "游戏获利监控": [
        ("页面功能", "用于监控玩家游戏获利异常情况。"),
        ("主要操作", "按玩家、游戏或时间筛选，查看获利、投注和风险状态。"),
        ("使用说明", "发现异常后结合游戏记录和风控规则进一步核查。"),
    ],
    "商户信息": [
        ("页面功能", "用于查看商户基础资料和账户状态。"),
        ("主要操作", "查看商户名称、账号信息、联系方式和状态配置。"),
        ("使用说明", "资料变更需与平台或管理员确认，避免影响后台使用。"),
    ],
    "商户充值": [
        ("页面功能", "用于商户账户充值。"),
        ("主要操作", "选择充值方式，填写金额并提交充值申请。"),
        ("使用说明", "提交前确认金额、收款信息和到账规则。"),
    ],
    "充值记录": [
        ("页面功能", "用于查询商户充值记录。"),
        ("主要操作", "按订单号、状态或时间筛选，查看充值金额和到账状态。"),
        ("使用说明", "到账异常时结合支付凭证和平台处理结果核对。"),
    ],
    "商户账变记录": [
        ("页面功能", "用于查看商户账户资金变动。"),
        ("主要操作", "按账变类型、时间或订单号筛选，查看金额变化和余额。"),
        ("使用说明", "排查商户资金问题时结合充值、账单和扣费记录核对。"),
    ],
    "商户账单": [
        ("页面功能", "用于查看商户账单和费用结算情况。"),
        ("主要操作", "按账期或状态筛选，查看账单金额、费用项目和结算状态。"),
        ("使用说明", "对账时需结合账变记录、充值记录和平台结算规则。"),
    ],
    "埋点统计": [
        ("页面功能", "用于查看站点埋点事件统计。"),
        ("主要操作", "按事件、页面、渠道或时间筛选，查看触发次数和转化数据。"),
        ("使用说明", "用于分析页面行为和推广效果，需确认埋点配置已生效。"),
    ],
    "会话列表": [
        ("页面功能", "用于查看用户访问或行为会话。"),
        ("主要操作", "按用户、会话 ID、渠道或时间筛选，查看会话轨迹。"),
        ("使用说明", "排查转化或行为问题时，可结合事件流水一起分析。"),
    ],
    "事件流水": [
        ("页面功能", "用于查看埋点事件明细流水。"),
        ("主要操作", "按事件、用户、页面或时间筛选，查看单次事件详情。"),
        ("使用说明", "用于排查埋点触发、参数上报和用户行为路径。"),
    ],
    "账号管理": [
        ("页面功能", "用于管理后台账号和角色归属。"),
        ("主要操作", "新增、编辑、启用、停用账号，分配系统预设角色。"),
        ("使用说明", "账号权限影响后台可见菜单和操作范围，调整前需确认职责。"),
    ],
    "管理员登录日志": [
        ("页面功能", "用于查看后台管理员登录记录。"),
        ("主要操作", "按账号、IP、时间或状态筛选，查看登录结果。"),
        ("使用说明", "排查账号安全或异常登录时优先查看本页面。"),
    ],
    "管理员操作日志": [
        ("页面功能", "用于查看后台关键操作记录。"),
        ("主要操作", "按账号、模块、操作类型或时间筛选，查看操作详情。"),
        ("使用说明", "追溯误操作时结合操作时间、操作人和业务页面记录核对。"),
    ],
    "系统异常日志": [
        ("页面功能", "用于查看后台系统异常和接口错误记录。"),
        ("主要操作", "按时间、模块或异常类型筛选，查看错误信息和处理状态。"),
        ("使用说明", "用于定位系统问题，必要时将异常信息提供给技术人员排查。"),
    ],
}


SPECIAL_SECTIONS = {
    "平台公告": {
        "image": PLATFORM_IMAGE,
        "caption": "图：平台公告列表，用于查看公告通知和站内消息。",
        "intro": "平台公告用于查看平台发布的公告、通知和站内消息。",
        "items": [
            ("类型", "按平台公告、通知、短信、风控预警、站内信等类型筛选。"),
            ("标题", "输入标题关键字查询消息。"),
            ("状态", "按未读、已读筛选消息。"),
            ("搜索/重置", "搜索按条件刷新列表；重置清空筛选。"),
            ("全部已读", "将当前未读消息批量标记为已读。"),
            ("列表", "展示类型、标题、时间、状态和操作入口。"),
            ("查看", "打开消息详情，查看完整标题和内容。"),
        ],
    },
    "消息推送": {
        "image": PUSH_IMAGE,
        "caption": "图：消息推送列表，用于管理消息、公告和跑马灯。",
        "intro": "消息推送用于创建和管理面向玩家展示的消息、公告和跑马灯。",
        "items": [
            ("消息类型", "按消息、公告、跑马灯筛选记录。"),
            ("状态", "按未开始、已发送、已结束、已撤回筛选。"),
            ("标题", "输入标题关键字查询推送内容。"),
            ("搜索/重置", "搜索按条件刷新列表；重置清空筛选。"),
            ("新增", "打开新增弹窗，配置消息类型、弹窗频次、收件人、标题、显示内容、发送时间和结束时间。"),
            ("列表", "展示消息编号、标题、消息类型、收件人、发送时间、结束时间、状态、操作人和操作时间。"),
            ("详情", "查看单条推送的完整配置和内容。"),
            ("撤回", "对已发送或需下线的推送执行撤回。"),
        ],
    },
}


def set_run(run, *, size: float = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def insert_after(paragraph, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    return new_para


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_body_after(anchor, text: str):
    paragraph = insert_after(anchor, "Normal")
    run = paragraph.add_run(text)
    set_run(run)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullet_after(anchor, label: str, text: str):
    paragraph = insert_after(anchor, "List Bullet")
    label_run = paragraph.add_run(label)
    set_run(label_run, bold=True, color=BLUE)
    body_run = paragraph.add_run("：" + text)
    set_run(body_run)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_picture_after(anchor, path: Path):
    paragraph = insert_after(anchor, "Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(6.5))
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def add_caption_after(anchor, text: str):
    paragraph = insert_after(anchor, "Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run(run, size=9, color=CAPTION_GRAY)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def section_bounds(document: Document, title: str) -> tuple[int, int]:
    start = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == title:
            start = index
            break
    if start is None:
        raise RuntimeError(f"未找到章节：{title}")
    end = len(document.paragraphs)
    for index in range(start + 1, len(document.paragraphs)):
        if document.paragraphs[index].style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            end = index
            break
    return start, end


def clear_section_at(document: Document, start: int, end: int):
    for paragraph in list(document.paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)
    heading = document.paragraphs[start]
    for run in heading.runs:
        set_run(run, size=12.5, bold=True, color=BLUE)
    return heading


def clear_section(document: Document, title: str):
    start, end = section_bounds(document, title)
    return clear_section_at(document, start, end)


def is_placeholder(paragraphs) -> bool:
    texts = [paragraph.text.strip() for paragraph in paragraphs if paragraph.text.strip()]
    return (
        len(texts) == 3
        and texts[0].startswith("页面介绍：用于处理")
        and texts[1].startswith("常用操作：进入页面后")
        and texts[2].startswith("后续补充：")
    )


def is_legacy_summary(paragraphs) -> bool:
    texts = [paragraph.text.strip() for paragraph in paragraphs if paragraph.text.strip()]
    return (
        len(texts) == 3
        and texts[0].startswith("页面功能：")
        and texts[1].startswith("主要操作：")
        and texts[2].startswith("使用说明：")
    )


def has_image(paragraphs) -> bool:
    return any(run._element.xpath(".//w:drawing") for paragraph in paragraphs for run in paragraph.runs)


def feature_label_from_title(title: str) -> str:
    if title.endswith("配置"):
        return title.replace("配置", "配置")
    if title.endswith("设置"):
        return title.replace("设置", "设置")
    if title.endswith("列表"):
        return title.replace("列表", "列表")
    if title.endswith("记录"):
        return title.replace("记录", "记录")
    if title.endswith("报表"):
        return title.replace("报表", "报表")
    if title.endswith("统计"):
        return title.replace("统计", "统计")
    if title.endswith("查询"):
        return title.replace("查询", "查询")
    if title.endswith("管理"):
        return title.replace("管理", "管理")
    if title.endswith("日志"):
        return title.replace("日志", "日志")
    if title.endswith("账单"):
        return title.replace("账单", "账单")
    if title.endswith("信息"):
        return title.replace("信息", "信息")
    return title


def operation_label_from_text(text: str) -> str:
    if "筛选" in text and "查看" in text:
        return "筛选与查看"
    if "筛选" in text:
        return "筛选条件"
    if any(word in text for word in ["新增", "编辑", "启用", "停用", "删除", "预览", "撤回", "复制"]):
        return "操作入口"
    if any(word in text for word in ["设置", "配置", "维护"]):
        return "配置内容"
    if "查看" in text and "详情" in text:
        return "详情查看"
    if "查看" in text:
        return "查看内容"
    return "操作说明"


def check_label_from_text(text: str) -> str:
    if text.startswith("保存前"):
        return "保存前检查"
    if text.startswith("发布前"):
        return "发布前检查"
    if text.startswith("上线前"):
        return "上线前检查"
    if text.startswith("投放前"):
        return "投放前检查"
    if text.startswith("提交前"):
        return "提交前检查"
    if "对账" in text:
        return "对账要点"
    if "核对" in text:
        return "核对要点"
    if "影响" in text:
        return "影响范围"
    if "分析" in text:
        return "分析要点"
    if "排查" in text:
        return "排查要点"
    return "使用要点"


def add_special(document: Document, title: str) -> None:
    data = SPECIAL_SECTIONS[title]
    image = data["image"]
    if not image.exists():
        raise FileNotFoundError(image)
    last = clear_section(document, title)
    last = add_body_after(last, data["intro"])
    last = add_picture_after(last, image)
    last = add_caption_after(last, data["caption"])
    for label, text in data["items"]:
        last = add_bullet_after(last, label, text)


def add_summary_at(document: Document, start: int, end: int, title: str) -> None:
    last = clear_section_at(document, start, end)
    items = PAGE_SUMMARIES[title]
    if not items:
        return
    generated = [
        (feature_label_from_title(title), items[0][1]),
        (operation_label_from_text(items[1][1]), items[1][1]),
        (check_label_from_text(items[2][1]), items[2][1]),
    ]
    for label, text in generated:
        last = add_bullet_after(last, label, text)


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{DOCX.stem}.simplify-pages-{stamp}.bak.docx"
    shutil.copy2(DOCX, backup)

    document = Document(DOCX)
    changed = []

    for title in SPECIAL_SECTIONS:
        add_special(document, title)
        changed.append(title)

    index = 0
    while index < len(document.paragraphs):
        paragraph = document.paragraphs[index]
        if paragraph.style.name != "Heading 3":
            index += 1
            continue

        title = paragraph.text.strip()
        if title in SPECIAL_SECTIONS or title not in PAGE_SUMMARIES:
            index += 1
            continue

        end = len(document.paragraphs)
        for cursor in range(index + 1, len(document.paragraphs)):
            if document.paragraphs[cursor].style.name in {"Heading 1", "Heading 2", "Heading 3"}:
                end = cursor
                break

        section = document.paragraphs[index + 1 : end]
        if not has_image(section):
            add_summary_at(document, index, end, title)
            changed.append(title)

        index += 1

    try:
        document.save(DOCX)
    except PermissionError:
        print(f"PERMISSION_ERROR backup={backup}")
        raise

    print(f"simplified {len(changed)} sections; backup={backup}")
    print(",".join(changed))


if __name__ == "__main__":
    main()
