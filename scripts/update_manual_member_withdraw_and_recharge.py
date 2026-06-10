from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = next(ROOT.glob("*.docx"))
BACKUP_DIR = ROOT / "backups"

FONT_NAME = "Microsoft YaHei"
BLUE = RGBColor(47, 85, 151)


SECTIONS = {
    "会员提现设置": [
        (
            "页面说明",
            "会员提现设置用于维护会员提现规则、出款限制与审核相关参数，图1展示主表，图2展示修改弹窗，后续图片对应主表按钮打开的说明弹窗或配置弹窗。",
        ),
        (
            "图1主表",
            "主表用于集中查看当前提现规则的配置结果，通常可直接核对站点、币种、提现方式、限额区间、手续费、状态和最后更新时间等信息，并通过操作列进入修改或查看关联弹窗。",
        ),
        (
            "筛选与列表",
            "如主表顶部提供站点、状态或提现方式等筛选项，可先按条件缩小范围后再执行搜索；列表字段应重点核对单笔最小/最大提现金额、每日提现次数、手续费规则、稽核条件和启停状态是否符合当前出款口径。",
        ),
        (
            "图2修改弹窗",
            "修改弹窗用于维护当前提现规则的核心参数，按页面从上到下依次填写或复核提现金额范围、每日次数、手续费收取方式、稽核倍数、审核条件、到账限制、备注和状态开关，确认无误后再保存。",
        ),
        (
            "金额与次数",
            "调整金额区间或次数限制时，应同时考虑前台提现体验、财务出款效率与风险控制要求；如存在单笔上限、单日上限、免费提现次数或超次收费规则，应按弹窗中的字段分别维护，避免只改其中一项导致口径不一致。",
        ),
        (
            "手续费与稽核",
            "涉及手续费、行政费、稽核倍数或流水校验的字段时，应明确收费触发条件、扣费方式和未达条件时的处理结果；保存前建议结合财务中心的提现审核和打码管理口径再次复核。",
        ),
        (
            "按钮弹窗",
            "图2之后的按钮弹窗用于说明某个规则项的详细配置、查看条件说明或执行确认操作；进入此类弹窗后，应重点确认弹窗标题、当前规则归属、字段是否只影响当前配置以及确认按钮提交后的生效范围。",
        ),
        (
            "使用要点",
            "提现规则变更会直接影响会员申请提现、风控审核与财务出款结果，调整后应同步检查提现订单列表中的审核口径是否一致；对于涉及状态停用、限额收紧或收费上调的修改，建议在业务低峰期处理并保留变更说明。",
        ),
    ],
    "会员充值配置": [
        (
            "页面说明",
            "会员充值配置用于维护会员充值入口、金额档位和前台展示规则，图1展示主表，图2展示修改弹窗，后续图片对应主表按钮打开的说明弹窗或配置弹窗。",
        ),
        (
            "图1主表",
            "主表用于集中查看当前充值配置的启用情况，通常可核对站点、币种、充值方式、金额档位、赠送或优惠提示、状态和更新时间等信息，并通过操作列进入修改或查看详情。",
        ),
        (
            "筛选与列表",
            "如页面提供站点、充值方式、状态或币种等筛选条件，可先搜索再维护目标配置；列表字段应重点检查充值最小/最大金额、快捷金额档位、默认推荐金额、文案提示和状态是否与实际支付通道能力一致。",
        ),
        (
            "图2修改弹窗",
            "修改弹窗用于维护充值配置的核心内容，按页面顺序设置充值金额范围、快捷金额选项、赠送说明、温馨提示、展示顺序、启停状态及备注信息，保存前应逐项确认前台是否可读、是否便于会员选择。",
        ),
        (
            "金额档位",
            "快捷金额或固定充值档位应覆盖常用充值场景，金额之间建议保持合理梯度；如支持自定义输入金额，应同时确认最小值、最大值和超限提示文案，避免会员提交后因金额不符被拦截。",
        ),
        (
            "展示与提示",
            "涉及前台展示文案、优惠提示、赠送说明或推荐标签时，应确保内容与当前活动规则、支付方式和站点策略一致；若某配置只对部分站点或部分充值方式生效，应在对应字段中明确限定范围。",
        ),
        (
            "按钮弹窗",
            "图2之后的按钮弹窗通常用于补充说明金额档位、查看配置明细、确认启停或处理关联设置；进入后应重点核对当前配置对象、操作影响范围以及确认提交后是否立即在前台生效。",
        ),
        (
            "使用要点",
            "充值配置调整后，应同步检查支付通道配置、充值订单列表和前台充值页的展示效果；对于停用充值方式、修改金额上限或更新优惠提示的操作，建议先确认对应支付通道可用，再执行保存。",
        ),
    ],
}


def set_run(run, *, size: float = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def insert_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_bullet_after(anchor, label: str, text: str):
    paragraph = insert_after(anchor, style="List Bullet")
    label_run = paragraph.add_run(label)
    set_run(label_run, size=10, bold=True, color=BLUE)
    body_run = paragraph.add_run("：" + text)
    set_run(body_run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def reset_section(document: Document, title: str):
    start = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == title:
            start = index
            break
    if start is None:
        raise RuntimeError(f"未找到章节：{title}")

    end = len(document.paragraphs)
    for index in range(start + 1, len(document.paragraphs)):
        if document.paragraphs[index].style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            end = index
            break

    for paragraph in list(document.paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)

    anchor = document.paragraphs[start]
    for run in anchor.runs:
        set_run(run, size=12.5, bold=True, color=BLUE)
    return anchor


def main() -> None:
    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{DOCX.stem}.member-withdraw-recharge-{stamp}.bak.docx"
    shutil.copy2(DOCX, backup)

    document = Document(DOCX)
    for title, items in SECTIONS.items():
        anchor = reset_section(document, title)
        last = anchor
        for label, text in items:
            last = add_bullet_after(last, label, text)

    document.save(DOCX)
    print(f"updated manual sections; backup={backup}")


if __name__ == "__main__":
    main()
