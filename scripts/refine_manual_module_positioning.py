from __future__ import annotations

from datetime import datetime
import shutil
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = next(ROOT.glob("*.docx"))
BACKUP_DIR = ROOT / "backups"

MODULE_POSITIONING = {
    "四、站点配置": "站点配置是站点基础参数与外围接入配置的入口，左侧菜单下的模板、站点、品牌、域名、客服、分享、第三方登录、支付通道、广告埋点和音乐等页面都归在这里；需要查站点展示、登录接入、支付或埋点配置时，先从这一章定位对应菜单。",
    "五、运营中心": "运营中心聚合站点维护、公告、建议、投诉、消息、宣传、客服、任务和落地页等运营处理页面；要找站点状态、活动通知、用户反馈或客服配置时，先从这一章进入对应功能页。",
    "六、优惠活动": "优惠活动集中管理活动列表、转盘、票券、活动报表和全局配置，主要用于查活动、改玩法、核票券和查看活动效果；按左侧菜单找到对应活动页后，再进入具体配置项。",
    "七、财务中心": "财务中心是提现、充值、通道、人工加扣款、账变和打码规则的资金入口；处理出入款审核、资金流水核对或通道查看时，先在这一章按菜单名找到对应页面。",
    "八、数据报表": "数据报表按经营、游戏、充值、活动、任务和VIP等口径汇总统计，是看日报和下钻明细的入口；需要核对数据时，先进入对应报表，再按时间、站点或维度筛选。",
    "九、游戏中心": "游戏中心集中管理游戏类型、频道、子游戏、统计和厂商数据，主要用于维护大厅展示和查看游戏表现；要找某类游戏配置或厂商报表时，先从这一章进入对应页面。",
    "十、用户管理": "用户管理聚合在线玩家、会员列表、VIP设置和反馈处理，是查会员状态和跟进用户问题的入口；按左侧菜单找到目标页面后，可继续查看在线、资料或反馈信息。",
    "十一、代理中心": "代理中心用于查看代理列表、代理配置、领取记录和代理数据，是跟进代理层级、佣金和结算情况的入口；需要定位某个代理或核对代理业绩时，先从这一章进入对应页面。",
    "十二、风险中心": "风险中心集中查看黑名单、刷子监控和游戏获利监控，是排查异常账号、异常行为和套利风险的入口；发现风险线索后，先从这里找到对应监控页再继续下钻。",
    "十三、商户中心": "商户中心是商户资料、充值、账变和账单的总入口，适合核对主体信息、资金流水和结算结果；涉及商户基础资料或费用问题时，先从这一章定位到具体页面。",
    "十四、埋点配置": "埋点配置用于查看埋点统计、会话列表和事件流水，是核对事件上报、访问路径和触发结果的入口；需要排查页面埋点时，先从这里进入对应日志或统计页。",
    "十五、人事中心": "人事中心用于管理后台账号、登录日志、操作日志和异常日志，是账号权限、审计追踪和异常排查的入口；查账号归属或操作记录时，先从这一章进入对应页面。",
}


def backup(docx_path: Path) -> Path:
    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    target = BACKUP_DIR / f"{docx_path.stem}.module-positioning-{stamp}.bak{docx_path.suffix}"
    shutil.copy2(docx_path, target)
    return target


def main() -> None:
    doc = Document(DOCX_PATH)
    current_h1 = None
    changed = []

    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        style = paragraph.style.name if paragraph.style else ""

        if style == "Heading 1":
            current_h1 = text
            continue

        if style == "Heading 2" and text == "模块定位":
            replacement = MODULE_POSITIONING.get(current_h1 or "")
            if replacement is None:
                raise RuntimeError(f"missing module positioning text for chapter: {current_h1}")

            target_index = index + 1
            while target_index < len(doc.paragraphs) and not doc.paragraphs[target_index].text.strip():
                target_index += 1
            if target_index >= len(doc.paragraphs):
                raise RuntimeError(f"missing body paragraph after module positioning heading: {current_h1}")

            doc.paragraphs[target_index].text = replacement
            changed.append((current_h1, target_index))

    if len(changed) != len(MODULE_POSITIONING):
        missing = sorted(set(MODULE_POSITIONING) - {chapter for chapter, _ in changed})
        raise RuntimeError(f"unexpected replacement count: {len(changed)}; missing={missing}")

    backup_path = backup(DOCX_PATH)
    doc.save(DOCX_PATH)
    print(f"updated={len(changed)}")
    print(f"backup={backup_path}")


if __name__ == "__main__":
    main()
