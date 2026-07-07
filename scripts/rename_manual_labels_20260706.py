from copy import deepcopy
from datetime import datetime
from pathlib import Path
import shutil

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "B端后台操作手册.docx"


def s(*codes):
    return "".join(chr(code) for code in codes)


REPLACEMENTS = [
    (s(0x6D88, 0x606F, 0x63A8, 0x9001), s(0x6D88, 0x606F, 0x516C, 0x544A)),
    (s(0x9ED8, 0x8BA4, 0x770B, 0x677F), s(0x9996, 0x9875)),
]


def iter_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)
    for section in document.sections:
        for story in [section.header, section.footer]:
            for paragraph in story.paragraphs:
                yield paragraph
            for table in story.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from iter_paragraphs(cell)


def replace_in_paragraph(paragraph):
    before = paragraph.text
    changed = False
    for old, new in REPLACEMENTS:
        for run in paragraph.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                changed = True

    # Very short labels are normally contained in a single run. This fallback
    # handles the rare split-run case while keeping the paragraph style.
    if any(old in paragraph.text for old, _ in REPLACEMENTS):
        final_text = paragraph.text
        for old, new in REPLACEMENTS:
            final_text = final_text.replace(old, new)
        if final_text != paragraph.text:
            first_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            template = deepcopy(first_run._r.rPr) if first_run._r.rPr is not None else None
            for run in list(paragraph.runs):
                run._element.getparent().remove(run._element)
            run = paragraph.add_run(final_text)
            if template is not None:
                run._r.insert(0, template)
            changed = True

    return changed, before, paragraph.text


def count_text(document):
    counts = {old: 0 for old, _ in REPLACEMENTS}
    for paragraph in iter_paragraphs(document):
        text = paragraph.text
        for old, _ in REPLACEMENTS:
            counts[old] += text.count(old)
    return counts


def main():
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    backup = DOCX.with_name(f"B端后台操作手册.{datetime.now():%Y%m%d-%H%M%S}.bak.docx")
    shutil.copy2(DOCX, backup)

    document = Document(DOCX)
    before_counts = count_text(document)
    changed_paragraphs = 0
    examples = []

    for paragraph in iter_paragraphs(document):
        changed, before, after = replace_in_paragraph(paragraph)
        if changed:
            changed_paragraphs += 1
            if len(examples) < 8:
                examples.append((before, after))

    after_counts = count_text(document)
    document.save(DOCX)

    print(f"backup: {backup.name}")
    print(f"changed_paragraphs: {changed_paragraphs}")
    print("before_counts:")
    for old, count in before_counts.items():
        print(f"  {old}: {count}")
    print("after_counts:")
    for old, count in after_counts.items():
        print(f"  {old}: {count}")
    print("examples:")
    for before, after in examples:
        print(f"  {before} -> {after}")


if __name__ == "__main__":
    main()
