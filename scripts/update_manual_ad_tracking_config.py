from datetime import datetime
from pathlib import Path
import shutil

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"
ASSET_DIR = ROOT / "custom" / "assets" / "manual-ad-tracking-config"
TITLE = "广告埋点配置"


SECTIONS = [
    (
        "页面说明",
        {
            "path": ASSET_DIR / "ad-tracking-config-list.png",
            "caption": "图：广告埋点配置列表，用于维护不同渠道的落地页、跳转地址和事件标识。",
            "width": 6.5,
        },
        [
            "广告埋点配置用于维护站点推广渠道的埋点参数，支持按渠道类型查询配置，并在列表中查看渠道编码、渠道名称、落地页域名、跳转地址、像素ID、安装类型、事件标识和创建时间等信息。",
            "该页面主要服务于广告投放、渠道归因和转化事件回传。新增或修改配置后，前台落地页和推广链接可能会按配置携带对应渠道参数，运营人员应在上线前核对渠道、域名、事件标识和跳转地址是否一致。",
        ],
    ),
    (
        "筛选与列表字段",
        None,
        [
            "渠道类型：用于按 facebook、adjust、google 等渠道类型筛选列表。需要排查某个渠道数据时，先选择渠道类型再点击“搜索”；点击“重置”可恢复默认查询条件。",
            "渠道ID：系统生成的配置唯一标识，通常用于技术排查或接口核对。列表中较长内容会省略显示，核对完整值时可进入修改弹窗或结合后台数据记录确认。",
            "渠道编码：渠道的业务编码，用于区分不同推广渠道或投放批次。新增配置后应确认编码与投放侧、落地页侧约定一致。",
            "渠道名称：展示渠道配置的名称，例如 Google、Facebook测试、adjust测试等，建议使用便于运营识别的渠道或活动名称。",
            "渠道类型：展示当前配置所属广告平台或归因类型。不同渠道类型会影响后续需要填写的像素ID、安装类型和事件标识字段。",
            "落地页域名：配置该渠道使用的落地页域名。保存前应确认域名已在站点或推广域名中配置完成，并可正常访问。",
            "跳转地址：用户从推广入口进入后需要跳转的地址，可包含渠道参数。地址较长时列表会省略显示，保存前应重点检查域名、路径和参数是否正确。",
            "像素ID：用于 Facebook、Google 等像素或广告平台识别的跟踪 ID。为空或填写错误时，可能导致注册、充值等事件无法被广告平台正确接收。",
            "安装类型：用于 Adjust 等归因场景区分 H5、App 或其他安装/访问类型。配置时应与实际推广入口和客户端形态一致。",
            "事件标识：包括注册、首充、复充等事件标识，用于向归因或广告平台回传关键转化事件。填写前应按投放平台要求核对事件名称、大小写和参数口径。",
            "备注：记录配置用途、投放批次、测试说明或特殊约定，便于后续排查和交接。",
            "创建时间：展示配置创建时间，用于判断配置是否为当前投放周期内新增或历史遗留配置。",
            "操作：提供“修改”和“删除”。修改用于调整渠道配置；删除用于移除不再使用的配置，删除前应确认该渠道已停止投放或已有替代配置。",
        ],
    ),
    (
        "新增与编辑配置",
        None,
        [
            "新增配置：点击“新增”打开新增埋点配置弹窗，按页面提示选择或填写渠道名称、落地页域名、跳转地址和渠道类型。带必填标识的字段必须补齐后再提交。",
            "渠道类型切换：弹窗中切换渠道类型后，表单会展示对应的配置项。Facebook、Google 等渠道通常重点填写像素ID；Adjust 类配置通常需要维护安装类型以及注册、首充、复充事件标识。",
            "编辑配置：在列表操作列点击“修改”，进入编辑埋点配置弹窗后可调整域名、跳转地址、像素ID、安装类型、事件标识和备注等信息。提交后应回到列表核对展示值是否更新。",
            "删除配置：点击“删除”前应确认该配置不再被当前推广链接、落地页或广告计划使用。删除后相关渠道可能无法继续完成归因或事件回传，必要时先停用投放或保留历史配置。",
        ],
    ),
    (
        "注意事项",
        None,
        [
            "落地页域名、跳转地址和广告平台参数需要保持一致，任一环节配置错误都可能造成用户无法进入正确页面、渠道无法归因或广告平台收不到转化事件。",
            "涉及像素ID和事件标识的调整应由运营、投放或技术人员共同复核。修改线上渠道前，建议先复制当前值留档，确认新配置验证通过后再用于正式投放。",
            "若广告平台未收到注册、首充或复充事件，优先检查渠道类型、像素ID、事件标识、安装类型、跳转地址参数、落地页域名访问状态，以及前台页面是否已经加载对应埋点脚本。",
            "同一渠道存在多条配置时，应通过渠道编码、渠道名称、备注和创建时间区分投放批次，避免误改正在使用的渠道配置。",
        ],
    ),
]


def font(size, bold=False):
    candidates = [
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for item in candidates:
        path = Path(item)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def text(draw, xy, value, fill="#2f3a4a", size=20, bold=False, anchor=None):
    draw.text(xy, value, fill=fill, font=font(size, bold), anchor=anchor)


def rounded(draw, box, fill, outline=None, radius=8, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def make_overview_image():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / "ad-tracking-config-list.png"
    img = Image.new("RGB", (1680, 760), "#f4f6f9")
    draw = ImageDraw.Draw(img)

    rounded(draw, (18, 18, 1656, 706), "#ffffff", "#e2e7ef", 6)
    text(draw, (38, 46), "渠道类型:", size=18)
    rounded(draw, (118, 34, 250, 72), "#ffffff", "#d6dde8", 4)
    text(draw, (132, 45), "请选择渠道类型", fill="#9aa5b5", size=16)
    rounded(draw, (266, 34, 354, 72), "#3290ff", "#3290ff", 4)
    text(draw, (310, 53), "搜索", fill="#ffffff", size=17, anchor="mm")
    rounded(draw, (370, 34, 458, 72), "#ffffff", "#d6dde8", 4)
    text(draw, (414, 53), "重置", fill="#4b5563", size=17, anchor="mm")
    rounded(draw, (476, 34, 564, 72), "#3290ff", "#3290ff", 4)
    text(draw, (520, 53), "新增", fill="#ffffff", size=17, anchor="mm")

    x0, y0 = 32, 92
    widths = [120, 110, 110, 110, 150, 170, 100, 100, 132, 132, 132, 105, 160, 92]
    headers = [
        "渠道ID", "渠道编码", "渠道名称", "渠道类型", "落地页域名", "跳转地址", "像素ID", "安装类型",
        "事件标识(注册)", "事件标识(首充)", "事件标识(复充)", "备注", "创建时间", "操作",
    ]
    rows = [
        ["da4a2bfb-d85...", "1111", "谷歌", "google", "jyowhite.cc", "jyowhite.cc?channel...", "1111", "-", "-", "-", "-", "", "2026-06-06 17:19:58", "修改  删除"],
        ["5a893607-b8c...", "123456", "苟富贵", "adjust", "ay.jyowhite.cc", "ay.jyowhite.cc?chan...", "", "App", "-", "-", "-", "", "2026-06-06 17:01:34", "修改  删除"],
        ["f06526c9-6c1c...", "101104", "Google", "google", "ae.jyowhite.cc", "ae.jyowhite.cc?chann...", "1233", "-", "-", "-", "-", "", "2026-06-06 16:58:45", "修改  删除"],
        ["aed1a170-8f6c...", "101103", "ces", "facebook", "tuiguanawda.jyowhit...", "tuiguanawda.jyowhit...", "123", "-", "-", "-", "-", "222", "2026-05-26 13:24:13", "修改  删除"],
        ["dc96f1cb-298...", "101102", "Facebook测试", "facebook", "tuiguanawda.jyowhit...", "tuiguanawda.jyowhit...", "1112", "-", "-", "-", "-", "测试", "2026-04-23 14:47:03", "修改  删除"],
        ["a6f02773-b97...", "101101", "adjust测试", "adjust", "tuiguanawda.jyowhit...", "tuiguanawda.jyowhit...", "", "H5", "111", "22", "33", "cs", "2026-04-23 13:58:49", "修改  删除"],
    ]

    x = x0
    for idx, w in enumerate(widths):
        draw.rectangle((x, y0, x + w, y0 + 46), fill="#f7f8fa", outline="#e3e7ee")
        text(draw, (x + w / 2, y0 + 24), headers[idx], fill="#374151", size=16, bold=True, anchor="mm")
        x += w

    for r, row in enumerate(rows):
        y = y0 + 46 + r * 48
        x = x0
        for c, w in enumerate(widths):
            draw.rectangle((x, y, x + w, y + 48), fill="#ffffff", outline="#e3e7ee")
            fill = "#2388ff" if c == len(widths) - 1 else "#334155"
            if c == len(widths) - 1:
                text(draw, (x + 34, y + 25), "修改", fill="#2388ff", size=15, anchor="mm")
                text(draw, (x + 68, y + 25), "删除", fill="#f35f68", size=15, anchor="mm")
            else:
                text(draw, (x + w / 2, y + 25), row[c], fill=fill, size=15, anchor="mm")
            x += w

    draw.rectangle((38, 676, 1606, 688), fill="#d1d5db")
    rounded(draw, (38, 676, 1568, 688), "#8c9198", "#8c9198", 6)
    text(draw, (1270, 724), "|‹    ‹‹    ‹", fill="#b0b8c5", size=18)
    rounded(draw, (1362, 710, 1416, 742), "#ffffff", "#d6dde8", 4)
    text(draw, (1389, 726), "1", fill="#334155", size=16, anchor="mm")
    text(draw, (1432, 726), "/ 1     ›    ››    ›|", fill="#8b95a5", size=18, anchor="lm")
    rounded(draw, (1532, 710, 1628, 742), "#ffffff", "#d6dde8", 4)
    text(draw, (1580, 726), "20条/页", fill="#334155", size=15, anchor="mm")
    text(draw, (1640, 726), "共 6 条记录", fill="#334155", size=15, anchor="lm")
    img.save(path)
    return path


def set_run_font(run, size=10, bold=False, color=None):
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def set_heading_style(paragraph):
    for run in paragraph.runs:
        set_run_font(run, 10.5, True, RGBColor(31, 78, 121))


def set_caption_style(paragraph):
    paragraph.alignment = 1
    for run in paragraph.runs:
        set_run_font(run, 9, False, RGBColor(127, 127, 127))


def add_label_text(paragraph, value):
    label, sep, body = value.partition("：")
    if sep:
        run = paragraph.add_run(label + sep)
        set_run_font(run, 10, True, RGBColor(31, 78, 121))
        run = paragraph.add_run(body)
        set_run_font(run, 10)
    else:
        run = paragraph.add_run(value)
        set_run_font(run, 10)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_target_range(document):
    site_config_seen = False
    start = None
    for idx, paragraph in enumerate(document.paragraphs):
        text_value = paragraph.text.strip()
        if paragraph.style.name == "Heading 1" and text_value == "四、站点配置":
            site_config_seen = True
        elif paragraph.style.name == "Heading 1" and site_config_seen:
            break
        elif site_config_seen and paragraph.style.name == "Heading 3" and text_value == TITLE:
            start = idx
            break
    if start is None:
        raise RuntimeError("未在“四、站点配置”下找到广告埋点配置章节")

    end = len(document.paragraphs)
    for idx in range(start + 1, len(document.paragraphs)):
        paragraph = document.paragraphs[idx]
        if paragraph.style.name in ("Heading 1", "Heading 2", "Heading 3") and paragraph.text.strip():
            end = idx
            break
    return start, end


def update_manual():
    make_overview_image()
    BACKUP_DIR.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = BACKUP_DIR / f"B端后台操作手册.ad-tracking-config-{timestamp}.bak.docx"
    shutil.copy2(DOCX_PATH, backup_path)

    document = Document(DOCX_PATH)
    start, end = find_target_range(document)
    anchor = document.paragraphs[start]
    for paragraph in list(document.paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)

    current = anchor
    for heading, image, bullets in SECTIONS:
        current = insert_paragraph_after(current, style="Heading 4")
        current.add_run(heading)
        set_heading_style(current)

        if image:
            if not image["path"].exists():
                raise FileNotFoundError(image["path"])
            current = insert_paragraph_after(current, style="Normal")
            current.alignment = 1
            current.add_run().add_picture(str(image["path"]), width=Inches(image["width"]))
            current = insert_paragraph_after(current, image["caption"], style="Caption")
            set_caption_style(current)

        for bullet in bullets:
            current = insert_paragraph_after(current, style="List Bullet")
            add_label_text(current, bullet)

    try:
        document.save(DOCX_PATH)
    except PermissionError:
        print(f"PERMISSION_ERROR backup={backup_path}")
        raise

    return backup_path


if __name__ == "__main__":
    backup = update_manual()
    print(f"updated ad tracking config; backup={backup}")
