from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


DOCX_PATH = "B端后台操作手册.docx"
FONT_NAME = "微软雅黑"
DEEP_BLUE = RGBColor(37, 64, 97)

OTHER_PAGES = [
    (
        "品牌设置",
        [
            ("页面用途", "维护品牌名称、Logo、视觉信息和品牌展示内容。"),
            ("常用操作", "上传或替换品牌素材；编辑品牌文案；设置展示状态。"),
            ("使用要点", "素材尺寸、清晰度和品牌一致性需要提前确认。"),
        ],
    ),
    (
        "站点域名",
        [
            ("页面用途", "维护站点访问域名及相关状态。"),
            ("常用操作", "新增域名；编辑域名备注；启用或停用域名；查看域名状态。"),
            ("使用要点", "域名变更涉及访问入口，应与技术或运维流程配合。"),
        ],
    ),
    (
        "客服配置",
        [
            ("页面用途", "维护客服入口、客服渠道、在线状态和展示规则。"),
            ("常用操作", "新增客服渠道；编辑链接或账号；调整排序；启停客服入口。"),
            ("使用要点", "客服信息错误会直接影响用户咨询转化。"),
        ],
    ),
    (
        "分享配置",
        [
            ("页面用途", "维护分享标题、描述、图片和渠道参数。"),
            ("常用操作", "编辑分享文案；配置分享图片；设置渠道或落地页参数。"),
            ("使用要点", "分享配置应与活动、品牌或投放页面保持一致。"),
        ],
    ),
    (
        "第三方登录配置",
        [
            ("页面用途", "配置第三方登录方式和接入参数。"),
            ("常用操作", "查看登录渠道；填写应用标识、密钥或回调信息；启用或停用渠道。"),
            ("使用要点", "密钥类信息属于敏感配置，变更前需确认权限。"),
        ],
    ),
    (
        "支付通道配置",
        [
            ("页面用途", "维护充值或提现相关支付通道。"),
            ("常用操作", "新增通道；配置通道参数、费率、限额和适用范围；启停通道；调整排序。"),
            ("使用要点", "支付通道配置会影响真实资金链路，保存前必须复核。"),
        ],
    ),
    (
        "广告埋点配置",
        [
            ("页面用途", "配置广告投放或渠道追踪所需埋点参数。"),
            ("常用操作", "新增广告埋点；维护渠道、事件、参数和状态；查看配置是否生效。"),
            ("使用要点", "参数名称应与投放平台和数据统计口径一致。"),
        ],
    ),
    (
        "音乐管理",
        [
            ("页面用途", "维护站点音乐资源或背景音乐配置。"),
            ("常用操作", "上传或编辑音乐信息；设置启用状态；调整展示或播放规则。"),
            ("使用要点", "注意版权、音量体验和不同终端兼容性。"),
        ],
    ),
]


def set_run_font(run, size=10, bold=False, color=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    p = paragraph._parent.add_paragraph()
    p._p = new_p
    if style:
        p.style = style
    if text:
        p.add_run(text)
    return p


def add_heading_after(anchor, title):
    p = add_after(anchor, style="Heading 3")
    run = p.add_run(title)
    set_run_font(run, size=10.5, bold=True, color=DEEP_BLUE)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_bullet_after(anchor, prefix, rest):
    p = add_after(anchor, style="List Bullet")
    run = p.add_run(prefix)
    set_run_font(run, size=10, bold=True, color=DEEP_BLUE)
    run = p.add_run("：" + rest)
    set_run_font(run, size=10, bold=False)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    return p


doc = Document(DOCX_PATH)

existing_titles = {p.text.strip() for p in doc.paragraphs}
missing = [title for title, _ in OTHER_PAGES if title not in existing_titles]
if not missing:
    print("other site config pages already present")
    raise SystemExit

anchor = next(p for p in doc.paragraphs if p.text.strip() == "操作方式：点击新增创建悬浮框，点击编辑维护按钮项；保存前检查左右位置是否冲突、按钮数量是否过多、移动端是否遮挡主要内容。")
last = anchor
for title, bullets in OTHER_PAGES:
    last = add_heading_after(last, title)
    for prefix, rest in bullets:
        last = add_bullet_after(last, prefix, rest)

doc.save(DOCX_PATH)
print("restored", ", ".join(missing))
