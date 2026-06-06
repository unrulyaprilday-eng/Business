from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"
ASSET_DIR = ROOT / "custom" / "assets" / "manual-customer-service-config"
ONLINE_IMAGE = ASSET_DIR / "customer-service-online.png"
OTHER_IMAGE = ASSET_DIR / "customer-service-other.png"
MODAL_IMAGE = ASSET_DIR / "customer-service-other-edit.png"

FONT_NAME = "微软雅黑"
BLUE = RGBColor(47, 85, 151)
CAPTION_GRAY = RGBColor(112, 126, 146)


def image_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for item in candidates:
        path = Path(item)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def draw_text(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, font, fill=(17, 24, 39)) -> None:
    draw.text(xy, text, font=font, fill=fill)


def text_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill=(17, 24, 39)) -> None:
    left, top, right, bottom = box
    bbox = draw.textbbox((0, 0), text, font=font)
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw_text(draw, (left + (right - left - width) // 2, top + (bottom - top - height) // 2 - 1), text, font, fill)


def rounded(draw: ImageDraw.ImageDraw, box, radius=4, fill=(255, 255, 255), outline=(218, 224, 233), width=1) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def switch(draw: ImageDraw.ImageDraw, x: int, y: int, on: bool = True) -> None:
    fill = (103, 180, 255) if on else (203, 213, 225)
    draw.rounded_rectangle((x, y, x + 38, y + 22), radius=12, fill=fill)
    knob_x = x + 19 if on else x + 3
    draw.ellipse((knob_x, y + 3, knob_x + 16, y + 19), fill=(255, 255, 255))


def im_icon(draw: ImageDraw.ImageDraw, x: int, y: int, im: str, size: int = 28) -> None:
    colors = {
        "Telegram": (53, 170, 222),
        "WhatsApp": (38, 212, 83),
        "Facebook": (24, 119, 242),
        "Line": (0, 185, 0),
        "Online": (36, 155, 211),
    }
    labels = {
        "Telegram": "➤",
        "WhatsApp": "☎",
        "Facebook": "f",
        "Line": "LINE",
        "Online": "Q",
    }
    draw.ellipse((x, y, x + size, y + size), fill=colors.get(im, (53, 170, 222)))
    font = image_font(9 if im == "Line" else int(size * 0.56), bold=True)
    text_center(draw, (x, y, x + size, y + size), labels.get(im, "?"), font, (255, 255, 255))


def draw_tabs(draw: ImageDraw.ImageDraw, active: str, width: int) -> None:
    f = image_font(14)
    blue = (22, 119, 255)
    border = (217, 222, 232)
    draw.line((0, 44, width, 44), fill=border, width=1)
    draw_text(draw, (22, 20), "在线客服", f, blue if active == "online" else (48, 55, 68))
    draw_text(draw, (100, 20), "其他客服", f, blue if active == "other" else (48, 55, 68))
    if active == "online":
        draw.line((18, 44, 74, 44), fill=blue, width=2)
    else:
        draw.line((100, 44, 145, 44), fill=blue, width=2)


def create_online_image() -> None:
    width, height = 884, 440
    img = Image.new("RGB", (width, height), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    f = image_font(14)
    f_small = image_font(12)
    border = (213, 219, 229)
    muted = (139, 149, 165)
    draw_tabs(draw, "online", width)

    draw_text(draw, (47, 79), "在线客服开关", f)
    switch(draw, 146, 74, True)

    draw_text(draw, (74, 154), "客服头像", f)
    draw.rectangle((146, 121, 202, 177), fill=(7, 21, 35))
    im_icon(draw, 151, 123, "Online", 52)
    draw_text(draw, (143, 189), "请上传64*64规格或等比例，PNG、JPG、GIF图片，100KB以内。", f_small, muted)

    draw_text(draw, (74, 234), "前端标题", f)
    rounded(draw, (142, 218, 430, 248), fill=(242, 243, 245), outline=border)
    draw_text(draw, (150, 226), "测试", f, (171, 179, 191))

    draw_text(draw, (74, 299), "前端说明", f)
    rounded(draw, (142, 266, 430, 335), fill=(242, 243, 245), outline=border)
    draw_text(draw, (150, 276), "测试说明", f, (171, 179, 191))

    rounded(draw, (38, 365, 90, 395), fill=(64, 145, 247), outline=(64, 145, 247))
    text_center(draw, (38, 365, 90, 395), "编辑", f, (255, 255, 255))
    img.save(ONLINE_IMAGE)


def create_other_image() -> None:
    width, height = 973, 441
    img = Image.new("RGB", (width, height), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    f = image_font(13)
    f_bold = image_font(13, True)
    blue = (22, 119, 255)
    border = (228, 232, 238)
    head = (248, 250, 252)
    draw_tabs(draw, "other", width)

    rounded(draw, (777, 67, 846, 97), fill=(64, 145, 247), outline=(64, 145, 247))
    text_center(draw, (777, 67, 846, 97), "+ 新增", f, (255, 255, 255))

    x0, y0 = 36, 112
    widths = [44, 63, 90, 369, 73, 64, 108]
    headers = ["#", "排序", "IM", "配置客服", "图标", "开关", "操作"]
    x = x0
    for w, h in zip(widths, headers):
        draw.rectangle((x, y0, x + w, y0 + 42), fill=head, outline=border)
        text_center(draw, (x, y0, x + w, y0 + 42), h, f_bold)
        x += w

    rows = [
        ("1", "1", "Telegram", [("dzh", "https://t.me/h_l_j_hui")], "Telegram"),
        ("2", "2", "WhatsApp", [("11", "11"), ("22", "22")], "WhatsApp"),
        ("3", "3", "Facebook", [("aa", "")], "Facebook"),
        ("4", "4", "Line", [], "Line"),
    ]
    y = y0 + 42
    for row in rows:
        x = x0
        for w in widths:
            draw.rectangle((x, y, x + w, y + 43), fill=(255, 255, 255), outline=border)
            x += w
        text_center(draw, (x0, y, x0 + widths[0], y + 43), row[0], f)
        text_center(draw, (x0 + widths[0], y, x0 + widths[0] + widths[1], y + 43), row[1], f)
        text_center(draw, (x0 + 107, y, x0 + 233, y + 43), row[2], f)
        cx = x0 + sum(widths[:3]) + 8
        if row[3]:
            line_y = y + (8 if len(row[3]) > 1 else 13)
            for name, link in row[3]:
                im_icon(draw, cx, line_y, row[4], 18)
                draw_text(draw, (cx + 25, line_y + 2), name, f)
                if link:
                    draw_text(draw, (cx + 55, line_y + 2), link, f, blue)
                line_y += 20
        else:
            text_center(draw, (x0 + sum(widths[:3]), y, x0 + sum(widths[:4]), y + 43), "-", f)
        ix = x0 + sum(widths[:4]) + 23
        im_icon(draw, ix, y + 10, row[4], 28)
        switch(draw, x0 + sum(widths[:5]) + 14, y + 11, True)
        draw_text(draw, (x0 + sum(widths[:6]) + 25, y + 16), "编辑", f, blue)
        draw_text(draw, (x0 + sum(widths[:6]) + 60, y + 16), "删除", f, (245, 101, 101))
        y += 43
    img.save(OTHER_IMAGE)


def create_modal_image() -> None:
    width, height = 693, 835
    img = Image.new("RGB", (width, height), (240, 240, 240))
    draw = ImageDraw.Draw(img)
    f = image_font(13)
    f_bold = image_font(13, True)
    f_small = image_font(12)
    blue = (22, 119, 255)
    border = (214, 221, 231)
    muted = (139, 149, 165)

    rounded(draw, (4, 6, 687, 829), radius=2, fill=(255, 255, 255), outline=(160, 166, 176))
    draw.line((4, 45, 687, 45), fill=border)
    draw_text(draw, (13, 22), "编辑其他客服", f_bold)
    draw_text(draw, (661, 18), "×", image_font(24), (79, 88, 101))

    draw_text(draw, (51, 79), "排序", f)
    rounded(draw, (92, 66, 637, 97), fill=(255, 255, 255), outline=border)
    draw_text(draw, (99, 77), "2", f)
    draw.rectangle((637, 66, 671, 82), fill=(248, 250, 252), outline=border)
    draw.rectangle((637, 82, 671, 97), fill=(248, 250, 252), outline=border)
    text_center(draw, (637, 66, 671, 82), "+", f)
    text_center(draw, (637, 82, 671, 97), "-", f)

    draw_text(draw, (52, 130), "*", f, (239, 68, 68))
    draw_text(draw, (63, 130), "IM", f)
    rounded(draw, (92, 117, 671, 148), fill=(255, 255, 255), outline=border)
    draw_text(draw, (99, 128), "WhatsApp", f)
    draw_text(draw, (657, 124), "▾", f, (148, 163, 184))

    draw_text(draw, (42, 221), "* 图标", f)
    im_icon(draw, 96, 170, "WhatsApp", 88)
    draw.ellipse((172, 165, 190, 183), fill=(248, 113, 113))
    text_center(draw, (172, 165, 190, 183), "×", f, (255, 255, 255))
    draw_text(draw, (91, 270), "请上传64*64规格或等比例，PNG、JPG、GIF图片，100KB以内。", f_small, muted)

    draw_text(draw, (26, 476), "配置客服", f)
    rounded(draw, (92, 295, 671, 682), fill=(255, 255, 255), outline=border)
    draw.rectangle((93, 296, 670, 327), fill=(248, 250, 252))
    for label, x in [("昵称", 103), ("链接", 250), ("在线时间", 400), ("图标", 548)]:
        draw_text(draw, (x, 310), label, f_bold, (100, 116, 139))

    for top, name, link in [(339, "11", "11"), (516, "22", "22")]:
        rounded(draw, (102, top + 64, 237, top + 94), fill=(255, 255, 255), outline=border)
        draw_text(draw, (110, top + 73), name, f)
        rounded(draw, (250, top + 64, 384, top + 94), fill=(255, 255, 255), outline=border)
        draw_text(draw, (258, top + 73), link, f)
        rounded(draw, (397, top + 64, 531, top + 94), fill=(255, 255, 255), outline=border)
        draw_text(draw, (405, top + 73), "00:00-21:00", f)
        im_icon(draw, 548, top, "WhatsApp", 88)
        draw.ellipse((623, top - 5, 641, top + 13), fill=(248, 113, 113))
        text_center(draw, (623, top - 5, 641, top + 13), "×", f, (255, 255, 255))
        draw_text(draw, (543, top + 97), "请上传64*64规格或", f_small, muted)
        draw_text(draw, (543, top + 116), "等比例，PNG、JPG、", f_small, muted)
        draw_text(draw, (543, top + 135), "GIF图片，100KB", f_small, muted)
        draw_text(draw, (543, top + 154), "以内。", f_small, muted)
        draw_text(draw, (645, top + 103), "删", f_small, (245, 101, 101))

    rounded(draw, (92, 691, 155, 716), fill=(255, 255, 255), outline=border)
    text_center(draw, (92, 691, 155, 716), "+ 添加", f)

    draw_text(draw, (52, 747), "开关", f)
    switch(draw, 97, 740, True)
    rounded(draw, (562, 790, 615, 820), fill=(255, 255, 255), outline=border)
    text_center(draw, (562, 790, 615, 820), "取消", f)
    rounded(draw, (622, 790, 674, 820), fill=(64, 145, 247), outline=(64, 145, 247))
    text_center(draw, (622, 790, 674, 820), "确定", f, (255, 255, 255))
    img.save(MODAL_IMAGE)


def create_images() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    create_online_image()
    create_other_image()
    create_modal_image()


def set_run(run, *, size: float = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def insert_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_heading_after(anchor, text: str, level: int):
    paragraph = insert_after(anchor, style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_run(run, size=11 if level >= 4 else 12.5, bold=True, color=BLUE)
    paragraph.paragraph_format.space_after = Pt(5)
    return paragraph


def add_body_after(anchor, text: str):
    paragraph = insert_after(anchor, style="Normal")
    run = paragraph.add_run(text)
    set_run(run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullet_after(anchor, label: str, text: str):
    paragraph = insert_after(anchor, style="List Bullet")
    label_run = paragraph.add_run(label)
    set_run(label_run, size=10, bold=True, color=BLUE)
    body_run = paragraph.add_run("：" + text)
    set_run(body_run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_picture_after(anchor, image_path: Path, width: float = 6.45):
    paragraph = insert_after(anchor, style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def add_caption_after(anchor, text: str):
    paragraph = insert_after(anchor, style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run(run, size=9, color=CAPTION_GRAY)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def reset_section(document: Document):
    start = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == "客服配置":
            start = index
            break
    if start is None:
        raise RuntimeError("未找到“客服配置”章节。")

    end = len(document.paragraphs)
    for index in range(start + 1, len(document.paragraphs)):
        style_name = document.paragraphs[index].style.name
        if style_name in {"Heading 1", "Heading 2", "Heading 3"}:
            end = index
            break

    for paragraph in list(document.paragraphs[start + 1:end]):
        delete_paragraph(paragraph)

    anchor = document.paragraphs[start]
    for run in anchor.runs:
        set_run(run, size=12.5, bold=True, color=BLUE)
    return anchor


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    create_images()

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{DOCX.stem}.customer-service-config-{stamp}.bak.docx"
    shutil.copy2(DOCX, backup)

    document = Document(DOCX)
    anchor = reset_section(document)
    last = add_body_after(
        anchor,
        "客服配置用于维护前台展示给会员的客服入口。商户可在“在线客服”中配置站内客服入口的头像、标题、说明和启用状态，也可在“其他客服”中维护 Telegram、WhatsApp、Facebook、Line 等第三方 IM 客服渠道。修改前应先确认当前维护的是哪一类客服，保存后检查前台入口、客服链接和在线时间是否与实际服务安排一致。",
    )

    last = add_heading_after(last, "在线客服", 4)
    last = add_picture_after(last, ONLINE_IMAGE, 6.35)
    last = add_caption_after(last, "图22：客服配置-在线客服页签，用于维护前台在线客服入口。")
    for label, text in [
        ("在线客服开关", "控制前台在线客服入口是否展示。启用后前台按当前头像、标题和说明展示入口；停用后用户侧不再显示该在线客服入口，需确认是否已有其他客服渠道兜底。"),
        ("客服头像", "上传前台展示的在线客服头像，建议使用 64*64 或等比例图片，支持 PNG、JPG、GIF，大小控制在 100KB 以内。编辑状态下可删除并重新上传。"),
        ("前端标题", "配置用户侧看到的客服入口标题，示例为“测试”。标题应简短明确，通常用于说明客服入口名称或服务类型。"),
        ("前端说明", "配置用户侧客服入口的说明文案，示例为“测试说明”。可说明服务范围、在线时间或联系前需要准备的信息。"),
        ("编辑与保存", "页面默认展示为只读状态，点击“编辑”后可修改头像、标题、说明和开关；保存前复核图片规格、文案语言和开关状态，取消则放弃本次调整。"),
    ]:
        last = add_bullet_after(last, label, text)

    last = add_heading_after(last, "其他客服", 4)
    last = add_picture_after(last, OTHER_IMAGE, 6.45)
    last = add_caption_after(last, "图23：客服配置-其他客服页签，用于维护 Telegram、WhatsApp、Facebook、Line 等 IM 客服。")
    for label, text in [
        ("新增", "点击右上角“新增”打开其他客服配置弹窗，用于新增一个 IM 类型及其下属客服账号。新增前应确认同类客服是否已存在，避免前台重复展示。"),
        ("排序", "控制其他客服在前台或列表中的展示顺序，数字越小通常越靠前。调整排序后需检查重点渠道是否排在预期位置。"),
        ("IM", "显示客服渠道类型，当前原型包含 Telegram、WhatsApp、Facebook、Line。不同 IM 对应不同图标和链接填写习惯。"),
        ("配置客服", "展示该 IM 下的客服昵称、链接及在线时间。单个 IM 可配置多名客服；若暂无客服账号，列表以“-”展示。"),
        ("图标", "展示该 IM 在前台使用的入口图标，需与 IM 类型保持一致，避免用户误点或识别错误。"),
        ("开关", "控制单个其他客服渠道是否启用。关闭后该渠道不应在用户侧展示；批量调整前需确认是否影响当前在线支持能力。"),
        ("操作", "“编辑”进入该行配置弹窗，可调整排序、IM、图标、客服明细和开关；“删除”用于移除该客服渠道，删除前应确认无前台正在使用的链接。"),
    ]:
        last = add_bullet_after(last, label, text)

    last = add_heading_after(last, "新增或编辑其他客服", 4)
    last = add_picture_after(last, MODAL_IMAGE, 5.6)
    last = add_caption_after(last, "图24：编辑其他客服弹窗，用于维护 IM 渠道、图标、客服账号和在线时间。")
    for label, text in [
        ("弹窗标题", "从列表点击“编辑”时标题显示“编辑其他客服”；点击“新增”时显示“新增其他客服”。可通过标题确认当前是在新增还是修改已有配置。"),
        ("排序与 IM", "排序为数字输入；IM 为必填下拉项，选择后应同步检查图标和客服明细中的图标是否一致。"),
        ("图标上传", "图标为必填项，使用 64*64 或等比例 PNG、JPG、GIF 图片，大小不超过 100KB。删除旧图标后需重新上传，避免保存后前台缺失入口图。"),
        ("配置客服明细", "每行维护一名客服，包括昵称、链接、在线时间、图标和删除操作。点击“添加”可增加一行；在线时间建议按 00:00-21:00 这类格式填写，便于用户判断可联系时段。"),
        ("开关", "控制当前 IM 渠道整体是否启用。即使客服明细已填写，开关关闭时前台也不应展示该渠道。"),
        ("取消与确定", "点击“取消”或右上角关闭不保存本次修改；点击“确定”保存弹窗内配置并回到其他客服列表，保存后应检查列表展示、开关状态和链接是否符合预期。"),
    ]:
        last = add_bullet_after(last, label, text)

    try:
        document.save(DOCX)
    except PermissionError:
        print(f"backup={backup}")
        raise

    print(f"backup={backup}")
    print(f"assets={ASSET_DIR}")


if __name__ == "__main__":
    main()
