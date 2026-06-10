from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"
ASSET_DIR = ROOT / "custom" / "assets" / "manual-member-complaint-list"
IMAGE_PATH = ASSET_DIR / "member-complaint-list-page.png"

TITLE = "会员投诉列表"
FONT_NAME = "微软雅黑"
BLUE = RGBColor(47, 85, 151)
CAPTION_GRAY = RGBColor(112, 126, 146)


def set_run(run, *, size: float = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def insert_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_body_after(anchor, text: str):
    paragraph = insert_after(anchor, style="Normal")
    run = paragraph.add_run(text)
    set_run(run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullet_after(anchor, label: str, text: str):
    paragraph = insert_after(anchor, style="List Bullet")
    label_run = paragraph.add_run(label)
    set_run(label_run, size=10, bold=True, color=BLUE)
    body_run = paragraph.add_run("：" + text)
    set_run(body_run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_picture_after(anchor, image_path: Path, width: float = 6.45):
    paragraph = insert_after(anchor, style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def add_caption_after(anchor, text: str):
    paragraph = insert_after(anchor, style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run(run, size=9, color=CAPTION_GRAY)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def reset_section(document: Document):
    start = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == TITLE:
            start = index
            break
    if start is None:
        raise RuntimeError(f"未找到“{TITLE}”章节。")

    end = len(document.paragraphs)
    for index in range(start + 1, len(document.paragraphs)):
        if document.paragraphs[index].style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            end = index
            break

    for paragraph in list(document.paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)

    anchor = document.paragraphs[start]
    for run in anchor.runs:
        set_run(run, size=12.5, bold=True, color=BLUE)
    return anchor


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    if not IMAGE_PATH.exists():
        raise FileNotFoundError(IMAGE_PATH)

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{DOCX.stem}.member-complaint-list-{stamp}.bak.docx"
    shutil.copy2(DOCX, backup)

    document = Document(DOCX)
    anchor = reset_section(document)

    last = add_body_after(
        anchor,
        "会员投诉列表用于集中查看玩家提交的投诉记录，并按会员、投诉时间和处理状态跟进投诉进度，完成回复、采纳或忽略等处理动作。",
    )
    last = add_picture_after(last, IMAGE_PATH, 6.45)
    last = add_caption_after(last, "图：会员投诉列表页面，用于筛选投诉记录并执行处理或查看操作。")

    items = [
        ("会员ID", "支持按会员ID精确筛选投诉记录，适合在接到会员反馈后快速定位对应账号的投诉历史。"),
        ("投诉时间", "可按开始时间和结束时间筛选投诉提交区间，用于核对某个活动周期、账变时段或客服处理时段内的投诉情况。"),
        ("状态", "支持按全部、待处理、已采纳筛选列表；处理中的记录可优先跟进，已采纳记录可结合回复内容复核处理结果。"),
        ("搜索/重置", "搜索按当前筛选条件刷新列表，重置用于清空条件并恢复查看全部投诉记录。"),
        ("列表字段", "列表展示投诉账号、投诉账号ID、状态、投诉时间、投诉类型、投诉内容、附件、回复内容、操作人和操作时间，便于同时核对投诉来源、处理结论和责任人。"),
        ("附件与回复", "有附件的投诉可直接查看附件入口；已处理记录会在回复内容中展示处理意见，方便复查客服答复是否完整。"),
        ("处理", "点击待处理记录右侧的处理可打开处理投诉窗口，查看原始反馈内容和附件，并填写回复内容后执行采纳或忽略。"),
        ("查看", "已处理记录可通过查看进入详情窗口，重点核对回复内容、处理人和操作时间，确认处理过程留痕完整。"),
        ("使用要点", "处理投诉前建议先核对会员账号、投诉类型以及关联订单、充值、活动或游戏记录；回复内容应写清处理结论和依据，避免后续重复投诉。"),
    ]

    for label_text, body_text in items:
        last = add_bullet_after(last, label_text, body_text)

    try:
        document.save(DOCX)
    except PermissionError:
        print(f"PERMISSION_ERROR backup={backup}")
        raise

    print(f"updated member complaint list; backup={backup}; image={IMAGE_PATH}")


if __name__ == "__main__":
    main()
