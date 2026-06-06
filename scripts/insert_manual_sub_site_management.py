from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "\u0042\u7aef\u540e\u53f0\u64cd\u4f5c\u624b\u518c.docx"
BACKUP_DIR = ROOT / "backups"
ASSET_DIR = ROOT / "custom" / "assets" / "manual-site-config" / "sub-site-management"
FONT_NAME = "\u5fae\u8f6f\u96c5\u9ed1"

SECTION_TITLE = "\u7ad9\u70b9\u914d\u7f6e"
INSERT_BEFORE = "\u54c1\u724c\u8bbe\u7f6e"
PAGE_TITLE = "\u5b50\u7ad9\u70b9\u7ba1\u7406"


CONTENT = [
    ("Heading 3", PAGE_TITLE),
    (
        "Normal",
        "\u5b50\u7ad9\u70b9\u7ba1\u7406\u7528\u4e8e\u7ef4\u62a4\u540c\u4e00\u5546\u6237\u4e0b\u7684\u591a\u4e2a\u5b50\u7ad9\u70b9\uff0c\u652f\u6301\u6309\u72b6\u6001\u67e5\u8be2\u7ad9\u70b9\uff0c\u521b\u5efa\u65b0\u7ad9\u70b9\uff0c\u4e3a\u5b50\u7ad9\u70b9\u914d\u7f6e\u9875\u9762\u6a21\u677f\u3001\u57df\u540d\u3001\u7b2c\u4e09\u65b9\u7ed1\u5b9a\u3001\u5ba2\u670d\u3001\u6e38\u620f\u3001\u6d3b\u52a8\u3001\u4ee3\u7406\u3001\u4efb\u52a1\u3001\u4f1a\u5458\u8fd4\u5229\u3001\u7968\u5238\u548c\u83b7\u5229\u76d1\u63a7\u7b49\u529f\u80fd\u3002",
    ),
    ("Heading 4", "\u5217\u8868\u67e5\u8be2\u4e0e\u7ad9\u70b9\u914d\u7f6e"),
    ("Image", "sub-site-list.png|\u56fe11\uff1a\u5b50\u7ad9\u70b9\u7ba1\u7406\u5217\u8868\uff0c\u7528\u4e8e\u67e5\u770b\u7ad9\u70b9\u53ca\u5404\u9879\u914d\u7f6e\u5165\u53e3\u3002|6.5"),
    (
        "List Bullet",
        "\u7b5b\u9009\u6761\u4ef6\uff1a\u9875\u9762\u5de6\u4e0a\u89d2\u652f\u6301\u6309\u72b6\u6001\u7b5b\u9009\uff0c\u9009\u62e9\u72b6\u6001\u540e\u70b9\u51fb\u641c\u7d22\u67e5\u8be2\uff0c\u70b9\u51fb\u91cd\u7f6e\u6e05\u7a7a\u5df2\u9009\u6761\u4ef6\u5e76\u6062\u590d\u9ed8\u8ba4\u5217\u8868\u3002",
    ),
    (
        "List Bullet",
        "\u5217\u8868\u5b57\u6bb5\uff1a\u8868\u683c\u5c55\u793a\u7ad9\u70b9\u7f16\u53f7\u3001\u7ad9\u70b9\u540d\u79f0\u3001\u9875\u9762\u6a21\u677f\u3001\u9875\u9762\u8bbe\u7f6e\u3001\u57df\u540d\u914d\u7f6e\u3001\u7b2c\u4e09\u65b9\u7ed1\u5b9a\u914d\u7f6e\u3001\u5ba2\u670d\u914d\u7f6e\u3001\u6e38\u620f\u914d\u7f6e\u3001\u6d3b\u52a8\u914d\u7f6e\u3001\u4ee3\u7406\u3001\u4efb\u52a1\u3001\u4f1a\u5458\u8fd4\u5229\u3001\u7968\u5238\u3001\u83b7\u5229\u76d1\u63a7\u548c\u64cd\u4f5c\u3002",
    ),
    (
        "List Bullet",
        "\u914d\u7f6e\u5165\u53e3\uff1a\u9875\u9762\u6a21\u677f\u5217\u663e\u793a\u5df2\u7ed1\u5b9a\u6a21\u677f\uff0c\u672a\u914d\u7f6e\u9879\u901a\u5e38\u663e\u793a\u4e3a\u53bb\u914d\u7f6e\u94fe\u63a5\uff0c\u70b9\u51fb\u540e\u8fdb\u5165\u5bf9\u5e94\u914d\u7f6e\u9875\u9762\u6216\u914d\u7f6e\u5f39\u7a97\u3002",
    ),
    (
        "List Bullet",
        "\u884c\u5185\u64cd\u4f5c\uff1a\u6bcf\u884c\u53f3\u4fa7\u63d0\u4f9b\u4fee\u6539\u5165\u53e3\uff0c\u7528\u4e8e\u8c03\u6574\u7ad9\u70b9\u540d\u79f0\u3001\u9875\u9762\u6a21\u677f\u548c\u72b6\u6001\u7b49\u57fa\u7840\u4fe1\u606f\u3002",
    ),
    ("Heading 4", "\u521b\u5efa\u7ad9\u70b9"),
    ("Image", "create-direct.png|\u56fe12\uff1a\u521b\u5efa\u7ad9\u70b9-\u76f4\u63a5\u521b\u5efa\uff0c\u9700\u586b\u5199\u7ad9\u70b9\u540d\u79f0\u5e76\u53ef\u9009\u62e9\u9875\u9762\u6a21\u677f\u3002|4.8"),
    ("Image", "create-copy.png|\u56fe13\uff1a\u521b\u5efa\u7ad9\u70b9-\u590d\u5236\u7ad9\u70b9\uff0c\u5148\u9009\u62e9\u8981\u590d\u5236\u7684\u7ad9\u70b9\u518d\u586b\u5199\u65b0\u7ad9\u70b9\u4fe1\u606f\u3002|4.8"),
    (
        "List Bullet",
        "\u5165\u53e3\uff1a\u70b9\u51fb\u5217\u8868\u53f3\u4e0a\u89d2\u521b\u5efa\u7ad9\u70b9\u6253\u5f00\u521b\u5efa\u5f39\u7a97\u3002",
    ),
    (
        "List Bullet",
        "\u76f4\u63a5\u521b\u5efa\uff1a\u521b\u5efa\u6a21\u5f0f\u9009\u62e9\u76f4\u63a5\u521b\u5efa\u65f6\uff0c\u9700\u586b\u5199\u7ad9\u70b9\u540d\u79f0\uff0c\u5e76\u53ef\u901a\u8fc7\u9009\u62e9\u6a21\u677f\u7ed1\u5b9a\u9875\u9762\u6a21\u677f\uff1b\u672a\u9009\u62e9\u65f6\u663e\u793a\u6682\u672a\u9009\u62e9\u3002",
    ),
    (
        "List Bullet",
        "\u590d\u5236\u7ad9\u70b9\uff1a\u521b\u5efa\u6a21\u5f0f\u9009\u62e9\u590d\u5236\u7ad9\u70b9\u65f6\uff0c\u9700\u5148\u5728\u590d\u5236\u7ad9\u70b9\u4e0b\u62c9\u6846\u4e2d\u9009\u62e9\u8981\u590d\u5236\u7684\u5df2\u6709\u7ad9\u70b9\uff0c\u518d\u586b\u5199\u65b0\u7ad9\u70b9\u540d\u79f0\u548c\u9875\u9762\u6a21\u677f\u3002",
    ),
    (
        "List Bullet",
        "\u6309\u94ae\u8bf4\u660e\uff1a\u70b9\u51fb\u53d6\u6d88\u5173\u95ed\u5f39\u7a97\u4e14\u4e0d\u4fdd\u5b58\uff0c\u70b9\u51fb\u786e\u5b9a\u63d0\u4ea4\u521b\u5efa\uff1b\u63d0\u4ea4\u524d\u9700\u786e\u8ba4\u5fc5\u586b\u9879\u5df2\u5b8c\u6210\uff0c\u7ad9\u70b9\u540d\u79f0\u672a\u4e0e\u5df2\u6709\u7ad9\u70b9\u6df7\u6dc6\u3002",
    ),
    ("Heading 4", "\u8bbe\u7f6e\u6700\u5c0f\u7528\u6237ID"),
    ("Image", "min-user-id.png|\u56fe14\uff1a\u8bbe\u7f6e\u6700\u5c0f\u7528\u6237ID\uff0c\u7528\u4e8e\u63a7\u5236\u5b50\u7ad9\u70b9\u7528\u6237ID\u8d77\u7b97\u8303\u56f4\u3002|3.2"),
    (
        "List Bullet",
        "\u7528\u9014\uff1a\u521b\u5efa\u6216\u7ef4\u62a4\u5b50\u7ad9\u70b9\u65f6\uff0c\u53ef\u901a\u8fc7\u8bbe\u7f6e\u6700\u5c0f\u7528\u6237ID\u63a7\u5236\u8be5\u7ad9\u70b9\u4e0b\u7528\u6237ID\u7684\u8d77\u7b97\u8303\u56f4\u3002",
    ),
    (
        "List Bullet",
        "\u5b57\u6bb5\u8bf4\u660e\uff1a\u6700\u5c0f\u7528\u6237ID\u9ed8\u8ba4\u793a\u4f8b\u4e3a10000\uff0c\u53ef\u901a\u8fc7\u8f93\u5165\u6846\u6216\u53f3\u4fa7\u52a0\u51cf\u63a7\u4ef6\u8c03\u6574\u3002",
    ),
    (
        "List Bullet",
        "\u6821\u9a8c\u89c4\u5219\uff1a\u82e5\u8bbe\u7f6e\u503c\u5c0f\u4e8e1000\uff0c\u7cfb\u7edf\u5c06\u63091000\u8ba1\u7b97\uff0c\u4fdd\u5b58\u524d\u5e94\u786e\u8ba4\u8be5\u503c\u4e0e\u7ad9\u70b9\u7528\u6237\u7f16\u53f7\u89c4\u5212\u4e00\u81f4\u3002",
    ),
    ("Heading 4", "\u7f16\u8f91\u7ad9\u70b9"),
    ("Image", "edit-site.png|\u56fe15\uff1a\u7f16\u8f91\u7ad9\u70b9\uff0c\u53ef\u4fee\u6539\u7ad9\u70b9\u540d\u79f0\u3001\u9875\u9762\u6a21\u677f\u548c\u7ad9\u70b9\u72b6\u6001\u3002|4.8"),
    (
        "List Bullet",
        "\u5165\u53e3\uff1a\u5728\u5217\u8868\u64cd\u4f5c\u5217\u70b9\u51fb\u4fee\u6539\uff0c\u6253\u5f00\u7f16\u8f91\u7ad9\u70b9\u5f39\u7a97\u3002",
    ),
    (
        "List Bullet",
        "\u7ad9\u70b9\u540d\u79f0\uff1a\u53ef\u4fee\u6539\u5b50\u7ad9\u70b9\u7684\u5c55\u793a\u540d\u79f0\uff0c\u4fee\u6539\u540e\u9700\u68c0\u67e5\u5217\u8868\u3001\u524d\u53f0\u5c55\u793a\u548c\u8fd0\u8425\u6c9f\u901a\u4e2d\u7684\u547d\u540d\u662f\u5426\u4e00\u81f4\u3002",
    ),
    (
        "List Bullet",
        "\u9875\u9762\u6a21\u677f\uff1a\u70b9\u51fb\u9009\u62e9\u6a21\u677f\u53ef\u66f4\u6362\u5df2\u7ed1\u5b9a\u7684\u9875\u9762\u6a21\u677f\uff0c\u5df2\u9009\u6a21\u677f\u4f1a\u5728\u6309\u94ae\u53f3\u4fa7\u663e\u793a\uff0c\u4f8b\u5982\u9996\u9875\u3002",
    ),
    (
        "List Bullet",
        "\u72b6\u6001\uff1a\u652f\u6301\u9009\u62e9\u51bb\u7ed3\u3001\u5df2\u5220\u9664\u3001\u5efa\u8bbe\u4e2d\u3001\u6b63\u5e38\u5f00\u542f\u3001\u7ef4\u62a4\u4e2d\u7b49\u72b6\u6001\uff1b\u5207\u6362\u72b6\u6001\u524d\u9700\u8bc4\u4f30\u5bf9\u524d\u53f0\u8bbf\u95ee\u3001\u7528\u6237\u64cd\u4f5c\u548c\u76f8\u5173\u914d\u7f6e\u7684\u5f71\u54cd\u3002",
    ),
    (
        "List Bullet",
        "\u4fdd\u5b58\u64cd\u4f5c\uff1a\u70b9\u51fb\u786e\u5b9a\u4fdd\u5b58\u4fee\u6539\uff0c\u70b9\u51fb\u53d6\u6d88\u6216\u53f3\u4e0a\u89d2\u5173\u95ed\u653e\u5f03\u672c\u6b21\u7f16\u8f91\u3002",
    ),
    ("Heading 4", "\u4f7f\u7528\u8981\u70b9"),
    (
        "List Bullet",
        "\u65b0\u5efa\u5b50\u7ad9\u70b9\u540e\uff0c\u5e94\u6309\u9875\u9762\u6a21\u677f\u3001\u9875\u9762\u8bbe\u7f6e\u3001\u57df\u540d\u3001\u5ba2\u670d\u3001\u6e38\u620f\u3001\u6d3b\u52a8\u548c\u7968\u5238\u7b49\u987a\u5e8f\u8865\u9f50\u5fc5\u8981\u914d\u7f6e\uff0c\u907f\u514d\u524d\u53f0\u9875\u9762\u53ef\u8bbf\u95ee\u4f46\u529f\u80fd\u672a\u914d\u597d\u3002",
    ),
    (
        "List Bullet",
        "\u590d\u5236\u7ad9\u70b9\u53ef\u63d0\u5347\u642d\u5efa\u6548\u7387\uff0c\u4f46\u4fdd\u5b58\u540e\u4ecd\u9700\u9010\u9879\u68c0\u67e5\u57df\u540d\u3001\u7b2c\u4e09\u65b9\u7ed1\u5b9a\u3001\u5ba2\u670d\u548c\u6d3b\u52a8\u914d\u7f6e\uff0c\u907f\u514d\u628a\u539f\u7ad9\u70b9\u7684\u6e20\u9053\u6216\u6d3b\u52a8\u5165\u53e3\u76f4\u63a5\u5e26\u5230\u65b0\u7ad9\u70b9\u3002",
    ),
    (
        "List Bullet",
        "\u5bf9\u5df2\u4e0a\u7ebf\u7ad9\u70b9\u8c03\u6574\u72b6\u6001\u6216\u6a21\u677f\u524d\uff0c\u5efa\u8bae\u5148\u786e\u8ba4\u5f71\u54cd\u65f6\u95f4\u3001\u53d7\u5f71\u54cd\u7528\u6237\u8303\u56f4\u548c\u56de\u9000\u65b9\u6848\u3002",
    ),
]


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def rect(draw, box, fill, outline="#dcdfe6", width=1):
    draw.rectangle(box, fill=fill, outline=outline, width=width)


def button(draw, box, text, fill="#409eff", outline="#409eff", text_fill="white"):
    rect(draw, box, fill, outline)
    x1, y1, x2, y2 = box
    draw.text(((x1 + x2) / 2, (y1 + y2) / 2), text, fill=text_fill, font=font(14), anchor="mm")


def input_box(draw, box, placeholder="", text="", arrow=False):
    rect(draw, box, "#ffffff", "#dcdfe6")
    x1, y1, x2, y2 = box
    value = text or placeholder
    color = "#606266" if text else "#b4bcc8"
    draw.text((x1 + 12, (y1 + y2) / 2), value, fill=color, font=font(13), anchor="lm")
    if arrow:
        draw.polygon([(x2 - 22, y1 + 14), (x2 - 12, y1 + 14), (x2 - 17, y1 + 22)], fill="#c0c4cc")


def draw_modal(draw, size, title):
    w, h = size
    rect(draw, (0, 0, w - 1, h - 1), "#ffffff", "#cfd4dc")
    rect(draw, (0, 0, w - 1, 40), "#ffffff", "#dcdfe6")
    draw.text((12, 20), title, fill="#303133", font=font(14, True), anchor="lm")
    draw.line((w - 28, 14, w - 16, 26), fill="#606266", width=1)
    draw.line((w - 16, 14, w - 28, 26), fill="#606266", width=1)


def draw_label(draw, x, y, text, required=False):
    if required:
        draw.text((x - 10, y), "*", fill="#f56c6c", font=font(13), anchor="rm")
    draw.text((x, y), text, fill="#606266", font=font(14), anchor="lm")


def save_sub_site_list():
    img = Image.new("RGB", (1645, 385), "#ffffff")
    d = ImageDraw.Draw(img)
    draw_label(d, 14, 49, "\u72b6\u6001\uff1a")
    input_box(d, (52, 31, 182, 62), "\u8bf7\u9009\u62e9\u72b6\u6001", arrow=True)
    button(d, (196, 31, 251, 62), "\u641c\u7d22")
    button(d, (272, 31, 324, 62), "\u91cd\u7f6e", "#ffffff", "#dcdfe6", "#303133")
    button(d, (1582, 31, 1660, 62), "\u521b\u5efa\u7ad9\u70b9")

    headers = [
        "\u7ad9\u70b9\u7f16\u53f7",
        "\u7ad9\u70b9\u540d\u79f0",
        "\u9875\u9762\u6a21\u677f",
        "\u9875\u9762\u8bbe\u7f6e",
        "\u57df\u540d\u914d\u7f6e",
        "\u7b2c\u4e09\u65b9\u7ed1\u5b9a\u914d\u7f6e",
        "\u5ba2\u670d\u914d\u7f6e",
        "\u6e38\u620f\u914d\u7f6e",
        "\u6d3b\u52a8\u914d\u7f6e",
        "\u4ee3\u7406",
        "\u4efb\u52a1",
        "\u4f1a\u5458\u8fd4\u5229",
        "\u7968\u5238",
        "\u83b7\u5229\u76d1\u63a7",
        "\u64cd\u4f5c",
    ]
    widths = [135, 135, 135, 99, 99, 117, 99, 99, 99, 99, 99, 99, 99, 99, 134]
    x = 13
    y = 79
    h = 43
    for header, width in zip(headers, widths):
        rect(d, (x, y, x + width, y + h), "#f7f8fa", "#e6e8ee")
        d.text((x + width / 2, y + h / 2), header, fill="#303133", font=font(13, True), anchor="mm")
        x += width
    rows = [
        ["104", "\u6d4b\u8bd5003", "\u9996\u9875"],
        ["103", "\u6d4b\u8bd5\u7ad903", "\u9996\u9875"],
        ["102", "\u6d4b\u8bd5\u7ad92", "\u6d4b\u8bd5\u7ad9\u70b9 \u6a21\u677f"],
        ["101", "\u6d4b\u8bd5\u7ad9", "\u9996\u9875"],
    ]
    for r, row in enumerate(rows):
        x = 13
        y = 122 + r * 43
        for c, width in enumerate(widths):
            rect(d, (x, y, x + width, y + 43), "#ffffff", "#e6e8ee")
            if c == 0:
                text, color = row[0], "#606266"
            elif c == 1:
                text, color = row[1], "#606266"
            elif c == 2:
                text, color = row[2], "#1683ff"
            elif c == len(widths) - 1:
                text, color = "\u4fee\u6539", "#1683ff"
            else:
                text, color = "\u53bb\u914d\u7f6e", "#1683ff"
            d.text((x + width / 2, y + 22), text, fill=color, font=font(13), anchor="mm")
            x += width
    img.save(ASSET_DIR / "sub-site-list.png")


def save_create_direct():
    img = Image.new("RGB", (538, 258), "#ffffff")
    d = ImageDraw.Draw(img)
    draw_modal(d, img.size, "\u521b\u5efa\u7ad9\u70b9")
    draw_label(d, 40, 75, "\u521b\u5efa\u6a21\u5f0f")
    d.ellipse((105, 68, 119, 82), fill="#409eff")
    d.ellipse((110, 73, 114, 77), fill="#ffffff")
    d.text((126, 75), "\u76f4\u63a5\u521b\u5efa", fill="#1683ff", font=font(13), anchor="lm")
    d.ellipse((185, 68, 199, 82), outline="#cfd4dc", width=2)
    d.text((206, 75), "\u590d\u5236\u7ad9\u70b9", fill="#606266", font=font(13), anchor="lm")
    draw_label(d, 40, 126, "\u7ad9\u70b9\u540d\u79f0", True)
    input_box(d, (105, 111, 522, 141), "\u8bf7\u8f93\u5165\u7ad9\u70b9\u540d\u79f0")
    draw_label(d, 40, 176, "\u9875\u9762\u6a21\u677f")
    button(d, (105, 161, 199, 192), "\u270e \u9009\u62e9\u6a21\u677f", "#ffffff", "#dcdfe6", "#606266")
    d.text((213, 176), "\u6682\u672a\u9009\u62e9", fill="#b4bcc8", font=font(13), anchor="lm")
    button(d, (413, 216, 466, 247), "\u53d6\u6d88", "#ffffff", "#dcdfe6", "#606266")
    button(d, (473, 216, 525, 247), "\u786e\u5b9a")
    img.save(ASSET_DIR / "create-direct.png")


def save_create_copy():
    img = Image.new("RGB", (538, 300), "#ffffff")
    d = ImageDraw.Draw(img)
    draw_modal(d, img.size, "\u521b\u5efa\u7ad9\u70b9")
    draw_label(d, 40, 75, "\u521b\u5efa\u6a21\u5f0f")
    d.ellipse((105, 68, 119, 82), outline="#cfd4dc", width=2)
    d.text((126, 75), "\u76f4\u63a5\u521b\u5efa", fill="#606266", font=font(13), anchor="lm")
    d.ellipse((185, 68, 199, 82), fill="#1683ff")
    d.ellipse((190, 73, 194, 77), fill="#ffffff")
    d.text((206, 75), "\u590d\u5236\u7ad9\u70b9", fill="#1683ff", font=font(13), anchor="lm")
    draw_label(d, 40, 126, "\u590d\u5236\u7ad9\u70b9")
    input_box(d, (105, 111, 522, 141), "\u8bf7\u9009\u62e9\u8981\u590d\u5236\u7684\u7ad9\u70b9", arrow=True)
    draw_label(d, 40, 176, "\u7ad9\u70b9\u540d\u79f0", True)
    input_box(d, (105, 161, 522, 192), "\u8bf7\u8f93\u5165\u7ad9\u70b9\u540d\u79f0")
    draw_label(d, 40, 226, "\u9875\u9762\u6a21\u677f")
    button(d, (105, 213, 199, 244), "\u270e \u9009\u62e9\u6a21\u677f", "#ffffff", "#dcdfe6", "#606266")
    d.text((213, 228), "\u6682\u672a\u9009\u62e9", fill="#b4bcc8", font=font(13), anchor="lm")
    button(d, (413, 269, 466, 298), "\u53d6\u6d88", "#ffffff", "#dcdfe6", "#606266")
    button(d, (473, 269, 525, 298), "\u786e\u5b9a")
    img.save(ASSET_DIR / "create-copy.png")


def save_min_user_id():
    img = Image.new("RGB", (355, 213), "#ffffff")
    d = ImageDraw.Draw(img)
    draw_modal(d, img.size, "\u8bbe\u7f6e\u6700\u5c0f\u7528\u6237ID")
    draw_label(d, 64, 96, "\u6700\u5c0f\u7528\u6237ID")
    input_box(d, (141, 80, 325, 111), text="10000")
    rect(d, (289, 80, 325, 95), "#f5f7fa", "#dcdfe6")
    rect(d, (289, 96, 325, 111), "#f5f7fa", "#dcdfe6")
    d.text((307, 87), "+", fill="#606266", font=font(14), anchor="mm")
    d.text((307, 103), "-", fill="#606266", font=font(14), anchor="mm")
    d.text((135, 135), "\u82e5\u8bbe\u7f6e\u503c\u5c0f\u4e8e 1000\uff0c\u7cfb\u7edf\u5c06\u6309 1000 \u8ba1\u7b97", fill="#909399", font=font(12))
    button(d, (233, 173, 286, 204), "\u53d6\u6d88", "#ffffff", "#dcdfe6", "#606266")
    button(d, (293, 173, 345, 204), "\u786e\u5b9a")
    img.save(ASSET_DIR / "min-user-id.png")


def save_edit_site():
    img = Image.new("RGB", (543, 431), "#ffffff")
    d = ImageDraw.Draw(img)
    draw_modal(d, img.size, "\u7f16\u8f91\u7ad9\u70b9")
    draw_label(d, 43, 78, "\u7ad9\u70b9\u540d\u79f0", True)
    input_box(d, (108, 63, 526, 93), text="\u6d4b\u8bd5003")
    draw_label(d, 43, 128, "\u9875\u9762\u6a21\u677f")
    button(d, (108, 113, 202, 144), "\u270e \u9009\u62e9\u6a21\u677f", "#ffffff", "#dcdfe6", "#606266")
    d.text((215, 128), "\u5df2\u9009\u62e9\u6a21\u677f\uff1a", fill="#606266", font=font(13), anchor="lm")
    d.text((315, 128), "\u9996\u9875", fill="#1683ff", font=font(13), anchor="lm")
    draw_label(d, 70, 178, "\u72b6\u6001")
    input_box(d, (108, 163, 526, 194), text="\u6b63\u5e38\u5f00\u542f", arrow=True)
    rect(d, (108, 198, 526, 399), "#ffffff", "#dcdfe6")
    for i, text in enumerate(["\u51bb\u7ed3", "\u5df2\u5220\u9664", "\u5efa\u8bbe\u4e2d", "\u6b63\u5e38\u5f00\u542f", "\u7ef4\u62a4\u4e2d"]):
        y = 216 + i * 31
        color = "#1683ff" if text == "\u6b63\u5e38\u5f00\u542f" else "#606266"
        d.text((116, y), text, fill=color, font=font(13), anchor="lm")
    img.save(ASSET_DIR / "edit-site.png")


def generate_images():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    save_sub_site_list()
    save_create_direct()
    save_create_copy()
    save_min_user_id()
    save_edit_site()


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT_NAME
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)


def format_inserted_paragraph(paragraph) -> None:
    style_name = paragraph.style.name
    if style_name == "Normal":
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            set_run_font(run, size=10, bold=False)
    elif style_name == "Caption":
        paragraph.alignment = 1
        paragraph.paragraph_format.space_after = Pt(6)
        for run in paragraph.runs:
            set_run_font(run, size=9, bold=False, color="808080")
    elif style_name == "List Bullet":
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = 1.12
        for run in paragraph.runs:
            set_run_font(run, size=10, bold=False)


def insert_paragraph_before(paragraph, text: str, style: str):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    inserted = paragraph._parent.add_paragraph()
    inserted._p = new_p
    inserted.style = style
    if text:
        run = inserted.add_run(text)
        if style == "Heading 3":
            set_run_font(run, size=12.5, bold=True, color="2F5597")
        elif style == "Heading 4":
            set_run_font(run, size=11, bold=True, color="44546A")
        else:
            set_run_font(run, size=10, bold=(style.startswith("Heading")))
    format_inserted_paragraph(inserted)
    return inserted


def insert_image_before(paragraph, payload: str):
    filename, caption, width = payload.split("|")
    image_path = ASSET_DIR / filename
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    inserted = paragraph._parent.add_paragraph()
    inserted._p = new_p
    inserted.alignment = 1
    inserted.add_run().add_picture(str(image_path), width=Inches(float(width)))

    cap_p = OxmlElement("w:p")
    paragraph._p.addprevious(cap_p)
    cap = paragraph._parent.add_paragraph()
    cap._p = cap_p
    cap.style = "Caption"
    run = cap.add_run(caption)
    set_run_font(run, size=9, bold=False, color="808080")
    format_inserted_paragraph(cap)


def find_insert_index(document: Document) -> int:
    in_site_config = False
    for i, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name
        if style_name == "Heading 1":
            in_site_config = text == SECTION_TITLE or text.endswith(SECTION_TITLE)
        elif in_site_config and style_name == "Heading 3" and text == INSERT_BEFORE:
            return i
    raise RuntimeError("insert point not found")


def remove_existing_section(document: Document) -> None:
    start = None
    end = None
    for i, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == PAGE_TITLE:
            start = i
            break
    if start is None:
        return
    for j in range(start + 1, len(document.paragraphs)):
        style_name = document.paragraphs[j].style.name
        if style_name in {"Heading 1", "Heading 2", "Heading 3"}:
            end = j
            break
    if end is None:
        end = len(document.paragraphs)
    for paragraph in list(document.paragraphs[start:end]):
        p = paragraph._element
        p.getparent().remove(p)


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{DOCX.stem}.sub-site-backup.{stamp}.docx"
    shutil.copy2(DOCX, backup)

    document = Document(DOCX)
    generate_images()
    remove_existing_section(document)
    insert_index = find_insert_index(document)
    marker = document.paragraphs[insert_index]
    for style, text in CONTENT:
        if style == "Image":
            insert_image_before(marker, text)
        else:
            insert_paragraph_before(marker, text, style)

    document.save(DOCX)
    print(f"backup={backup}")
    print(f"inserted={PAGE_TITLE}")


if __name__ == "__main__":
    main()
