from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"
ASSET_DIR = ROOT / "custom" / "assets" / "manual-platform-suggestion"
IMAGE_PATH = ASSET_DIR / "chat-screenshot.png"

TITLE = "给平台提建议"
FONT_NAME = "微软雅黑"
BLUE = RGBColor(47, 85, 151)
CAPTION_GRAY = RGBColor(112, 126, 146)


def set_run(run, *, size: float = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def insert_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_body_after(anchor, text: str):
    paragraph = insert_after(anchor, style="Normal")
    run = paragraph.add_run(text)
    set_run(run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullet_after(anchor, label: str, text: str):
    paragraph = insert_after(anchor, style="List Bullet")
    label_run = paragraph.add_run(label)
    set_run(label_run, size=10, bold=True, color=BLUE)
    body_run = paragraph.add_run("：" + text)
    set_run(body_run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_picture_after(anchor, image_path: Path, width: float = 6.45):
    paragraph = insert_after(anchor, style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def add_caption_after(anchor, text: str):
    paragraph = insert_after(anchor, style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run(run, size=9, color=CAPTION_GRAY)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def reset_section(document: Document):
    start = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == TITLE:
            start = index
            break
    if start is None:
        raise RuntimeError(f"未找到“{TITLE}”章节。")

    end = len(document.paragraphs)
    for index in range(start + 1, len(document.paragraphs)):
        if document.paragraphs[index].style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            end = index
            break

    for paragraph in list(document.paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)

    anchor = document.paragraphs[start]
    for run in anchor.runs:
        set_run(run, size=12.5, bold=True, color=BLUE)
    return anchor


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    if not IMAGE_PATH.exists():
        raise FileNotFoundError(IMAGE_PATH)

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{DOCX.stem}.platform-suggestion-{stamp}.bak.docx"
    shutil.copy2(DOCX, backup)

    document = Document(DOCX)
    anchor = reset_section(document)

    last = add_body_after(
        anchor,
        "给平台提建议页面用于提交产品、运营或系统使用过程中的优化建议，并在列表中跟踪建议的处理状态、回复内容和附件记录。",
    )
    last = add_picture_after(last, IMAGE_PATH, 6.45)
    last = add_caption_after(last, "图：给平台提建议页面，用于按状态筛选建议记录并查看平台处理结果。")

    items = [
        ("时间范围", "可按开始时间和结束时间筛选建议提交记录，适合按活动期、版本发布时间或问题反馈时间段追踪处理情况。"),
        ("状态", "支持按全部、已采纳、待处理、已忽略筛选建议，便于优先跟进还未处理或已确认采纳的内容。"),
        ("搜索/重置", "搜索按当前筛选条件刷新列表；重置用于清空条件并恢复查看全部建议记录。"),
        ("新增", "点击新增后填写反馈内容，并可上传图片或视频附件，提交后系统会将建议写入列表等待平台处理。"),
        ("列表字段", "列表展示建议 ID、状态、提交时间、反馈内容、附件、回复内容、操作人和操作时间，方便核对建议进度与处理结果。"),
        ("状态说明", "已采纳表示平台已确认采纳建议；待处理表示建议已提交但还未完成反馈；已忽略表示平台暂不采用该建议。"),
        ("查看", "点击列表中的查看可打开详情窗口，查看完整反馈内容、附件预览和平台回复内容。"),
        ("使用要点", "提交建议时应尽量写清问题场景、复现步骤、期望效果和影响范围；如涉及页面异常，建议同时上传截图或录屏附件，便于平台快速定位。"),
    ]

    for label_text, body_text in items:
        last = add_bullet_after(last, label_text, body_text)

    try:
        document.save(DOCX)
    except PermissionError:
        print(f"PERMISSION_ERROR backup={backup}")
        raise

    print(f"updated platform suggestion; backup={backup}; image={IMAGE_PATH}")


if __name__ == "__main__":
    main()
