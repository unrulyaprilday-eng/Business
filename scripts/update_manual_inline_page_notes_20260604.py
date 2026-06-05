from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path("B端后台操作手册.docx")
FONT_NAME = "微软雅黑"


def set_font(paragraph, size: float | None = None, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        if size is not None:
            run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)
    set_font(paragraph, size=10)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_text(doc: Document, text: str, start: int = 0) -> int:
    for index, paragraph in enumerate(doc.paragraphs[start:], start):
        if paragraph.text.strip() == text:
            return index
    raise ValueError(f"not found: {text}")


def find_heading1_after(doc: Document, start: int) -> int:
    for index, paragraph in enumerate(doc.paragraphs[start + 1 :], start + 1):
        if paragraph.style.name == "Heading 1" and paragraph.text.strip():
            return index
    return len(doc.paragraphs)


def blank_paragraph(paragraph) -> None:
    paragraph.clear()
    set_font(paragraph, size=10)


def strip_standalone_heading(doc: Document, heading_text: str, remove_texts: set[str]) -> None:
    start = find_text(doc, heading_text)
    delete_paragraph(doc.paragraphs[start])
    for paragraph in doc.paragraphs[start:]:
        text = paragraph.text.strip()
        if paragraph.style.name == "Heading 1" and text:
            break
        if text in remove_texts:
            if any(run._element.xpath(".//w:drawing") for run in paragraph.runs):
                blank_paragraph(paragraph)
            else:
                delete_paragraph(paragraph)
            break


def main() -> None:
    backup = DOCX.with_suffix(f".inline-page-notes-{datetime.now():%Y%m%d-%H%M%S}.bak")
    shutil.copy2(DOCX, backup)

    doc = Document(DOCX)

    page_notes = {
        "站点应用设置：维护站点基础展示、首页入口、分享卡片、侧滑优惠中心、首页快捷操作和首页悬浮按钮等前台基础配置；其中网站图标用于维护 Logo 与 Favicon，币种管理用于维护货币图标和货币符号，分享卡片配置用于维护 Open Graph 分享标题、描述、图片和落地地址，侧滑优惠中心、首页快捷操作和首页悬浮按钮用于维护前台活动入口、快捷入口及悬浮按钮组。": "站点应用设置",
        "首页与模板管理：首页看板用于查看今日、昨日、本月等核心经营指标、趋势、待办事项和快捷入口；模板管理用于维护首页、登录页、VIP 页等前台展示方案，可配置组件内容、页面布局、主题预设、主色调、背景、按钮样式、图片、跳转链接和排序。": "首页与模板管理",
    }

    for replacement, prefix in page_notes.items():
        for paragraph in doc.paragraphs:
            if paragraph.text.strip().startswith(prefix + "："):
                set_paragraph_text(paragraph, replacement)
                paragraph.style = "List Bullet"
                break
        else:
            raise ValueError(f"page note not found: {prefix}")

    # Keep screenshots and card-level details, but remove the extra wrapper headings
    # and duplicate standalone page summaries now represented above.
    strip_standalone_heading(
        doc,
        "站点应用配置",
        {
            "站点配置用于维护站点基础展示、首页入口、分享卡片和活动入口等前台基础配置。配置保存后会影响前台品牌露出、用户访问入口、社交分享效果和首页运营位展示。"
        },
    )
    strip_standalone_heading(doc, "首页与模板管理", set())
    strip_standalone_heading(
        doc,
        "首页看板",
        {
            "登入后台后即可见首页看板，用于快速查看站点当下的即时状况。看板通常包含核心经营指标、趋势、待办事项和快捷入口，适合运营人员每日进入后台后先进行总览。"
        },
    )
    strip_standalone_heading(
        doc,
        "模板管理",
        {
            "模板管理用于维护前台页面或业务场景的展示模板，是日常运营中维护前台展示的重要页面。运营人员可配置首页、登录页、VIP 页等不同页面的展示方案，包括组件内容、页面布局、主题预设、主色调、背景、按钮样式等。"
        },
    )

    doc.save(DOCX)
    print(f"backup={backup}")


if __name__ == "__main__":
    main()
