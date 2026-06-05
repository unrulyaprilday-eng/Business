# -*- coding: utf-8 -*-
from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(".")
DOCX_PATH = BASE / "B端后台操作手册.docx"
ASSET_DIR = BASE / "custom" / "assets" / "manual-reorganized"
FONT_NAME = "微软雅黑"
DEEP_BLUE = RGBColor(37, 64, 97)
MUTED = RGBColor(96, 108, 128)


IMAGE_NAMES = [
    "login-top-nav.png",
    "login-toolbar.png",
    "home-dashboard.png",
    "template-management.png",
    "site-icons.png",
    "site-currency.png",
    "site-share-card.png",
    "site-slide-promo.png",
    "site-quick-actions.png",
    "site-float-buttons.png",
]


def backup_docx() -> Path:
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    backup = DOCX_PATH.with_suffix(f".docx.reorg-{stamp}.bak")
    shutil.copy2(DOCX_PATH, backup)
    return backup


def extract_existing_images() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    old_doc = Document(str(DOCX_PATH))
    blobs = []
    for paragraph in old_doc.paragraphs:
        for run in paragraph.runs:
            for drawing in run._element.xpath(".//w:drawing"):
                blips = drawing.xpath(".//a:blip")
                if not blips:
                    continue
                rid = blips[0].get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                part = old_doc.part.related_parts.get(rid)
                if part is not None:
                    blobs.append(part.blob)

    images = {}
    for index, name in enumerate(IMAGE_NAMES):
        if index >= len(blobs):
            break
        path = ASSET_DIR / name
        path.write_bytes(blobs[index])
        images[name] = path
    return images


def set_run_font(run, size=10, bold=False, color=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_doc_defaults(doc: Document):
    styles = doc.styles
    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(10)

    for style_name, size in [("Heading 1", 14), ("Heading 2", 12.5), ("Heading 3", 10.5)]:
        style = styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = DEEP_BLUE


def add_title(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=18, bold=True, color=DEEP_BLUE)
    paragraph.paragraph_format.space_after = Pt(12)


def add_heading(doc: Document, text: str, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    paragraph.paragraph_format.space_after = Pt(4)
    for run in paragraph.runs:
        set_run_font(run, size={1: 14, 2: 12.5, 3: 10.5}.get(level, 10.5), bold=True, color=DEEP_BLUE)
    return paragraph


def add_para(doc: Document, text: str, bold_prefix: str | None = None, color=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, bold=True, color=DEEP_BLUE)
        rest = text[len(bold_prefix) :]
        if rest:
            run = paragraph.add_run(rest)
            set_run_font(run, color=color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, color=color)
    return paragraph


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, bold=True, color=DEEP_BLUE)
        rest = text[len(bold_prefix) :]
        if rest:
            run = paragraph.add_run(rest)
            set_run_font(run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_number(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color=MUTED)
    return paragraph


def add_image(doc: Document, images: dict[str, Path], name: str, caption: str, width=6.25):
    path = images.get(name)
    if not path or not path.exists():
        add_para(doc, f"图片占位：{caption}", color=MUTED)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_template_block(doc: Document, module: dict):
    add_heading(doc, "模块定位", 2)
    add_para(doc, module["position"])

    add_heading(doc, "页面说明", 2)
    for item in module["pages"]:
        add_bullet(doc, item, item.split("：", 1)[0] + "：" if "：" in item else None)

    add_heading(doc, "字段说明", 2)
    for item in module["fields"]:
        add_bullet(doc, item, item.split("：", 1)[0] + "：" if "：" in item else None)

    add_heading(doc, "操作流程", 2)
    for item in module["flows"]:
        add_number(doc, item)

    add_heading(doc, "注意事项", 2)
    for item in module["notes"]:
        add_bullet(doc, item)

    add_heading(doc, "常见问题", 2)
    for question, answer in module["faqs"]:
        add_bullet(doc, f"{question}：{answer}", f"{question}：")


MODULES = [
    {
        "chapter": "三、商户初始化配置",
        "position": "商户初始化配置用于完成新站点上线前的基础参数、品牌展示、域名入口、客服入口、登录分享和支付通道配置。该模块决定前台能否正常访问、用户能否完成登录注册、充值通道是否可用以及品牌展示是否符合交付要求。",
        "pages": [
            "站点设置：维护站点基础展示、首页入口、分享卡片、侧滑优惠中心、首页快捷操作和首页悬浮按钮等前台基础配置。",
            "品牌与图标配置：维护 Logo、Favicon、币种图标、货币符号和品牌素材，影响前台页头、浏览器标签、钱包、充值和活动金额展示。",
            "域名配置：维护站点访问域名、备用域名、状态和备注，保证用户访问入口可控。",
            "客服配置：维护在线客服、投诉入口、反馈入口和联系方式，保证用户问题可被及时承接。",
            "登录与分享配置：维护第三方登录、分享标题、分享描述、分享图片和分享落地地址。",
            "支付通道配置：维护充值和提现相关通道参数、限额、费率、状态、适用范围和排序。",
        ],
        "fields": [
            "基础信息：包含站点名称、站点状态、默认语言、默认币种、时区、维护状态等信息。",
            "品牌素材：包含 Logo、Favicon、分享图、入口图标、悬浮按钮图标等图片字段，需关注尺寸、格式、清晰度和品牌一致性。",
            "链接与路由：包含域名、客服链接、分享 URL、活动入口路由、快捷入口路由和第三方登录回调地址。",
            "状态与排序：启用状态控制前台是否展示；排序字段控制同类入口或配置项的展示优先级。",
            "支付参数：包含通道编码、商户号、密钥、费率、单笔限额、日限额、适用会员范围和回调配置。",
        ],
        "flows": [
            "进入商户初始化配置相关页面，先完成站点设置、品牌素材、域名和客服入口。",
            "配置登录、分享和前台入口，检查路由、图片、排序、启用状态与前台展示位置。",
            "配置支付通道并复核通道参数、金额限制、费率和适用范围。",
            "保存后按登录、访问、分享、充值、客服入口的顺序做上线前检查。",
        ],
        "notes": [
            "站点设置是页面，不作为独立模块处理；手册中统一放在商户初始化配置下说明。",
            "支付通道、第三方登录密钥和域名配置属于高影响配置，修改前需确认权限和审批依据。",
            "图片字段保存后可能影响多个前台位置，需同时检查 PC、移动端和分享预览效果。",
        ],
        "faqs": [
            ("配置保存后前台没变化", "检查配置状态、生效时间、缓存、域名访问入口和是否保存到当前站点。"),
            ("分享出去的标题或图片不对", "检查分享卡片配置中的 og:title、og:description、og:image、og:url，并确认外部渠道缓存是否已刷新。"),
            ("支付通道无法下单", "先核对通道状态、限额、适用范围和密钥，再结合充值订单与支付通道数据排查。"),
        ],
    },
    {
        "chapter": "四、商户中心",
        "position": "商户中心用于查看和维护商户自身信息、商户余额、商户充值、充值记录、账变记录和账单数据。该模块服务于商户账户层面的资金核对与结算追踪，与玩家资金管理区分开。",
        "pages": [
            "商户信息：查看商户名称、状态、联系人、账户余额、结算信息和基础资料。",
            "商户充值：为商户账户进行充值或额度补充，提交前需核对商户、金额和备注。",
            "充值记录：查询商户充值历史，包含金额、状态、操作人、时间和备注。",
            "商户账变记录：追踪商户账户资金变动，核对变动前余额、变动金额、变动后余额和变动原因。",
            "商户账单：查看商户账期、费用、结算金额、账单状态和明细。",
        ],
        "fields": [
            "商户字段：包含商户名称、商户编号、状态、联系人、账户余额和结算配置。",
            "充值字段：包含充值金额、充值方式、充值状态、操作人、提交时间和备注。",
            "账变字段：包含账变类型、变动前余额、变动金额、变动后余额、关联单据和原因。",
            "账单字段：包含账期、费用类型、应收应付金额、结算状态、生成时间和确认状态。",
        ],
        "flows": [
            "进入商户中心，先查看商户信息和当前余额。",
            "需要补充额度时进入商户充值，确认金额、商户和备注后提交。",
            "充值或结算后，通过充值记录、商户账变记录和商户账单交叉核对。",
            "遇到账务争议时固定账期和时间范围，导出或记录关键单据用于复核。",
        ],
        "notes": [
            "商户中心处理的是商户账户资金，不等同于玩家充值、提现和会员余额。",
            "商户充值、账变和账单核对时，应固定同一账期、同一币种和同一站点。",
            "账单确认前建议与商户充值记录和商户账变记录交叉核对。",
        ],
        "faqs": [
            ("商户余额和账单金额不一致", "检查账期、账变记录、充值记录、费用口径和是否存在未确认账单。"),
            ("商户充值后余额没变化", "检查充值状态、操作站点、账变记录和是否提交成功。"),
            ("商户账单需要复核什么", "重点核对账期、费用项、结算金额、已充值金额和账变明细。"),
        ],
    },
    {
        "chapter": "五、日常运营管理",
        "position": "日常运营管理用于承载运营人员每天最常使用的内容发布、活动配置、任务配置、公告消息、宣传位、客服投诉和首页模板维护。该模块的重点是提高前台内容更新效率，并保证活动和消息不会误投、漏投或超预算。",
        "pages": [
            "首页与模板管理：通过首页看板查看经营概览，通过模板管理维护首页、登录页、VIP 页等前台页面展示方案。",
            "优惠活动：管理优惠活动列表、活动创建、幸运转盘、票券中心和优惠数据。",
            "任务中心：配置用户任务、任务条件、奖励、周期、参与范围和任务状态。",
            "公告/消息/宣传：维护平台公告、消息推送、宣传管理、落地页和推荐内容。",
            "客服与投诉：处理客服会话、会员投诉、有奖反馈和用户建议。",
        ],
        "fields": [
            "内容字段：包含标题、正文、图片、跳转链接、展示位置、展示时间和目标人群。",
            "活动字段：包含活动类型、参与条件、奖励规则、预算、领取限制、有效期和展示状态。",
            "任务字段：包含任务名称、完成条件、奖励类型、周期、参与范围、排序和状态。",
            "消息字段：包含发送对象、发送渠道、发送时间、标题、内容、状态和送达结果。",
            "客服字段：包含会员账号、问题类型、处理状态、处理人、处理备注和附件。",
        ],
        "flows": [
            "进入对应运营页面，先通过筛选条件确认当前内容或活动是否已存在。",
            "点击新增或编辑，填写基础信息、展示信息、目标范围和生效时间。",
            "保存前复核图片、链接、时间、状态、预算、目标人群和移动端展示效果。",
            "上线后通过活动报表、消息状态、客服记录或首页看板观察实际效果。",
        ],
        "notes": [
            "活动、任务和消息推送上线前必须复核目标范围，避免误发给非目标用户。",
            "优惠活动与票券配置会影响成本和用户权益，建议保留审批记录。",
            "模板管理调整会直接影响前台页面结构，修改后需重点检查首页、登录页和 VIP 页展示。",
        ],
        "faqs": [
            ("活动列表查不到活动", "检查活动状态、活动时间、活动类型和站点筛选条件。"),
            ("消息没有发送给目标用户", "检查发送对象、用户分群、发送时间、状态和是否命中权限或渠道限制。"),
            ("模板修改后展示异常", "检查组件是否启用、组件排序、图片尺寸、跳转链接和页面主题配置。"),
        ],
    },
    {
        "chapter": "六、用户与资金管理",
        "position": "用户与资金管理用于查询会员资料、维护 VIP 和会员规则、处理充值提现、执行人工加扣款，并追踪用户资金变动。该模块涉及用户资产和资金安全，是后台日常操作中风险最高、审核要求最严格的部分。",
        "pages": [
            "用户查询：通过所有会员、在线玩家列表等页面查询会员资料、状态、余额、标签和行为记录。",
            "VIP 与会员配置：维护 VIP 等级、会员充值配置、会员提现设置和相关权益规则。",
            "充值订单：查询充值订单状态、支付通道、金额、回调结果和异常订单。",
            "提现审核：处理提现订单列表和提现风控审核中的待审核申请。",
            "人工加扣款：对会员账户执行人工加款或扣款，要求填写明确原因并确认金额方向。",
            "账变追踪：通过账变记录、余额变动、账户余额详情和资金账变类型追踪完整资金链路。",
        ],
        "fields": [
            "用户字段：包含会员账号、UID、手机号、注册时间、会员等级、VIP 等级、状态、渠道和标签。",
            "订单字段：包含订单号、会员账号、金额、通道、状态、提交时间、完成时间、处理人和备注。",
            "审核字段：包含审核状态、风控提示、打码要求、账户信息、驳回原因和审核备注。",
            "账变字段：包含账变类型、变动前余额、变动金额、变动后余额、关联订单和来源模块。",
        ],
        "flows": [
            "用户问题先定位会员账号或 UID，再进入会员资料和资金相关页面交叉核对。",
            "充值异常优先查看充值订单，再查看支付通道数据、回调状态和账变记录。",
            "提现审核需核对账户信息、金额、打码、风控提示和历史资金记录，再决定通过、驳回或转风控。",
            "人工加扣款提交前复核会员、金额、方向、原因和审批依据，提交后检查账变是否生成。",
        ],
        "notes": [
            "涉及会员资产的操作必须遵循内部审批流程，并填写可追溯的备注。",
            "不要只凭单一页面判断资金问题，充值、提现、人工加扣款和账变记录需要联动核对。",
            "调整 VIP、提现或充值规则前，应确认对存量用户和当前活动的影响。",
        ],
        "faqs": [
            ("用户说充值到账不对", "核对充值订单状态、支付通道回调、会员账变和账户余额详情。"),
            ("提现审核无法通过", "检查会员账户、打码要求、风控命中原因、黑名单状态和审核权限。"),
            ("人工加扣款后余额没变", "检查提交是否成功、审批状态、账变记录和是否选错会员或站点。"),
        ],
    },
    {
        "chapter": "七、游戏与代理管理",
        "position": "游戏与代理管理用于维护游戏入口、游戏分类、子游戏、推荐频道、游戏数据、代理资料、代理规则和代理结算。该模块连接前台游戏体验、代理业务拓展和后续结算报表。",
        "pages": [
            "游戏配置：维护游戏类型设置、游戏频道管理、子游戏列表和推荐频道管理。",
            "游戏数据：查看游戏统计、游戏报表、游戏记录、游戏人数报表和游戏厂商报表。",
            "代理配置：维护代理列表、代理配置、返佣规则、占比分红规则和代理状态。",
            "代理结算：查看代理数据查询、代理领取记录、代理结算明细和代理返佣报表。",
        ],
        "fields": [
            "游戏字段：包含游戏名称、厂商、类型、图标、排序、上下架状态、推荐状态和展示频道。",
            "数据字段：包含投注人数、投注金额、有效投注、输赢、订单量、游戏局号和统计时间。",
            "代理字段：包含代理账号、层级、状态、返佣比例、结算周期、直属会员和业绩条件。",
            "结算字段：包含结算周期、结算金额、领取状态、业绩明细、返佣明细和处理备注。",
        ],
        "flows": [
            "新增或调整游戏前，先确认游戏类型、厂商、频道和前台展示位置。",
            "游戏上下架或排序调整后，检查活动、推荐频道和前台入口是否仍然匹配。",
            "代理规则上线前，先核算返佣比例、结算周期和适用层级。",
            "代理结算时结合代理数据、领取记录和结算明细核对金额。",
        ],
        "notes": [
            "下架游戏前需确认是否仍有关联活动、推荐频道或未完成订单。",
            "游戏数据统计可能存在同步延迟，处理争议时应结合游戏记录和第三方记录。",
            "代理规则变更会影响后续结算，应明确生效时间和适用范围。",
        ],
        "faqs": [
            ("前台看不到某个游戏", "检查游戏上下架状态、频道关联、排序、推荐状态和前台模板配置。"),
            ("游戏报表和记录不一致", "确认统计时间、时区、数据同步延迟和是否筛选了同一游戏厂商。"),
            ("代理结算金额有争议", "核对代理配置、统计周期、直属会员业绩、返佣报表和结算明细。"),
        ],
    },
    {
        "chapter": "八、报表与数据分析",
        "position": "报表与数据分析用于查看经营总览、用户增长、留存、LTV、活动效果、游戏表现、财务指标和代理数据。该模块为运营复盘、财务核对、投放优化和管理决策提供数据依据。",
        "pages": [
            "综合报表：查看综合数据报表、数据排行、充值分布、渠道媒体报表等经营总览数据。",
            "用户分析：查看用户分析、用户分群查询、用户分群报表和在线/活跃表现。",
            "留存/LTV：查看用户留存、付费留存率报表和 LTV，用于评估用户质量与投放回收。",
            "活动报表：查看活动统计报表、玩家活动报表、任务统计报表、玩家任务报表和优惠数据报表。",
            "游戏报表：查看投注报表、游戏报表、游戏记录、游戏厂商报表、游戏人数报表和游戏获利监控相关数据。",
            "财务报表：查看充值、提现、账变、支付通道数据、商户账单和余额变动明细。",
        ],
        "fields": [
            "筛选字段：包含时间范围、站点、渠道、会员、活动、游戏、厂商、代理和订单状态。",
            "指标字段：包含新增、活跃、充值、提现、投注、有效投注、输赢、优惠成本、留存率和 LTV。",
            "维度字段：包含日期、渠道、用户分群、活动、任务、游戏、厂商、代理、VIP 等级和订单类型。",
            "导出字段：包含导出范围、导出格式、导出权限和数据口径说明。",
        ],
        "flows": [
            "先确认统计目的，再选择对应报表和固定统计时间范围。",
            "设置筛选维度后点击查询，优先查看汇总指标，再进入明细或排行定位对象。",
            "需要复盘时记录时间口径、筛选条件和数据来源，避免跨报表混用口径。",
            "需要导出时先缩小范围，确认数据量和权限后再导出。",
        ],
        "notes": [
            "不同报表可能存在统计口径、刷新频率和时区差异，跨页面对比前需先确认口径。",
            "报表数据用于经营分析时，应固定同一时间范围和同一筛选条件。",
            "排行类数据适合发现重点对象，不建议单独作为风控或处罚依据。",
        ],
        "faqs": [
            ("报表数据和订单明细不一致", "检查统计时间、时区、状态口径、数据延迟和筛选条件。"),
            ("导出数据过大或失败", "缩短时间范围、减少筛选维度，并确认当前账号是否有导出权限。"),
            ("活动效果看哪个报表", "先看活动统计报表和优惠数据报表，再结合用户分析、充值分布和 LTV 判断后续价值。"),
        ],
    },
    {
        "chapter": "九、风控管理",
        "position": "风控管理用于识别和处理玩家侧风险，包括黑名单、刷子监控、提现风控和游戏获利异常。该模块重点关注玩家行为、活动套利、资金风险和异常获利，不包含后台账号安全或人事账号管理。",
        "pages": [
            "黑名单：维护高风险玩家、设备、IP 或其他识别对象，控制其登录、参与活动或提现等行为。",
            "刷子监控：识别异常薅羊毛、套利、批量注册、活动滥用等行为。",
            "提现风控：处理命中风控规则的提现申请，并记录审核结论。",
            "游戏获利监控：监控玩家在游戏中的异常获利、异常频率或异常投注表现。",
            "风险排查说明：结合玩家资料、充值提现、账变、游戏记录和活动记录判断风险结果。",
        ],
        "fields": [
            "风险对象：包含会员账号、UID、IP、设备、银行卡、手机号、代理账号和命中规则。",
            "风险状态：包含待处理、处理中、已通过、已驳回、已加入黑名单和已解除。",
            "证据字段：包含命中原因、关联订单、资金流水、登录记录、操作记录和处理备注。",
            "处理字段：包含处理人、处理时间、处理方式、驳回原因、解除原因和复核备注。",
        ],
        "flows": [
            "收到风险提示后，先定位玩家、订单、活动、游戏记录或提现申请。",
            "结合黑名单、刷子监控、提现风控、账变记录、充值提现记录和游戏记录交叉核对。",
            "根据证据执行通过、驳回、加入黑名单、解除限制或提交技术排查。",
            "处理完成后填写明确备注，便于后续审计和复盘。",
        ],
        "notes": [
            "异常获利或命中监控不等同违规，需结合规则、证据和历史行为判断。",
            "加入黑名单或驳回提现属于高影响动作，应保留清晰处理依据。",
            "操作日志是追溯误操作和配置变更的重要依据，不应随意删除或忽略。",
        ],
        "faqs": [
            ("用户被误判怎么办", "复核命中规则、历史行为、资金流水和人工备注，确认后按权限解除限制。"),
            ("提现风控需要看哪些信息", "重点核对会员资料、打码要求、账变、充值提现历史、黑名单和刷子监控记录。"),
            ("游戏获利异常怎么判断", "结合游戏记录、投注频率、输赢波动、充值提现和历史行为综合判断。"),
        ],
    },
    {
        "chapter": "十、人事中心",
        "position": "人事中心用于维护后台账号、选择系统预设角色，并查看管理员登录日志、管理员操作日志和系统异常日志。商户侧不自行配置角色权限，只能在已设置好的角色中为后台账号选择适用角色。",
        "pages": [
            "账号管理：新增、编辑、停用后台账号，维护姓名、登录账号、联系方式、状态和角色。",
            "角色选择说明：商户创建或维护后台账号时，只能选择系统已配置好的角色；不同角色对应不同菜单、按钮和数据范围。",
            "修改密码：后台账号按规则修改自身密码，或由管理员按流程重置。",
            "管理员登录日志：查看后台账号登录时间、IP、登录结果和异常登录情况。",
            "管理员操作日志：追踪后台关键配置、资金、活动、账号和审核操作。",
            "系统异常日志：查看系统异常记录，辅助定位功能异常或接口问题。",
        ],
        "fields": [
            "账号字段：包含登录账号、姓名、手机号、邮箱、状态、角色、创建时间和最近登录时间。",
            "角色字段：包含角色名称、角色说明、菜单范围和数据范围；权限由系统预设，商户只负责选择。",
            "登录日志字段：包含账号、登录时间、IP、浏览器、登录结果和失败原因。",
            "操作日志字段：包含操作人、操作时间、模块、动作、操作对象、操作前后内容和 IP。",
            "异常日志字段：包含异常时间、异常模块、异常等级、错误内容和处理状态。",
        ],
        "flows": [
            "新增后台账号时，填写账号基础信息并选择适合岗位的预设角色。",
            "账号离职、转岗或不再使用时，及时停用账号或调整角色。",
            "出现登录异常时，先查看管理员登录日志，再结合账号状态和 IP 判断是否存在风险。",
            "需要追踪后台操作时，进入管理员操作日志，按账号、模块、对象和时间范围筛选。",
        ],
        "notes": [
            "商户不能自行修改角色权限；如角色权限不满足业务需要，应联系平台管理员或上级管理方调整。",
            "账号停用前应确认是否仍有待处理业务，避免影响日常操作。",
            "管理员操作日志是问题追溯和审计的重要依据，处理争议时应保留关键日志信息。",
        ],
        "faqs": [
            ("新账号看不到菜单", "检查账号状态、选择的角色、站点范围和该角色是否包含对应菜单。"),
            ("角色权限不够怎么办", "商户侧不能自行改权限，需要联系平台管理员或上级管理方调整预设角色。"),
            ("如何排查异常登录", "查看管理员登录日志中的 IP、时间、结果和失败原因，必要时停用账号并重置密码。"),
        ],
    },
]


def add_manual_overview(doc: Document):
    add_heading(doc, "一、手册说明", 1)
    for title, body in [
        ("适用对象", "本手册适用于商户后台的运营、客服、财务、风控、代理管理、数据分析和管理员等后台使用人员。"),
        ("后台角色说明", "不同角色只能访问自身权限范围内的菜单和数据。运营侧重活动、模板、消息和内容；财务侧重充值、提现、账变和通道；风控侧重黑名单、刷子监控和审核；管理员侧重账号、权限和日志。"),
        ("常见操作约定", "查询用于按条件筛选数据；重置用于清空筛选条件；新增/编辑通常在弹窗或表单中完成；启用/停用控制配置是否生效；导出需先确认数据范围；审核类操作必须填写明确处理意见。"),
        ("数据与权限说明", "后台数据受站点、角色、菜单权限和数据权限共同控制。同一页面在不同账号下看到的按钮、字段和数据范围可能不同。报表数据可能存在统计延迟，资金和订单问题应以订单明细与账变记录交叉核对。"),
    ]:
        add_heading(doc, title, 2)
        add_para(doc, body)


def add_login_navigation(doc: Document, images: dict[str, Path]):
    add_heading(doc, "二、登录与基础导航", 1)
    add_heading(doc, "登录", 2)
    add_bullet(doc, "打开后台地址进入登录界面，输入账号和密码后登录。登录后仅展示当前角色有权限访问的菜单。")
    add_bullet(doc, "第一次登录会要求绑定谷歌验证器，之后登录需要输入正确的谷歌验证码才可登录成功。")

    add_heading(doc, "菜单导航", 2)
    add_image(doc, images, "login-top-nav.png", "图1：顶部导航栏与基础入口。")
    add_bullet(doc, "左上角菜单图标用于展开或收起页面目录。")
    add_bullet(doc, "页面目录是主要导航入口，可先点击一级模块，再进入对应页面处理业务。")
    add_bullet(doc, "顶部时间区域用于查看当前系统时间及对应时区；商户余额区域用于查看当前商户账户可用额度。")
    add_bullet(doc, "点击“充值”按钮可进入商户充值页面；点击站点下拉框可切换当前操作和查看的站点。")
    add_image(doc, images, "login-toolbar.png", "图2：顶部导航栏右侧区域。")
    for item in [
        "点击主题开关，可切换页面显示模式。",
        "点击声音图标，可开启或关闭系统提示音。",
        "点击背景样式图标，可切换页面背景显示效果。",
        "点击界面尺寸选项，可在默认、中、小、迷你之间切换页面显示尺寸。",
        "点击语言入口，可切换系统显示语言。",
        "点击通知图标，可查看系统通知或消息提醒。",
        "点击站点下拉框，可切换当前操作和查看的站点。",
        "点击头像图标，可退出当前登录账号。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "筛选、表格、分页、弹窗、状态开关等通用操作", 2)
    for item in [
        "筛选：输入时间、账号、状态、渠道、类型等条件后点击查询；结果不符合预期时先重置再重新查询。",
        "表格：表格通常承载列表数据，支持查看状态、金额、时间、操作人和行内操作。",
        "分页：列表数据较多时通过分页切换；排查问题时注意当前页码和每页条数。",
        "弹窗：新增、编辑、审核、确认等操作通常在弹窗中完成，关闭前需确认是否保存。",
        "状态开关：启用后配置生效，停用后配置隐藏或失效；停用前应确认是否影响线上业务。",
    ]:
        add_bullet(doc, item, item.split("：", 1)[0] + "：")


def add_existing_home_and_template(doc: Document, images: dict[str, Path]):
    add_heading(doc, "首页与模板管理", 2)
    add_heading(doc, "首页看板", 3)
    add_para(doc, "登入后台后即可见首页看板，用于快速查看站点当下的即时状况。看板通常包含核心经营指标、趋势、待办事项和快捷入口，适合运营人员每日进入后台后先进行总览。")
    add_image(doc, images, "home-dashboard.png", "图3：首页看板，用于查看站点核心数据与待处理事项。")
    for item in [
        "模块定位：首页看板用于经营总览和异常提醒，不替代具体报表页面的明细核对。",
        "页面说明：重点查看今日、昨日、本月等关键指标，并通过快捷入口进入高频业务页面。",
        "字段说明：常见字段包含注册人数、充值金额、提现金额、投注金额、输赢、待审核订单和异常提醒。",
        "操作流程：进入后台后先查看看板指标，再根据待办或异常入口进入对应模块处理。",
        "注意事项：看板数据通常存在刷新周期，遇到口径差异时以具体报表和订单明细为准。",
        "常见问题：看板数据与报表不同步时，优先检查统计时间、刷新频率、站点筛选和数据口径。",
    ]:
        add_bullet(doc, item, item.split("：", 1)[0] + "：")

    add_heading(doc, "模板管理", 3)
    add_image(doc, images, "template-management.png", "图4：模板管理页面，用于维护首页、登录页和 VIP 页展示方案。")
    add_para(doc, "模板管理用于维护前台页面或业务场景的展示模板，是日常运营中维护前台展示的重要页面。运营人员可配置首页、登录页、VIP 页等不同页面的展示方案，包括组件内容、页面布局、主题预设、主色调、背景、按钮样式等。")
    for item in [
        "模块定位：模板管理负责前台页面装修和展示方案维护，调整后会直接影响用户看到的页面结构。",
        "页面说明：左侧组件区用于切换首页模板、登录模板、VIP 模板并选择组件；中间预览区用于查看页面装修效果；右侧设置区用于配置页面主题或组件属性。",
        "字段说明：常见字段包含组件可见性、组件样式、图片、颜色、跳转链接、排序、播放间隔、按钮样式和背景设置。",
        "操作流程：选择页面类型，添加或选中组件，在右侧配置属性，检查中间预览效果，确认无误后保存并检查前台展示。",
        "注意事项：未选中组件时右侧为页面全局设置；选中组件时右侧为组件属性，不要把全局主题配置误认为组件配置。",
        "常见问题：组件已配置但前台不展示时，检查组件可见性、排序、页面模板、跳转链接和是否保存到当前站点。",
    ]:
        add_bullet(doc, item, item.split("：", 1)[0] + "：")


def add_site_setting_detail(doc: Document, images: dict[str, Path]):
    add_heading(doc, "站点设置", 2)
    add_para(doc, "站点设置用于维护站点基础展示、首页入口、分享卡片和活动入口等前台基础配置。配置保存后会影响前台品牌露出、用户访问入口、社交分享效果和首页运营位展示。")
    details = [
        ("网站图标", "site-icons.png", "图5：网站图标配置卡片，用于维护站点 Logo 与 Favicon。", [
            "Logo：上传或替换站点主 Logo，通常展示在前台页头、登录页或品牌露出位置；需按页面提示控制图片尺寸、格式和文件大小。",
            "Favicon：上传浏览器标签页图标，建议使用 32*32 或等比例图片，保证在浏览器标签、收藏夹等位置清晰展示。",
            "操作方式：点击右上角编辑进入编辑态，替换图片后保存；保存前检查图片是否清晰、是否符合品牌规范。",
        ]),
        ("币种管理", "site-currency.png", "图6：币种管理配置卡片，用于维护站点货币图标和货币符号。", [
            "货币图标：上传前台展示用的币种图标，常用于钱包、金额、充值或活动奖励等场景。",
            "货币符号：配置金额前后展示的币种符号，需与站点实际结算币种保持一致。",
            "操作方式：点击编辑后调整图标或符号，保存前确认币种展示不会与活动、充值、提现页面口径冲突。",
        ]),
        ("分享卡片配置", "site-share-card.png", "图7：分享卡片配置卡片，用于维护 Social Share / Open Graph 分享信息。", [
            "og:title：设置外部渠道分享时展示的标题，建议简洁表达品牌或活动主题。",
            "og:description：设置分享描述文案，用于补充说明站点、活动或推广卖点。",
            "og:image：上传分享预览图，需检查图片比例、清晰度和裁切效果。",
            "og:url：设置分享落地地址，需确认链接可访问且与当前活动或站点入口一致。",
        ]),
        ("侧滑优惠中心", "site-slide-promo.png", "图8：侧滑优惠中心配置卡片，用于维护前台侧滑活动入口。", [
            "活动卡片：每个卡片对应一个侧滑入口，包含标识、名称、路由、排序、图标和启用状态。",
            "新增：点击右上角新增创建新的侧滑入口，填写名称、路由、排序并上传图标。",
            "启用状态：开关开启后前台展示，关闭后隐藏；排序数字越小通常展示越靠前。",
        ]),
        ("首页快捷操作配置", "site-quick-actions.png", "图9：首页快捷操作配置卡片，用于维护首页快捷入口。", [
            "快捷入口：配置首页常用功能入口，如分享赚钱、待领取、利息宝、实时返水、任务中心、VIP 等。",
            "展示信息：每个入口包含名称、路由、排序、颜色、图标和启用状态。",
            "操作方式：通过新增、编辑、删除维护入口列表；保存前检查路由、图标和排序。",
        ]),
        ("首页悬浮按钮配置", "site-float-buttons.png", "图10：首页悬浮按钮配置卡片，用于维护首页左右侧悬浮按钮组。", [
            "悬浮框：按左侧或右侧配置悬浮按钮组，每组包含位置、排序、按钮数和启用状态。",
            "按钮项：每个按钮项可配置图标、名称、跳转目标和开关状态。",
            "操作方式：点击新增创建悬浮框，点击编辑维护按钮项；保存前检查左右位置是否冲突、移动端是否遮挡主要内容。",
        ]),
    ]
    for title, image, caption, bullets in details:
        add_heading(doc, title, 3)
        add_image(doc, images, image, caption)
        for item in bullets:
            add_bullet(doc, item, item.split("：", 1)[0] + "：")


def add_business_flows(doc: Document):
    add_heading(doc, "十一、典型业务流程", 1)
    flows = [
        ("新站点上线流程", [
            "完成站点设置、品牌素材、域名、客服、登录、分享和支付通道配置。",
            "检查首页模板、登录模板、VIP 模板和主要前台入口。",
            "使用测试账号完成登录、充值、提现申请、客服入口、活动入口和分享入口检查。",
            "上线后通过首页看板、综合报表、支付通道数据和异常日志观察运行状态。",
        ]),
        ("发布活动流程", [
            "进入优惠活动或新增活动页面，填写活动类型、时间、参与条件、奖励规则和展示内容。",
            "配置活动入口、票券或任务关联，确认预算、领取限制和目标人群。",
            "保存后在优惠活动列表检查状态，再通过活动报表、优惠数据报表和用户分析跟踪效果。",
        ]),
        ("处理提现流程", [
            "进入提现订单列表，筛选待审核或处理中订单。",
            "打开订单详情，核对会员、金额、账户、打码情况、风控提示和历史资金记录。",
            "无异常时通过；存在风险时转入提现风控审核或驳回，并填写处理原因。",
        ]),
        ("排查用户资金问题", [
            "通过所有会员定位会员，或进入账变记录直接筛选会员账号。",
            "结合充值订单列表、提现订单列表、人工加扣款、账户余额详情和账变记录核对完整资金链路。",
            "需要追责时进入管理员操作日志，按时间、账号和模块筛选相关操作。",
        ]),
        ("查看活动效果", [
            "固定活动名称、统计时间和渠道范围，先查看活动统计报表或优惠数据报表。",
            "结合玩家活动报表、充值分布、用户分析、留存和 LTV 判断活动质量。",
            "记录活动成本、参与人数、转化表现和后续复盘结论。",
        ]),
    ]
    for title, steps in flows:
        add_heading(doc, title, 2)
        for step in steps:
            add_number(doc, step)


def add_global_faq(doc: Document):
    add_heading(doc, "十二、常见问题与排查", 1)
    faqs = [
        ("登录后看不到菜单", "确认账号角色权限、站点权限和菜单权限；如需开通，联系管理员在账号或角色权限中调整。"),
        ("列表查不到数据", "检查时间范围、状态、账号、渠道等筛选条件，必要时重置后重新查询。"),
        ("配置保存后未生效", "确认配置状态、生效时间、适用范围、缓存和当前操作站点。"),
        ("订单状态异常", "查看订单详情，结合支付通道数据、账变记录、操作日志和异常日志核对。"),
        ("报表数据不一致", "确认统计时间、时区、口径、筛选条件和数据延迟。"),
        ("需要追踪误操作", "进入人事中心的管理员操作日志，按账号、时间、模块、对象和操作类型筛选。"),
        ("活动成本异常", "检查活动规则、领取限制、票券核销、人工调整和活动报表口径。"),
        ("用户提现被拦截", "检查提现风控、黑名单、打码要求、资金流水和审核备注。"),
    ]
    for question, answer in faqs:
        add_heading(doc, question, 2)
        add_para(doc, answer)


def add_core_modules(doc: Document, images: dict[str, Path]):
    for module in MODULES:
        add_heading(doc, module["chapter"], 1)
        add_template_block(doc, module)
        if module["chapter"] == "三、商户初始化配置":
            add_site_setting_detail(doc, images)
        if module["chapter"] == "五、日常运营管理":
            add_existing_home_and_template(doc, images)


def main():
    backup = backup_docx()
    images = extract_existing_images()
    doc = Document()
    set_doc_defaults(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    add_title(doc, "商户后台操作手册")
    add_manual_overview(doc)
    add_login_navigation(doc, images)
    add_core_modules(doc, images)
    add_business_flows(doc)
    add_global_faq(doc)

    doc.save(str(DOCX_PATH))
    print(f"backup={backup}")
    print(f"images={len(images)}")
    print(f"output={DOCX_PATH}")


if __name__ == "__main__":
    main()
