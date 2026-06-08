from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"
ASSET_DIR = ROOT / "custom" / "assets" / "manual-payment-channel-config"
LIST_IMAGE = ASSET_DIR / "payment-channel-list.png"
MODAL_IMAGE = ASSET_DIR / "payment-channel-modal.png"

FONT_NAME = "微软雅黑"
BLUE = RGBColor(47, 85, 151)
CAPTION_GRAY = RGBColor(112, 126, 146)


CONTENT = [
    (
        "页面介绍：支付通道配置用于维护后台可用的在线支付平台，支持按支付通道筛选查询，查看通道类型、通道名称、代收渠道、代付渠道、货币类型、创建时间、状态、商户号、使用站点、备注等信息，并通过操作列进入修改或删除。",
        "List Bullet",
    ),
    ("IMAGE", str(LIST_IMAGE), "图：支付通道配置列表。"),
    (
        "新增支付平台：点击“新增支付平台”后，在弹窗中选择支付通道、填写通道名称和货币类型；不同支付通道需要填写的参数略有差异，例如代收渠道、代付渠道、代收商户ID、代收商户Key、代付商户号、代付商户Key等，请根据页面展示的必填标识、占位提示和可选项逐项填写。",
        "List Bullet",
    ),
    ("IMAGE", str(MODAL_IMAGE), "图：新增在线支付平台参数配置。"),
    (
        "渠道选择：代收渠道和代付渠道以复选项方式配置，可按实际接入能力选择 Cashapp、Btcpay、Paypal、Applepay、Googlepay、Card、Chime、ACH 等选项；如某个通道不展示对应渠道或商户参数，表示该支付通道暂不需要填写该项。",
        "List Bullet",
    ),
    (
        "启用配置：新增或编辑时可设置“是否启用”。启用后系统会调用第三方接口检测配置，需确认商户号、Key、渠道、币种、站点范围等信息正确，避免因参数错误导致支付通道不可用。",
        "List Bullet",
    ),
    (
        "站点范围：列表中的“使用站点”用于确认该支付通道对哪些站点生效，显示“全部站点”时代表全部子站点可使用，显示站点编号时仅对指定站点生效。上线前应结合币种、通道状态和站点范围做复核。",
        "List Bullet",
    ),
    (
        "常用操作：需要调整通道信息时点击“修改”，按页面提示补充或修正参数后确认保存；确认通道不再使用时再执行“删除”。删除前应先确认该通道没有正在使用的充值、提现或代付流程。",
        "List Bullet",
    ),
    (
        "注意事项：支付通道参数以页面实时展示为准，不同支付通道、币种和渠道组合可能出现不同必填项；遇到保存失败、检测失败或前台无法支付时，优先检查必填参数、商户密钥、币种、启用状态、使用站点和第三方接口配置。",
        "List Bullet",
    ),
]


def image_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for item in candidates:
        path = Path(item)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def draw_text(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, font, fill=(31, 41, 55)) -> None:
    draw.text(xy, text, font=font, fill=fill)


def text_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill=(31, 41, 55)) -> None:
    left, top, right, bottom = box
    bbox = draw.textbbox((0, 0), text, font=font)
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.text((left + (right - left - width) // 2, top + (bottom - top - height) // 2 - 1), text, font=font, fill=fill)


def rounded(draw: ImageDraw.ImageDraw, box, radius=4, fill=(255, 255, 255), outline=(214, 221, 231), width=1) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def checkbox(draw: ImageDraw.ImageDraw, x: int, y: int, label: str, checked: bool = True) -> None:
    blue = (64, 145, 247)
    if checked:
        rounded(draw, (x, y, x + 14, y + 14), radius=2, fill=blue, outline=blue)
        draw.line((x + 3, y + 7, x + 6, y + 10, x + 11, y + 4), fill=(255, 255, 255), width=2)
    else:
        rounded(draw, (x, y, x + 14, y + 14), radius=2, fill=(255, 255, 255), outline=(194, 202, 214))
    draw_text(draw, (x + 22, y - 2), label, image_font(13), (22, 119, 255))


def create_list_image() -> None:
    width, height = 1658, 482
    img = Image.new("RGB", (width, height), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    f = image_font(13)
    f_bold = image_font(13, True)
    blue = (22, 119, 255)
    border = (229, 233, 240)
    head = (248, 249, 251)
    muted = (149, 160, 176)

    draw_text(draw, (22, 12), "在线支付类型配置", f, blue)
    draw.line((0, 32, width, 32), fill=(229, 232, 238), width=1)
    draw_text(draw, (8, 62), "支付通道:", f)
    rounded(draw, (68, 49, 230, 80), fill=(255, 255, 255), outline=(213, 220, 230))
    draw_text(draw, (78, 58), "请选择支付通道", f, muted)
    draw.polygon([(213, 60), (223, 60), (218, 67)], fill=(190, 198, 210))
    rounded(draw, (245, 49, 314, 80), fill=(36, 142, 255), outline=(36, 142, 255))
    text_center(draw, (245, 49, 314, 80), "搜索", f_bold, (255, 255, 255))
    rounded(draw, (335, 49, 403, 80), fill=(255, 255, 255), outline=(213, 220, 230))
    text_center(draw, (335, 49, 403, 80), "重置", f, (75, 85, 99))
    rounded(draw, (1551, 49, 1654, 80), fill=(36, 142, 255), outline=(36, 142, 255))
    text_center(draw, (1551, 49, 1654, 80), "新增支付平台", f_bold, (255, 255, 255))

    headers = ["通道类型", "通道名称", "代收渠道", "代付渠道", "货币类型", "创建时间", "状态", "商户号", "使用站点", "备注", "操作"]
    widths = [150, 144, 170, 170, 125, 180, 126, 162, 144, 168, 128]
    rows = [
        ["ercpay", "发过誓", "-", "-", "USD", "2026-06-06 14:07:25", "停用", "", "109", "", "修改   删除"],
        ["succuspay", "抖音", "Cashapp, Btcpay, Paypal, ...", "Paypal, Card, Cashapp, Ch...", "USD", "2026-06-06 14:06:25", "启用", "345623457", "全部站点", "", "修改   删除"],
        ["itransfer", "人工", "-", "-", "USD", "2026-06-06 14:01:08", "启用", "", "全部站点", "", "修改   删除"],
        ["UsdtPay", "微信", "TRC20, ERC20", "TRC20", "BRL", "2026-06-06 13:53:32", "启用", "", "101", "微信", "修改   删除"],
        ["UsdtPay", "usdt1", "-", "-", "MXN", "2026-06-06 13:46:01", "启用", "", "全部站点", "", "修改   删除"],
        ["ercpay", "USDTERC", "-", "-", "USD", "2026-05-28 18:46:41", "启用", "", "101, 102, 103", "", "修改   删除"],
        ["itransfer", "Itransfer", "-", "-", "USD", "2026-05-27 13:05:55", "启用", "", "101", "", "修改   删除"],
        ["manual", "Manual Deposit", "Manual Deposit", "-", "USD", "2026-05-21 16:32:46", "启用", "", "全部站点", "test", "修改   删除"],
    ]
    x0, y0 = 8, 87
    row_h = 43
    x = x0
    for w, header in zip(widths, headers):
        draw.rectangle((x, y0, x + w, y0 + row_h), fill=head, outline=border)
        text_center(draw, (x, y0, x + w, y0 + row_h), header, f_bold)
        x += w
    y = y0 + row_h
    for row in rows:
        x = x0
        for w, cell in zip(widths, row):
            draw.rectangle((x, y, x + w, y + row_h), fill=(255, 255, 255), outline=border)
            color = (82, 196, 26) if cell == "启用" else (255, 77, 79) if cell == "停用" else blue if cell in ("109", "101", "全部站点", "101, 102, 103", "修改   删除") else (37, 48, 69)
            text_center(draw, (x + 4, y, x + w - 4, y + row_h), cell, f, color)
            x += w
        y += row_h
    img.save(LIST_IMAGE)


def create_modal_image() -> None:
    width, height = 633, 759
    img = Image.new("RGB", (width, height), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    f = image_font(13)
    f_bold = image_font(13, True)
    f_small = image_font(12)
    red = (255, 77, 79)
    blue = (22, 119, 255)
    border = (214, 221, 231)
    muted = (166, 176, 190)

    draw.rectangle((0, 0, width - 1, height - 1), outline=(148, 155, 166), width=2)
    draw.line((0, 40, width, 40), fill=border, width=1)
    draw_text(draw, (10, 16), "新增在线支付平台", f_bold)
    draw_text(draw, (605, 10), "×", image_font(24), (79, 88, 101))

    def field(y: int, label: str, value: str = "", required: bool = False, textarea: bool = False, select: bool = False) -> None:
        if required:
            draw_text(draw, (18, y + 11), "*", f, red)
            label_x = 31
        else:
            label_x = 18
        draw_text(draw, (label_x, y + 11), label, f)
        h = 69 if textarea else 31
        rounded(draw, (127, y, 617, y + h), fill=(255, 255, 255), outline=border)
        draw_text(draw, (134, y + 9), value, f, muted if value.startswith("请输入") else (37, 48, 69))
        if select:
            draw.polygon([(599, y + 11), (611, y + 11), (605, y + 18)], fill=(190, 198, 210))

    field(62, "支付通道", "SuccusPay", True, select=True)
    field(113, "通道名称", "请输入自定义通道名称", True)
    field(164, "货币类型", "USD", False, select=True)

    draw_text(draw, (18, 209), "代收渠道", f)
    draw.rectangle((127, 210, 617, 267), fill=(247, 248, 250), outline=None)
    for x, y, label in [(138, 220, "Cashapp"), (299, 220, "Btcpay"), (460, 220, "Paypal"), (138, 245, "Applepay"), (299, 245, "Googlepay")]:
        checkbox(draw, x, y, label)

    draw_text(draw, (18, 282), "代付渠道", f)
    draw.rectangle((127, 276, 617, 337), fill=(247, 248, 250), outline=None)
    for x, y, label in [(138, 286, "Paypal"), (299, 286, "Card"), (460, 286, "Cashapp"), (138, 311, "Chime"), (299, 311, "ACH")]:
        checkbox(draw, x, y, label)

    field(356, "代收商户ID", "请输入代收商户ID", True)
    field(406, "代收商户Key", "请输入代收商户Key", True)
    field(456, "代付商户号", "请输入代付商户号", True)
    field(506, "代付商户Key", "请输入代付商户Key", True)
    field(554, "备注", "请输入备注", False, textarea=True)

    draw_text(draw, (18, 661), "是否启用", f)
    draw.rounded_rectangle((132, 642, 170, 662), radius=10, fill=(160, 163, 168))
    draw.ellipse((134, 644, 150, 660), fill=(255, 255, 255))
    draw_text(draw, (126, 679), "提示：启用后将会调用第三方接口检测配置", f, red)
    rounded(draw, (503, 719, 555, 750), fill=(255, 255, 255), outline=border)
    text_center(draw, (503, 719, 555, 750), "取消", f, (75, 85, 99))
    rounded(draw, (572, 719, 625, 750), fill=(36, 142, 255), outline=(36, 142, 255))
    text_center(draw, (572, 719, 625, 750), "确定", f_bold, (255, 255, 255))
    img.save(MODAL_IMAGE)


def create_images() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    create_list_image()
    create_modal_image()


def set_run_font(run, *, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(10)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def style_paragraph(paragraph) -> None:
    text = paragraph.text
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    if "：" in text:
        label, body = text.split("：", 1)
        label_run = paragraph.add_run(label + "：")
        set_run_font(label_run, bold=True, color=BLUE)
        body_run = paragraph.add_run(body)
        set_run_font(body_run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_paragraph_after(paragraph, text: str, style: str):
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph.style = style
    new_paragraph.add_run(text)
    paragraph._element.addnext(new_paragraph._element)
    style_paragraph(new_paragraph)
    return new_paragraph


def insert_image_after(paragraph, image_path: str, caption: str):
    image_paragraph = paragraph._parent.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_run = image_paragraph.add_run()
    image_run.add_picture(image_path, width=Inches(6.35))
    paragraph._element.addnext(image_paragraph._element)

    caption_paragraph = paragraph._parent.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, color=CAPTION_GRAY)
    caption_run.font.size = Pt(9)
    image_paragraph._element.addnext(caption_paragraph._element)
    return caption_paragraph


def find_heading_index(document: Document, title: str) -> int:
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name.startswith("Heading") and paragraph.text.strip() == title:
            return index
    raise ValueError(f"未找到章节：{title}")


def main() -> None:
    create_images()
    BACKUP_DIR.mkdir(exist_ok=True)
    backup = BACKUP_DIR / f"B端后台操作手册.payment-channel-config-{datetime.now():%Y%m%d_%H%M%S}.bak.docx"
    shutil.copy2(DOCX, backup)

    document = Document(DOCX)
    heading_index = find_heading_index(document, "支付通道配置")
    paragraphs = document.paragraphs

    end_index = len(paragraphs)
    for index in range(heading_index + 1, len(paragraphs)):
        if paragraphs[index].style.name.startswith("Heading"):
            end_index = index
            break

    heading = paragraphs[heading_index]
    for paragraph in list(paragraphs[heading_index + 1 : end_index]):
        remove_paragraph(paragraph)

    cursor = heading
    for item in CONTENT:
        if item[0] == "IMAGE":
            cursor = insert_image_after(cursor, item[1], item[2])
        else:
            text, style = item
            cursor = insert_paragraph_after(cursor, text, style)

    document.save(DOCX)
    print(f"backup={backup}")
    print("updated=支付通道配置")


if __name__ == "__main__":
    main()
