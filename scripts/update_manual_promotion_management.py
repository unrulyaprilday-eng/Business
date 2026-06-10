from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"

TITLE = "宣传管理"
FONT_NAME = "微软雅黑"
BLUE = RGBColor(47, 85, 151)
CAPTION_GRAY = RGBColor(112, 126, 146)


def set_run(run, *, size: float = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def insert_after(paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_body_after(anchor: Paragraph, text: str) -> Paragraph:
    paragraph = insert_after(anchor, style="Normal")
    run = paragraph.add_run(text)
    set_run(run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullet_after(anchor: Paragraph, label: str, text: str) -> Paragraph:
    paragraph = insert_after(anchor, style="List Bullet")
    label_run = paragraph.add_run(label)
    set_run(label_run, size=10, bold=True, color=BLUE)
    body_run = paragraph.add_run("：" + text)
    set_run(body_run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_image_placeholder_after(anchor: Paragraph, caption: str) -> Paragraph:
    placeholder = insert_after(anchor, style="Normal")
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = placeholder.add_run("[图片]")
    set_run(run, size=10, color=CAPTION_GRAY)
    placeholder.paragraph_format.space_after = Pt(1)

    caption_para = insert_after(placeholder, style="Caption")
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_para.add_run(caption)
    set_run(caption_run, size=9, color=CAPTION_GRAY)
    caption_para.paragraph_format.space_after = Pt(6)
    return caption_para


def reset_section(document: Document) -> Paragraph:
    start = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == TITLE:
            start = index
            break
    if start is None:
        raise RuntimeError(f"未找到《{TITLE}》章节。")

    end = len(document.paragraphs)
    for index in range(start + 1, len(document.paragraphs)):
        style_name = document.paragraphs[index].style.name
        if style_name in {"Heading 1", "Heading 2", "Heading 3"}:
            end = index
            break

    for paragraph in list(document.paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)

    anchor = document.paragraphs[start]
    for run in anchor.runs:
        set_run(run, size=10.5, bold=True, color=BLUE)
    return anchor


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = BACKUP_DIR / f"{DOCX_PATH.stem}.promotion-management-{stamp}.bak.docx"
    shutil.copy2(DOCX_PATH, backup_path)

    document = Document(DOCX_PATH)
    anchor = reset_section(document)

    last = add_body_after(
        anchor,
        "宣传管理用于维护站点大厅 Banner 等宣传素材的展示顺序、展示图片和跳转内容，是运营统一配置首页宣传位的入口。新增、修改或删除宣传内容后，会直接影响前台宣传位展示，提交前应复核图片规格、跳转目标、活动有效期和排序结果。",
    )

    last = add_image_placeholder_after(
        last,
        "图：宣传管理列表页面，用于查看宣传记录、筛选条件和行内操作。",
    )

    for label, text in [
        ("页面操作", "页面顶部提供宣传类型筛选、搜索、重置和新增入口，适合先按宣传类型定位目标宣传位，再执行维护操作。"),
        ("筛选查询", "选择宣传类型后点击“搜索”刷新列表；点击“重置”可清空筛选条件并恢复全部记录。"),
        ("列表字段", "列表展示排序、类型、宣传图、路径链接、创建时间和操作等信息，便于快速确认当前展示内容。"),
        ("排序", "排序值通常用于控制前台展示顺序，数值越小越靠前；调整排序后应结合前台展示位再次核对实际顺序。"),
        ("路径链接", "列表中的路径链接用于显示当前宣传图对应的跳转目标，常见内容包括无、外部链接、活动、返水、代理和 VIP 等。"),
        ("修改", "点击“修改”可进入编辑弹窗，调整宣传类型、宣传图片、跳转类型、活动或外部链接以及排序。"),
        ("删除", "点击“删除”可移除当前宣传记录。删除前应确认该宣传位已下线或已有替代素材，避免前台出现空白宣传位。"),
    ]:
        last = add_bullet_after(last, label, text)

    last = add_image_placeholder_after(
        last,
        "图：宣传管理新增弹窗，用于配置宣传类型、上传宣传图片和设置跳转内容。",
    )

    for label, text in [
        ("宣传类型", "新增时应先选择宣传类型。不同宣传类型对应的图片尺寸要求可能不同，先选类型再上传图片可减少返工。"),
        ("宣传图片", "宣传图片为核心展示素材，上传前应检查清晰度、比例和文案是否正确，避免出现裁切、拉伸或旧活动图未替换的情况。"),
        ("跳转类型", "跳转类型支持无、外部链接、活动、返水、代理和 VIP 等配置，应按实际落地页面选择对应类型。"),
        ("活动选择", "当跳转类型选择活动时，应从活动列表中选择对应活动，并确认活动状态、名称和有效期无误。"),
        ("外部链接", "当跳转类型选择外部链接时，应填写完整可访问地址，并检查是否能够正常打开目标页面。"),
        ("排序设置", "弹窗中的排序字段用于维护当前宣传位的展示先后顺序，保存前应与同类宣传位统一核对。"),
        ("保存与取消", "点击“确定”后保存本次新增或修改；点击“取消”或关闭弹窗会放弃当前未保存内容。"),
        ("发布检查", "发布前重点检查宣传图片、跳转目标、活动配置、排序值和前台展示效果，避免出现图片正确但跳转错误的情况。"),
        ("维护提示", "修改已上线宣传素材时，应同步关注相关活动、落地页和公告排期，确保前台宣传内容与实际运营安排一致。"),
    ]:
        last = add_bullet_after(last, label, text)

    try:
        document.save(DOCX_PATH)
    except PermissionError:
        print(f"PERMISSION_ERROR backup={backup_path}")
        raise

    print(f"updated promotion management; backup={backup_path}")


if __name__ == "__main__":
    main()
