# -*- coding: utf-8 -*-
from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(".")
DOCX_PATH = BASE / "B端后台操作手册.docx"
ASSET_DIR = BASE / "custom" / "assets" / "manual-reorganized"
SOURCE_IMAGE_DOCX_CANDIDATES = [
    BASE / "B端后台操作手册.docx.nav-20260604-182335.bak",
    BASE / "B端后台操作手册.docx.labels-20260604-180329.bak",
    BASE / "B端后台操作手册.docx.labels-20260604-180217.bak",
    BASE / "B端后台操作手册.docx.menu-20260605-101251.bak",
    BASE / "B端后台操作手册.docx.bak",
    DOCX_PATH,
]
FONT_NAME = "微软雅黑"
DEEP_BLUE = RGBColor(37, 64, 97)
MUTED = RGBColor(96, 108, 128)

IMAGE_NAMES = [
    "login-top-nav.png",
    "login-toolbar.png",
    "home-dashboard.png",
    "template-management.png",
    "site-icons.png",
    "site-currency.png",
    "site-share-card.png",
    "site-slide-promo.png",
    "site-quick-actions.png",
    "site-float-buttons.png",
]

CAPTION_IMAGE_MAP = {
    "图1": "login-top-nav.png",
    "图2": "login-toolbar.png",
    "图3": "home-dashboard.png",
    "图4": "template-management.png",
    "图5": "site-icons.png",
    "图6": "site-currency.png",
    "图7": "site-share-card.png",
    "图8": "site-slide-promo.png",
    "图9": "site-quick-actions.png",
    "图10": "site-float-buttons.png",
}


MENU_MODULES = [
    ("三、首页", ["默认看板"]),
    ("四、站点配置", ["模板管理", "站点配置", "品牌设置", "站点域名", "客服配置", "分享配置", "第三方登录配置", "支付通道配置", "广告埋点配置", "音乐管理"]),
    ("五、运营中心", ["用户分群思维导图", "站点维护管理", "平台公告", "给平台提建议", "会员投诉列表", "消息推送", "宣传管理", "客服中心", "任务中心", "落地页管理", "活动排行统计"]),
    ("六、优惠活动", ["优惠活动列表", "幸运转盘配置", "新增活动", "活动分析", "优惠数据报表", "票券中心", "票券使用记录", "全局配置"]),
    ("七、财务中心", ["会员提现设置", "会员充值配置", "充值订单列表", "提现订单列表", "提现风控审核", "支付通道数据", "人工加扣款", "账变记录", "资金账变类型", "打码管理"]),
    ("八、数据报表", ["综合数据报表", "游戏人数据表", "优惠数据报表", "充值留存报表", "代理结算明细", "VIP数据统计", "数据排行", "游戏记录", "用户留存", "用户分析", "LTV", "活动统计报表", "充值分布表", "任务统计报表", "用户分群查询"]),
    ("九、游戏中心", ["推荐频道管理", "游戏类型设置", "游戏频道管理", "子游戏列表", "游戏统计", "游戏厂商统计"]),
    ("十、用户管理", ["在线玩家列表", "所有会员", "VIP设置", "有奖反馈"]),
    ("十一、代理中心", ["代理列表", "代理配置", "代理领取记录", "代理数据查询"]),
    ("十二、风险中心", ["黑名单", "刷子监控", "游戏获利监控"]),
    ("十三、商户中心", ["商户信息", "商户充值", "充值记录", "商户账变记录", "商户账单"]),
    ("十四、埋点配置", ["埋点统计", "会话列表", "事件流水"]),
    ("十五、人事中心", ["账号管理", "日志管理", "管理员登录日志", "管理员操作日志", "系统异常日志"]),
]


def backup_docx() -> Path:
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    backup = DOCX_PATH.with_suffix(f".docx.menu-{stamp}.bak")
    shutil.copy2(DOCX_PATH, backup)
    return backup


def extract_existing_images() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    for source in SOURCE_IMAGE_DOCX_CANDIDATES:
        if not source.exists():
            continue
        old_doc = Document(str(source))
        images = extract_images_by_caption(old_doc)
        if all(name in images for name in IMAGE_NAMES):
            print(f"image_source={source}")
            return images
    return images if "images" in locals() else {}


def paragraph_image_blob(doc: Document, paragraph):
    for run in paragraph.runs:
        for drawing in run._element.xpath(".//w:drawing"):
            blips = drawing.xpath(".//a:blip")
            if not blips:
                continue
            rid = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            part = doc.part.related_parts.get(rid)
            if part is not None:
                return part.blob
    return None


def next_caption(paragraphs, start_index: int) -> str:
    for paragraph in paragraphs[start_index + 1 : start_index + 4]:
        text = paragraph.text.strip()
        if text:
            return text
    return ""


def extract_images_by_caption(old_doc: Document) -> dict[str, Path]:
    images = {}
    paragraphs = old_doc.paragraphs
    for index, paragraph in enumerate(paragraphs):
        blob = paragraph_image_blob(old_doc, paragraph)
        if blob is None:
            continue
        caption = next_caption(paragraphs, index)
        key = next((prefix for prefix in CAPTION_IMAGE_MAP if caption.startswith(prefix + "：")), None)
        if key is None:
            continue
        name = CAPTION_IMAGE_MAP[key]
        path = ASSET_DIR / name
        path.write_bytes(blob)
        images[name] = path
    return images


def set_run_font(run, size=10, bold=False, color=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_doc_defaults(doc: Document):
    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(10)
    for style_name, size in [("Heading 1", 14), ("Heading 2", 12.5), ("Heading 3", 10.5), ("Heading 4", 10)]:
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = DEEP_BLUE


def add_title(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=18, bold=True, color=DEEP_BLUE)
    paragraph.paragraph_format.space_after = Pt(12)


def add_heading(doc: Document, text: str, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    paragraph.paragraph_format.space_after = Pt(4)
    for run in paragraph.runs:
        set_run_font(run, size={1: 14, 2: 12.5, 3: 10.5}.get(level, 10.5), bold=True, color=DEEP_BLUE)
    return paragraph


def add_para(doc: Document, text: str, color=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, color=color)
    return paragraph


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, bold=True, color=DEEP_BLUE)
        rest = text[len(bold_prefix):]
        if rest:
            run = paragraph.add_run(rest)
            set_run_font(run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_number(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run)


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size=9, color=MUTED)


def add_image(doc: Document, images: dict[str, Path], name: str, caption: str, width=6.25):
    path = images.get(name)
    if not path or not path.exists():
        add_para(doc, f"图片占位：{caption}", color=MUTED)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_field_bullet(doc: Document, text: str):
    add_bullet(doc, text, text.split("：", 1)[0] + "：" if "：" in text else None)


def add_page_brief(doc: Document, page: str):
    add_heading(doc, page, 3)
    for item in [
        f"页面介绍：用于处理“{page}”相关后台业务，具体说明可在后续补充截图时继续细化。",
        "常用操作：进入页面后先通过筛选条件定位数据或配置项，再执行页面支持的新增、编辑、启用、停用、审核、导出或查看详情等操作。",
        "后续补充：按页面截图补齐关键字段、按钮说明、弹窗流程、保存校验和常见问题。",
    ]:
        add_field_bullet(doc, item)


def add_module_intro(doc: Document):
    add_heading(doc, "模块定位", 2)
    add_para(doc, "本章节按后台真实菜单顺序整理，页面说明与后台左侧菜单保持一致，便于商户按菜单路径查找和操作。")
    add_heading(doc, "页面说明", 2)


def add_manual_overview(doc: Document):
    add_heading(doc, "一、手册说明", 1)
    for title, body in [
        ("适用对象", "本手册适用于商户后台的运营、客服、财务、风控、代理管理、数据分析和管理员等后台使用人员。"),
        ("后台角色说明", "商户创建或维护后台账号时，只能选择系统已配置好的角色；不同角色对应不同菜单、按钮和数据范围。若角色权限不满足业务需要，需要联系平台管理员或上级管理方调整。"),
        ("常见操作约定", "查询用于按条件筛选数据；重置用于清空筛选条件；新增/编辑通常在弹窗或表单中完成；启用/停用控制配置是否生效；导出需先确认数据范围；审核类操作必须填写明确处理意见。"),
        ("数据与权限说明", "后台数据受站点、角色、菜单权限和数据权限共同控制。同一页面在不同账号下看到的按钮、字段和数据范围可能不同。报表数据可能存在统计延迟，资金和订单问题应以订单明细与账变记录交叉核对。"),
    ]:
        add_heading(doc, title, 2)
        add_para(doc, body)


def add_login_navigation(doc: Document, images: dict[str, Path]):
    add_heading(doc, "二、登录与基础导航", 1)
    add_heading(doc, "登录", 2)
    add_bullet(doc, "打开后台地址进入登录界面，输入账号和密码后登录。登录后仅展示当前角色有权限访问的菜单。")
    add_bullet(doc, "第一次登录会要求绑定谷歌验证器，之后登录需要输入正确的谷歌验证码才可登录成功。")
    add_heading(doc, "菜单导航", 2)
    add_image(doc, images, "login-top-nav.png", "图1：顶部导航栏与基础入口。")
    add_bullet(doc, "左上角菜单图标用于展开或收起页面目录。")
    add_bullet(doc, "页面目录是主要导航入口，可先点击一级模块，再进入对应页面处理业务。")
    add_bullet(doc, "顶部时间区域用于查看当前系统时间及对应时区；商户余额区域用于查看当前商户账户可用额度。")
    add_bullet(doc, "点击“充值”按钮可进入商户充值页面；点击站点下拉框可切换当前操作和查看的站点。")
    add_image(doc, images, "login-toolbar.png", "图2：顶部导航栏右侧区域。")
    for item in [
        "点击主题开关，可切换页面显示模式。",
        "点击声音图标，可开启或关闭系统提示音。",
        "点击背景样式图标，可切换页面背景显示效果。",
        "点击界面尺寸选项，可在默认、中、小、迷你之间切换页面显示尺寸。",
        "点击语言入口，可切换系统显示语言。",
        "点击通知图标，可查看系统通知或消息提醒。",
        "点击站点下拉框，可切换当前操作和查看的站点。",
        "点击头像图标，可退出当前登录账号。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "筛选、表格、分页、弹窗、状态开关等通用操作", 2)
    for item in [
        "筛选：输入时间、账号、状态、渠道、类型等条件后点击查询；结果不符合预期时先重置再重新查询。",
        "表格：表格通常承载列表数据，支持查看状态、金额、时间、操作人和行内操作。",
        "分页：列表数据较多时通过分页切换；排查问题时注意当前页码和每页条数。",
        "弹窗：新增、编辑、审核、确认等操作通常在弹窗中完成，关闭前需确认是否保存。",
        "状态开关：启用后配置生效，停用后配置隐藏或失效；停用前应确认是否影响线上业务。",
    ]:
        add_field_bullet(doc, item)


def add_home_page(doc: Document, images: dict[str, Path]):
    add_heading(doc, "三、首页", 1)
    add_heading(doc, "默认看板", 2)
    add_para(doc, "登入后台后即可见首页看板，用于快速查看站点当下的即时状况。")
    add_image(doc, images, "home-dashboard.png", "图3：首页看板，用于查看站点核心数据与待处理事项。")
    for item in [
        "模块定位：首页用于经营总览和待办提醒，是商户进入后台后的默认查看入口。",
        "页面说明：重点查看今日、昨日、本月等关键指标，并通过快捷入口进入高频业务页面。",
        "字段说明：常见字段包含注册人数、充值金额、提现金额、投注金额、输赢、待审核订单和异常提醒。",
        "操作流程：进入后台后先查看看板指标，再根据待办或异常入口进入对应模块处理。",
        "注意事项：看板数据通常存在刷新周期，遇到口径差异时以具体报表和订单明细为准。",
        "常见问题：看板数据与报表不同步时，优先检查统计时间、刷新频率、站点筛选和数据口径。",
    ]:
        add_field_bullet(doc, item)


def add_template_management(doc: Document, images: dict[str, Path]):
    add_heading(doc, "模板管理", 3)
    add_image(doc, images, "template-management.png", "图4：模板管理页面，用于维护首页、登录页和 VIP 页展示方案。")
    add_para(doc, "模板管理用于维护前台页面的装修模板和展示方案，位于站点配置目录下。运营人员可在同一页面内选择首页、登录页、VIP 页模板，添加页面组件，预览手机端效果，并调整页面主题或组件属性。模板保存并启用后，会影响用户在前台看到的页面结构、视觉风格和功能入口。")

    add_heading(doc, "左侧组件区", 4)
    for item in [
        "页面类型：顶部可在首页模板、登录模板、VIP 模板之间切换，不同模板下展示可用的组件或样式。",
        "首页模板：按营销、导航、内容、基础等分组选择组件，例如轮播图、奖池、公告栏、快捷入口、楼层导航、个人引导、底部导航和游戏模块。",
        "登录模板：用于选择登录展示方式，包含弹框模式和页面模式。",
        "VIP模板：用于选择 VIP 页面展示样式，包含模板 A、模板 B、模板 C 等不同布局。",
        "操作方式：点击或拖拽组件到中间预览区；当前选中的 Tab、模板或组件会高亮显示。",
    ]:
        add_field_bullet(doc, item)

    add_heading(doc, "中间页面预览区", 4)
    for item in [
        "页面预览：中间区域用于实时预览当前页面装修效果，顶部可编辑页面名称，并提供创建/更改、撤销、重做等操作。",
        "预览内容：首页模板展示首页布局和已添加组件；登录模板展示弹框登录或页面登录；VIP 模板展示对应 VIP 页面样式。",
        "组件操作：未添加组件时显示添加提示；添加组件后按页面顺序展示。点击组件后可看到选中边框、组件标题以及上移、下移、删除等操作。",
        "联动规则：选中预览区组件后，右侧会同步切换为该组件的属性配置。",
    ]:
        add_field_bullet(doc, item)

    add_heading(doc, "右侧设置区", 4)
    for item in [
        "页面设置：未选中组件时，右侧用于配置页面全局主题和视觉样式，例如主题预设、自定义主色调、页面背景色、页面背景图片、高级颜色配置和按钮样式。",
        "组件属性：选中组件时，右侧用于配置当前组件的展示状态、样式、内容、图片、滚动、播放、颜色、跳转和排序等参数。",
        "应用更改：修改组件属性后，需要点击应用更改或保存，配置才会同步到中间预览区和最终模板。",
        "使用要点：区分页面全局设置和组件属性，避免把右侧全局主题配置误认为某个组件的单独配置。",
    ]:
        add_field_bullet(doc, item)

    add_heading(doc, "组件说明", 4)
    add_para(doc, "组件说明按模板管理中常见组件逐项整理。后续补充截图时，可在对应组件下继续追加字段截图、弹窗说明或配置示例。")
    component_details = [
        ("轮播图（Banner/Carousel）", [
            "功能用途：用于配置首页或活动页 Banner 轮播，是前台最主要的视觉入口之一，常用于展示活动、品牌、游戏推荐或充值引导。",
            "主要配置：组件名称、删除、可见性、轮播样式、自动播放、播放间隔、显示指示器、指示器样式、轮播项列表、圆角和应用更改。",
            "轮播项配置：每一帧可配置图片、跳转链接、打开方式、排序权重和生效时间；多张 Banner 同时配置时，应检查展示顺序和生效时间是否冲突。",
            "使用要点：重点检查图片尺寸、移动端裁切、跳转地址、打开方式和活动有效期；上线前确认自动播放间隔不会影响用户阅读。",
        ]),
        ("奖池（Jackpot）", [
            "功能用途：用于展示奖金池或动态金额，适合放在首页营销区域，用来强化活动氛围和奖金规模感。",
            "主要配置：组件名称、删除、可见性、背景图片、数值变动速度、奖金数字颜色、奖金字体大小、奖金字体粗细、数字动画效果、起始数值、货币符号、千分位分隔符和应用更改。",
            "展示配置：背景图片决定奖池区域视觉效果；数字颜色、字号、字重和动画决定金额可读性；起始数值和变动速度影响前台展示节奏。",
            "使用要点：检查金额格式、货币符号、千分位分隔和背景图对比度，避免数字不清晰或与实际活动口径冲突。",
        ]),
        ("公告栏（Notice/Marquee）", [
            "功能用途：用于展示平台公告、活动提醒、中奖播报或重要通知，通常放在首页靠前位置。",
            "主要配置：组件名称、删除、可见性、滚动方向、滚动速度、背景颜色、文字颜色、消息图标、自定义消息图标、公告内容列表、单行显示行数和应用更改。",
            "公告内容配置：每条公告可配置消息文本、跳转链接和排序；横向滚动适合单行短消息，纵向滚动适合多条公告轮播。",
            "使用要点：发布前检查文字长度、颜色对比、跳转链接和排序；重要公告建议避免滚动过快。",
        ]),
        ("快捷入口（Quick Entry）", [
            "功能用途：用于配置首页常用功能入口，帮助用户快速进入充值、提现、活动、客服、VIP、任务等高频页面。",
            "主要配置：组件名称、删除、可见性、展示样式、每行数量、行数、入口项列表、图标大小、文字大小和应用更改。",
            "入口项配置：每个入口可配置入口名称、图标、选中图标、跳转信息、打开方式、角标、排序权重和单项可见性。",
            "使用要点：检查入口名称是否简洁、图标是否匹配业务含义、跳转目标是否正确、角标是否误导用户，并控制入口数量避免首页拥挤。",
        ]),
        ("楼层导航（Floor Navigation）", [
            "功能用途：用于页面内分区导航，帮助用户快速跳转到指定楼层或组件区域，适合内容较长的首页或活动页。",
            "主要配置：组件名称、删除、可见性、展示样式、吸顶偏移、导航项列表、背景颜色、文字颜色、选中文字颜色、滚动动画和应用更改。",
            "导航项配置：每个导航项可配置楼层名称、锚点目标、图标、选中态样式和排序；锚点目标应与页面组件或区域 ID 对应。",
            "使用要点：检查锚点是否能准确定位、吸顶偏移是否遮挡内容、选中态是否清晰；移动端尤其要关注导航宽度和滚动体验。",
        ]),
        ("个人引导（User Guide / Profile Entry）", [
            "功能用途：用于展示用户登录前后的引导区域，承担登录注册引导、用户信息展示和快捷功能入口。",
            "主要配置：组件名称、删除、可见性、未登录展示内容、已登录展示内容、背景图和应用更改。",
            "未登录配置：可配置标题文案、副标题、登录按钮文案、注册按钮文案、按钮样式和默认头像。",
            "已登录配置：可配置是否显示用户昵称、用户头像、VIP 等级、账户余额、快捷功能和跳转信息。",
            "使用要点：未登录态重点检查登录/注册按钮是否清晰；已登录态重点检查余额、VIP、快捷入口和个人中心跳转是否准确。",
        ]),
        ("底部导航（Bottom Navigation / TabBar）", [
            "功能用途：用于配置移动端底部 Tab 栏，是用户切换首页、活动、游戏、会员中心等核心页面的主导航。",
            "主要配置：组件名称、删除、可见性、背景颜色、背景图片、文字颜色、选中文字颜色、导航项列表、中间突出项、中间突出图标、分隔线和应用更改。",
            "导航项配置：每个 Tab 可配置名称、默认图标、选中图标、跳转页面、角标、排序和单项可见性。",
            "使用要点：通常控制在 3 到 5 个 Tab；检查默认图标和选中图标是否成套、跳转页面是否正确、中间突出项是否遮挡内容。",
        ]),
        ("游戏模块（Game Module）", [
            "功能用途：用于展示游戏分类和游戏列表，是首页承接用户进入游戏的重要区域。",
            "主要配置：组件名称、删除、可见性、游戏分类、分类 Tab 展示、分类 Tab 样式、展示数量、排列样式、卡片样式、卡片圆角、加载更多、排序规则、搜索功能、收藏筛选、厂商筛选、空态提示和应用更改。",
            "展示配置：可选择热门、老虎机、真人、棋牌、体育、捕鱼等分类；卡片可按图片、名称、厂商等不同信息组合展示。",
            "使用要点：检查游戏分类、展示数量、排序规则和加载更多方式；无游戏时需配置清晰空态提示，避免前台出现空白区域。",
        ]),
        ("登录模板组件（Login Template）", [
            "功能用途：用于配置登录页或登录弹窗的展示方式和登录相关入口。",
            "主要配置：组件名称、删除、可见性、登录展示方式、登录方式、默认登录方式、登录按钮样式、注册按钮样式、第三方登录入口、公告弹框样式、返回顶部、记住密码、忘记密码入口、背景图/颜色、Logo 图片和应用更改。",
            "按钮与入口配置：登录按钮可配置文案、颜色、圆角和宽度；注册按钮可配置文案、颜色和展示位置；第三方登录入口可配置是否显示、支持平台和展示样式。",
            "公告与背景配置：公告弹框可配置标题、内容、图片和关闭方式；页面背景可使用图片或颜色，Logo 图片需与品牌设置保持一致。",
            "使用要点：检查默认登录方式、第三方登录入口、忘记密码入口和公告弹框是否符合当前站点登录策略。",
        ]),
        ("VIP模板组件（VIP Template）", [
            "功能用途：用于配置 VIP 页面整体布局、等级展示、权益说明和奖励领取入口。",
            "主要配置：组件名称、删除、可见性、模板布局、等级信息展示、升级条件、权益内容、奖励展示、规则说明和应用更改。",
            "等级配置：可配置当前等级展示、等级图标、等级名称、等级进度条、等级背景色或背景图。",
            "权益与奖励配置：可配置权益列表、权益图标、权益描述、适用等级、权益对比表、升级奖励、周/月奖励、返水比例和奖励领取按钮。",
            "使用要点：检查等级名称、升级条件、权益适用等级、奖励领取按钮和规则说明是否与实际 VIP 规则一致。",
        ]),
    ]
    for title, bullets in component_details:
        add_heading(doc, title, 4)
        for item in bullets:
            add_field_bullet(doc, item)
    add_field_bullet(doc, "上线检查：保存模板前重点检查模板类型、组件顺序、图片尺寸、跳转链接、颜色风格、按钮效果和手机端预览展示。")


def add_site_settings(doc: Document, images: dict[str, Path]):
    add_heading(doc, "站点配置", 3)
    add_para(doc, "站点配置用于维护站点基础展示、首页入口、分享卡片和活动入口等前台基础配置。")
    details = [
        ("网站图标", "site-icons.png", "图5：网站图标配置卡片，用于维护站点 Logo 与 Favicon。", [
            "Logo：上传或替换站点主 Logo，通常展示在前台页头、登录页或品牌露出位置；需按页面提示控制图片尺寸、格式和文件大小。",
            "Favicon：上传浏览器标签页图标，建议使用 32*32 或等比例图片，保证在浏览器标签、收藏夹等位置清晰展示。",
            "操作方式：点击右上角编辑进入编辑态，替换图片后保存；保存前检查图片是否清晰、是否符合品牌规范。",
        ]),
        ("币种管理", "site-currency.png", "图6：币种管理配置卡片，用于维护站点货币图标和货币符号。", [
            "货币图标：上传前台展示用的币种图标，常用于钱包、金额、充值或活动奖励等场景。",
            "货币符号：配置金额前后展示的币种符号，需与站点实际结算币种保持一致。",
            "操作方式：点击编辑后调整图标或符号，保存前确认币种展示不会与活动、充值、提现页面口径冲突。",
        ]),
        ("分享卡片配置", "site-share-card.png", "图7：分享卡片配置卡片，用于维护 Social Share / Open Graph 分享信息。", [
            "og:title：设置外部渠道分享时展示的标题，建议简洁表达品牌或活动主题。",
            "og:description：设置分享描述文案，用于补充说明站点、活动或推广卖点。",
            "og:image：上传分享预览图，需检查图片比例、清晰度和裁切效果。",
            "og:url：设置分享落地地址，需确认链接可访问且与当前活动或站点入口一致。",
        ]),
        ("侧滑优惠中心", "site-slide-promo.png", "图8：侧滑优惠中心配置卡片，用于维护前台侧滑活动入口。", [
            "活动卡片：每个卡片对应一个侧滑入口，包含标识、名称、路由、排序、图标和启用状态。",
            "新增：点击右上角新增创建新的侧滑入口，填写名称、路由、排序并上传图标。",
            "启用状态：开关开启后前台展示，关闭后隐藏；排序数字越小通常展示越靠前。",
        ]),
        ("首页快捷操作配置", "site-quick-actions.png", "图9：首页快捷操作配置卡片，用于维护首页快捷入口。", [
            "快捷入口：配置首页常用功能入口，如分享赚钱、待领取、利息宝、实时返水、任务中心、VIP 等。",
            "展示信息：每个入口包含名称、路由、排序、颜色、图标和启用状态。",
            "操作方式：通过新增、编辑、删除维护入口列表；保存前检查路由、图标和排序。",
        ]),
        ("首页悬浮按钮配置", "site-float-buttons.png", "图10：首页悬浮按钮配置卡片，用于维护首页左右侧悬浮按钮组。", [
            "悬浮框：按左侧或右侧配置悬浮按钮组，每组包含位置、排序、按钮数和启用状态。",
            "按钮项：每个按钮项可配置图标、名称、跳转目标和开关状态。",
            "操作方式：点击新增创建悬浮框，点击编辑维护按钮项；保存前检查左右位置是否冲突、移动端是否遮挡主要内容。",
        ]),
    ]
    for title, image, caption, bullets in details:
        add_heading(doc, title, 4)
        add_image(doc, images, image, caption)
        for item in bullets:
            add_field_bullet(doc, item)


def add_site_config_module(doc: Document, images: dict[str, Path], pages: list[str]):
    add_heading(doc, "四、站点配置", 1)
    add_module_intro(doc)
    add_template_management(doc, images)
    add_site_settings(doc, images)
    for page in pages[2:]:
        add_page_brief(doc, page)


def add_generic_menu_module(doc: Document, chapter: str, pages: list[str]):
    add_heading(doc, chapter, 1)
    add_module_intro(doc)
    for page in pages:
        add_page_brief(doc, page)


def add_hr_module(doc: Document, chapter: str, pages: list[str]):
    add_heading(doc, chapter, 1)
    add_module_intro(doc)
    add_page_brief(doc, "账号管理")
    add_heading(doc, "日志管理", 3)
    add_field_bullet(doc, "页面说明：日志管理用于查看后台账号登录、操作和系统异常相关记录。")
    for page in ["管理员登录日志", "管理员操作日志", "系统异常日志"]:
        add_page_brief(doc, page)


def add_business_flows(doc: Document):
    add_heading(doc, "十六、典型业务流程", 1)
    flows = [
        ("新站点上线流程", ["进入站点配置，完成模板管理、站点配置、品牌、域名、客服、分享、登录和支付相关配置。", "检查首页、登录、充值、客服、活动入口和分享入口。", "上线后通过首页看板、数据报表和异常日志观察运行状态。"]),
        ("发布活动流程", ["进入优惠活动，创建或编辑活动规则。", "检查活动展示入口、票券或任务关联、预算、领取限制和目标人群。", "上线后通过活动报表、优惠数据报表和用户分析跟踪效果。"]),
        ("处理提现流程", ["进入财务中心的提现订单列表，筛选待审核或处理中订单。", "核对会员、金额、账户、打码、风控提示和历史资金记录。", "无异常时通过；存在风险时转入提现风控审核或驳回，并填写处理原因。"]),
        ("排查用户资金问题", ["进入用户管理定位会员，或进入财务中心账变记录筛选会员账号。", "结合充值订单、提现订单、人工加扣款和账变记录核对完整资金链路。", "需要追责时进入人事中心的管理员操作日志查询相关操作。"]),
        ("查看活动效果", ["固定活动名称、统计时间和渠道范围。", "查看优惠活动、数据报表中的活动相关报表。", "结合用户分析、充值分布、留存和 LTV 判断活动质量。"]),
    ]
    for title, steps in flows:
        add_heading(doc, title, 2)
        for step in steps:
            add_number(doc, step)


def add_global_faq(doc: Document):
    add_heading(doc, "十七、常见问题与排查", 1)
    for question, answer in [
        ("登录后看不到菜单", "确认账号状态、所选角色、站点范围和该角色是否包含对应菜单。"),
        ("列表查不到数据", "检查时间范围、状态、账号、渠道等筛选条件，必要时重置后重新查询。"),
        ("配置保存后未生效", "确认配置状态、生效时间、适用范围、缓存和当前操作站点。"),
        ("订单状态异常", "查看订单详情，结合支付通道数据、账变记录、操作日志和异常日志核对。"),
        ("报表数据不一致", "确认统计时间、时区、口径、筛选条件和数据延迟。"),
        ("需要追踪误操作", "进入人事中心的管理员操作日志，按账号、时间、模块、对象和操作类型筛选。"),
    ]:
        add_heading(doc, question, 2)
        add_para(doc, answer)


def main():
    backup = backup_docx()
    images = extract_existing_images()
    doc = Document()
    set_doc_defaults(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    add_title(doc, "商户后台操作手册")
    add_manual_overview(doc)
    add_login_navigation(doc, images)
    add_home_page(doc, images)
    for chapter, pages in MENU_MODULES[1:]:
        if chapter == "四、站点配置":
            add_site_config_module(doc, images, pages)
        elif chapter == "十五、人事中心":
            add_hr_module(doc, chapter, pages)
        else:
            add_generic_menu_module(doc, chapter, pages)
    add_business_flows(doc)
    add_global_faq(doc)
    try:
        doc.save(str(DOCX_PATH))
    except PermissionError:
        print("blocked=word_file_open")
        print("请先关闭 B端后台操作手册.docx，然后回复我继续。")
        return
    print(f"backup={backup}")
    print(f"images={len(images)}")
    print(f"output={DOCX_PATH}")


if __name__ == "__main__":
    main()
