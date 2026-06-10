from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"
ASSET_DIR = ROOT / "custom" / "assets" / "manual-customer-center"
USER_SCREENSHOT = ASSET_DIR / "customer-center-screenshot.png"

FONT_NAME = "微软雅黑"
BLUE = RGBColor(47, 85, 151)
CAPTION_GRAY = RGBColor(112, 126, 146)


def set_run(run, *, size: float = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def insert_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_heading_after(anchor, text: str, level: int):
    paragraph = insert_after(anchor, style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_run(run, size=11 if level >= 4 else 12.5, bold=True, color=BLUE)
    paragraph.paragraph_format.space_after = Pt(5)
    return paragraph


def add_body_after(anchor, text: str):
    paragraph = insert_after(anchor, style="Normal")
    run = paragraph.add_run(text)
    set_run(run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullet_after(anchor, label: str, text: str):
    paragraph = insert_after(anchor, style="List Bullet")
    label_run = paragraph.add_run(label)
    set_run(label_run, size=10, bold=True, color=BLUE)
    body_run = paragraph.add_run("：" + text)
    set_run(body_run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_picture_after(anchor, image_path: Path, width: float = 6.45):
    paragraph = insert_after(anchor, style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def add_caption_after(anchor, text: str):
    paragraph = insert_after(anchor, style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run(run, size=9, color=CAPTION_GRAY)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def reset_section(document: Document):
    start = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == "客服中心":
            start = index
            break
    if start is None:
        raise RuntimeError("未找到“客服中心”章节。")

    end = len(document.paragraphs)
    for index in range(start + 1, len(document.paragraphs)):
        style_name = document.paragraphs[index].style.name
        if style_name in {"Heading 1", "Heading 2", "Heading 3"}:
            end = index
            break

    for paragraph in list(document.paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)

    anchor = document.paragraphs[start]
    for run in anchor.runs:
        set_run(run, size=12.5, bold=True, color=BLUE)
    return anchor


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{DOCX.stem}.customer-center-{stamp}.bak.docx"
    shutil.copy2(DOCX, backup)

    document = Document(DOCX)
    anchor = reset_section(document)
    last = add_body_after(
        anchor,
        "客服中心用于接待会员咨询、处理在线会话和跟进待回复问题，是客服日常处理用户沟通的主要页面。",
    )

    if USER_SCREENSHOT.exists():
        last = add_picture_after(last, USER_SCREENSHOT, 6.45)
        last = add_caption_after(last, "图：客服中心页面，用于查看会话队列、接入会员咨询并发送客服回复。")
    for label, text in [
        ("会话列表", "左侧展示当前会话、排队数量、搜索框和自动接待开关，可先从队列或搜索结果中定位目标会员。"),
        ("搜索", "支持按会员昵称、ID 或标签快速筛选会话，适合同时处理多位会员时快速切换。"),
        ("自动接待", "用于控制是否自动接入新会话；开启前先确认当前客服处理量，避免接入后回复不及时。"),
        ("聊天工作区", "右侧显示当前会话的会员信息、聊天记录、标签按钮、语言切换和消息输入区，用于集中完成沟通处理。"),
        ("发送回复", "确认当前选中的会话后，在底部输入内容并发送；涉及订单、活动、充值提现等问题时，先到对应页面核对再回复结果。"),
        ("工具入口", "可结合附件、表情或快捷用语提高回复效率，但发送前仍需按当前问题补充实际处理进度。"),
        ("问题跟进", "问题处理完成后可继续在当前会话反馈结果；如涉及风控、财务或人工审核，需说明已转交并保留会话记录便于跟进。"),
    ]:
        last = add_bullet_after(last, label, text)

    try:
        document.save(DOCX)
    except PermissionError:
        print(f"PERMISSION_ERROR backup={backup}")
        raise

    image_note = str(USER_SCREENSHOT) if USER_SCREENSHOT.exists() else "no screenshot inserted"
    print(f"updated customer center; backup={backup}; image={image_note}")


if __name__ == "__main__":
    main()
