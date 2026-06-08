from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"
ASSET_DIR = ROOT / "custom" / "assets" / "manual" / "platform-notice"

TITLE = "平台公告"


NOTICES = [
    ("平台公告", "平台结算规则升级公告", "2026-05-22 10:30", "未读"),
    ("通知", "钱包服务例行维护通知", "2026-05-21 22:00", "未读"),
    ("风控预警", "商户余额不足预警", "2026-05-20 16:18", "已读"),
    ("站内信", "运营顾问站内信提醒", "2026-05-18 09:45", "未读"),
    ("平台公告", "后台权限菜单调整公告", "2026-05-15 17:51", "已读"),
    ("平台公告", "商户后台版本发布公告", "2026-05-15 17:51", "已读"),
    ("短信", "短信签名审核结果通知", "2026-05-14 11:20", "未读"),
]


SECTIONS = [
    (
        "页面说明",
        {
            "path": ASSET_DIR / "platform-notice-list.png",
            "caption": "图：平台公告列表，用于查看平台、商户相关通知和站内消息。",
            "width": 6.5,
        },
        [
            "平台公告用于查看平台、系统和商户相关消息。商户可从顶部导航栏右侧的消息入口进入该页面，进入后按类型、标题和阅读状态筛选消息，并通过“查看”打开消息详情。",
            "页面消息类型包含平台公告、通知、风控预警、站内信和短信。列表同时展示消息标题、发送时间、阅读状态和操作入口，便于商户快速处理重要公告、维护通知、资金预警和运营提醒。",
        ],
    ),
    (
        "筛选条件",
        None,
        [
            "类型：用于按消息来源或业务类型筛选。可选择全部类型、平台公告、通知、短信、风控预警或站内信。排查某类消息时先选择类型，再结合状态缩小范围。",
            "标题：支持输入标题关键字查询，例如结算规则、钱包服务、余额不足、短信签名等。输入后点击“搜索”，也可按回车触发查询。",
            "状态：用于按未读或已读筛选。处理待办消息时建议选择“未读”，核对历史公告时可切换为“已读”或全部状态。",
            "搜索：按当前筛选条件刷新列表。筛选无结果时，先检查标题关键字是否过窄，再重置条件重新查询。",
            "重置：清空类型、标题和状态筛选，恢复展示全部消息列表。",
            "全部已读：将当前消息标记为已读。点击前应确认未读公告已经查看，避免遗漏维护通知、资金预警或审核结果等重要内容。",
        ],
    ),
    (
        "列表字段",
        None,
        [
            "类型：显示消息分类，例如平台公告、通知、风控预警、站内信、短信。类型可帮助商户判断消息优先级和处理责任。",
            "标题：显示消息主题。未读消息在列表中以更醒目的文字展示，建议优先处理与结算、钱包维护、风控预警和审核结果相关的标题。",
            "时间：显示消息发送时间，用于判断公告生效日期、维护窗口和消息处理时效。",
            "状态：显示未读或已读。查看消息详情后，未读消息会转为已读；也可通过“全部已读”批量处理。",
            "操作：点击“查看”打开消息详情，查看标题、类型和完整内容。查看后返回列表时，应根据消息内容继续到对应业务页面处理，例如结算中心、充值提现、权限管理或消息配置等。",
        ],
    ),
    (
        "查看详情",
        None,
        [
            "查看单条消息：在目标消息行点击“查看”，系统打开公告详情弹窗，弹窗中展示消息标题、类型和正文内容。",
            "阅读状态变化：未读消息打开详情后会自动标记为已读。若列表当前筛选为未读，查看后该消息可能从当前筛选结果中消失，这是正常的状态刷新结果。",
            "关闭详情：阅读完成后点击弹窗右上角关闭按钮或遮罩区域返回列表。若消息涉及维护时间、资金预警、短信签名审核或后台版本发布，应在关闭前确认后续处理动作。",
        ],
    ),
    (
        "注意事项",
        None,
        [
            "平台公告页是商户接收平台和系统消息的重要入口，尤其是钱包维护、结算规则调整、余额不足预警、权限菜单调整和短信审核结果，建议运营或管理员每日查看未读消息。",
            "“全部已读”只改变阅读状态，不代表业务已处理。批量标记前应先筛选重要类型并逐条查看，避免把需要执行的任务误认为已完成。",
            "若顶部导航栏右侧有未读消息提示，但进入页面后看不到记录，优先点击“重置”，再检查类型、标题和状态筛选是否限制了列表。",
            "如果公告内容要求进入其他模块处理，应按公告中的业务方向跳转到对应页面操作，并在处理完成后回到平台公告页确认相关消息状态。",
        ],
    ),
    (
        "常见问题",
        None,
        [
            "为什么点击“查看”后未读数量减少：打开详情会将该条消息标记为已读，未读数量随之减少；如果当前筛选条件为未读，列表也会同步移除该条消息。",
            "为什么搜索不到某条公告：先点击“重置”恢复全部列表，再只输入标题中的短关键字查询；如果仍无结果，确认消息类型和状态是否选错。",
            "收到风控预警后怎么处理：先查看详情确认预警原因，例如商户余额不足，再进入相关资金、提现或风控页面处理，处理完成后持续关注后续公告或站内信。",
        ],
    ),
]


def font(size, bold=False):
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def draw_text_center(draw, xy, text, fill, text_font):
    x1, y1, x2, y2 = xy
    bbox = draw.textbbox((0, 0), text, font=text_font)
    x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2
    y = y1 + (y2 - y1 - (bbox[3] - bbox[1])) / 2 - 1
    draw.text((x, y), text, fill=fill, font=text_font)


def draw_button(draw, xy, text, fill, outline, text_fill, text_font):
    draw.rounded_rectangle(xy, radius=3, fill=fill, outline=outline, width=1)
    draw_text_center(draw, xy, text, text_fill, text_font)


def create_notice_image():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / "platform-notice-list.png"

    image = Image.new("RGB", (1680, 630), "#ffffff")
    draw = ImageDraw.Draw(image)
    f12 = font(22)
    f13 = font(24)
    f14 = font(25)
    f14b = font(25, bold=True)
    f15b = font(27, bold=True)

    draw.rectangle((0, 0, 1680, 56), fill="#ffffff")
    y = 16
    draw.text((8, y + 6), "类型:", fill="#1f2d3d", font=f13)
    draw.rounded_rectangle((74, y, 164, y + 32), radius=3, fill="#ffffff", outline="#cfd8e6")
    draw.text((92, y + 7), "全部类型", fill="#6b7a90", font=f12)
    draw.text((146, y + 6), "⌄", fill="#9aa8ba", font=f13)

    draw.text((176, y + 6), "标题:", fill="#1f2d3d", font=f13)
    draw.rounded_rectangle((214, y, 374, y + 32), radius=3, fill="#ffffff", outline="#cfd8e6")
    draw.text((224, y + 7), "请输入标题", fill="#a2adbd", font=f12)

    draw.text((388, y + 6), "状态:", fill="#1f2d3d", font=f13)
    draw.rounded_rectangle((424, y, 544, y + 32), radius=3, fill="#ffffff", outline="#cfd8e6")
    draw.text((448, y + 7), "全部状态", fill="#6b7a90", font=f12)
    draw.text((524, y + 6), "⌄", fill="#9aa8ba", font=f13)

    draw_button(draw, (558, y, 614, y + 32), "搜索", "#2f8df6", "#2f8df6", "#ffffff", f12)
    draw_button(draw, (626, y, 684, y + 32), "重置", "#ffffff", "#cfd8e6", "#344256", f12)
    draw_button(draw, (698, y, 778, y + 32), "全部已读", "#59bf32", "#59bf32", "#ffffff", f12)

    table_x, table_y = 8, 56
    table_w = 1660
    row_h = 43
    col_widths = [98, 1255, 144, 80, 83]
    headers = ["类型", "标题", "时间", "状态", "操作"]

    draw.rectangle((table_x, table_y, table_x + table_w, table_y + row_h * 8), outline="#dfe6ef", width=1)
    draw.rectangle((table_x, table_y, table_x + table_w, table_y + row_h), fill="#f7f8fa", outline="#dfe6ef")
    x = table_x
    for idx, width in enumerate(col_widths):
        draw.line((x, table_y, x, table_y + row_h * 8), fill="#dfe6ef", width=1)
        draw_text_center(draw, (x, table_y, x + width, table_y + row_h), headers[idx], "#0f2036", f14b)
        x += width
    draw.line((x, table_y, x, table_y + row_h * 8), fill="#dfe6ef", width=1)

    for row_index, row in enumerate(NOTICES, start=1):
        y1 = table_y + row_h * row_index
        y2 = y1 + row_h
        draw.rectangle((table_x, y1, table_x + table_w, y2), fill="#ffffff", outline="#dfe6ef")
        x = table_x
        unread = row[3] == "未读"
        colors = ["#0f2036", "#0f2036", "#0f2036", "#ff6b2b", "#006dff"]
        weights = [f14b if unread else f14, f14b if unread else f14, f14b if unread else f14, f14, f14b]
        if row[3] == "已读":
            colors[3] = "#344256"
        for idx, width in enumerate(col_widths):
            draw.line((x, y1, x, y2), fill="#dfe6ef", width=1)
            draw_text_center(draw, (x, y1, x + width, y2), row[idx] if idx < 4 else "查看", colors[idx], weights[idx])
            x += width
        draw.line((x, y1, x, y2), fill="#dfe6ef", width=1)

    pager_y = 583
    pager_x = 1232
    for label in ["|‹", "‹‹", "‹", "1", "›", "››", "›|"]:
        w = 28 if label != "1" else 38
        draw.rounded_rectangle((pager_x, pager_y, pager_x + w, pager_y + 26), radius=3, fill="#ffffff", outline="#d6dde8")
        draw_text_center(draw, (pager_x, pager_y, pager_x + w, pager_y + 26), label, "#496176", f12)
        pager_x += w + 8
        if label == "1":
            draw.text((pager_x, pager_y + 4), "/ 1", fill="#344256", font=f12)
            pager_x += 28
    draw.rounded_rectangle((pager_x, pager_y, pager_x + 96, pager_y + 26), radius=3, fill="#ffffff", outline="#d6dde8")
    draw.text((pager_x + 8, pager_y + 5), "20条/页", fill="#496176", font=f12)
    draw.text((pager_x + 78, pager_y + 4), "⌄", fill="#8b97a8", font=f12)
    draw.text((pager_x + 104, pager_y + 5), "共 7 条记录", fill="#496176", font=f12)

    image.save(path)
    return path


def set_field_run_style(run):
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(31, 78, 121)


def set_body_run_style(run):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10)


def set_heading_style(paragraph):
    for run in paragraph.runs:
        run.bold = True
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor(31, 78, 121)
        run.font.size = Pt(10.5)


def set_caption_style(paragraph):
    paragraph.alignment = 1
    for run in paragraph.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(127, 127, 127)


def add_label_text(paragraph, text):
    label, sep, body = text.partition("：")
    if sep:
        label_run = paragraph.add_run(label + sep)
        set_field_run_style(label_run)
        body_run = paragraph.add_run(body)
        set_body_run_style(body_run)
    else:
        run = paragraph.add_run(text)
        set_body_run_style(run)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_target_range(document):
    start = None
    for idx, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name.startswith("Heading") and paragraph.text.strip() == TITLE:
            start = idx
            break
    if start is None:
        raise RuntimeError("未找到平台公告章节")

    start_style = document.paragraphs[start].style.name
    target_levels = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3, "Heading 4": 4}
    start_level = target_levels.get(start_style, 3)
    end = len(document.paragraphs)
    for idx in range(start + 1, len(document.paragraphs)):
        paragraph = document.paragraphs[idx]
        level = target_levels.get(paragraph.style.name)
        if level is not None and level <= start_level and paragraph.text.strip():
            end = idx
            break
    return start, end


def update_platform_notice_section():
    image_path = create_notice_image()
    BACKUP_DIR.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = BACKUP_DIR / f"B端后台操作手册.platform-notice-{timestamp}.bak.docx"
    shutil.copy2(DOCX_PATH, backup_path)

    document = Document(DOCX_PATH)
    start, end = find_target_range(document)

    anchor = document.paragraphs[start]
    for paragraph in list(document.paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)

    current = anchor
    for heading, image, bullets in SECTIONS:
        current = insert_paragraph_after(current, style="Heading 4")
        current.add_run(heading)
        set_heading_style(current)

        if image:
            if not image["path"].exists():
                raise FileNotFoundError(image["path"])
            current = insert_paragraph_after(current, style="Normal")
            current.alignment = 1
            current.add_run().add_picture(str(image["path"]), width=Inches(image["width"]))
            current = insert_paragraph_after(current, image["caption"], style="Caption")
            set_caption_style(current)

        for bullet in bullets:
            current = insert_paragraph_after(current, style="List Bullet")
            add_label_text(current, bullet)

    try:
        document.save(DOCX_PATH)
    except PermissionError:
        print(f"PERMISSION_ERROR backup={backup_path}")
        raise

    return backup_path, image_path


if __name__ == "__main__":
    backup, image = update_platform_notice_section()
    print(f"updated platform notice; backup={backup}; image={image}")
