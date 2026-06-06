from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"
ASSET_DIR = ROOT / "custom" / "assets" / "manual-site-config" / "site-domain"

PAGE_TITLE = "站点域名"
FONT_NAME = "微软雅黑"
BLUE = RGBColor(47, 85, 151)
CAPTION_GRAY = RGBColor(112, 126, 146)


def image_font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def text(draw: ImageDraw.ImageDraw, xy, value, font, fill="#172033", anchor=None):
    draw.text(xy, value, font=font, fill=fill, anchor=anchor)


def rect(draw: ImageDraw.ImageDraw, box, fill="#ffffff", outline="#dfe5ef", width=1):
    draw.rectangle(box, fill=fill, outline=outline, width=width)


def rounded(draw: ImageDraw.ImageDraw, box, fill="#ffffff", outline="#d6dce7", radius=4, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def button(draw: ImageDraw.ImageDraw, box, value, primary=False):
    fill = "#409eff" if primary else "#ffffff"
    outline = "#409eff" if primary else "#d6dce7"
    color = "#ffffff" if primary else "#2f3b52"
    rounded(draw, box, fill=fill, outline=outline, radius=4)
    x1, y1, x2, y2 = box
    text(draw, ((x1 + x2) / 2, (y1 + y2) / 2), value, image_font(14), color, "mm")


def input_box(draw: ImageDraw.ImageDraw, box, value, arrow=False):
    rounded(draw, box, fill="#ffffff", outline="#cfd8e6", radius=3)
    x1, y1, x2, y2 = box
    text(draw, (x1 + 13, (y1 + y2) / 2), value, image_font(13), "#647083", "lm")
    if arrow:
        draw.polygon([(x2 - 20, y1 + 14), (x2 - 10, y1 + 14), (x2 - 15, y1 + 21)], fill="#303744")


def copy_icon(draw: ImageDraw.ImageDraw, x, y):
    rounded(draw, (x, y, x + 13, y + 13), fill="#ffffff", outline="#5a9cff", radius=2)


def draw_tabs(draw: ImageDraw.ImageDraw, active: str):
    tabs = [("frontend", "前台域名管理", 10), ("backend", "后台域名管理", 112), ("promotion", "推广域名", 232)]
    for key, label, x in tabs:
        color = "#1677ff" if key == active else "#172033"
        text(draw, (x, 18), label, image_font(14), color, "lm")
        if key == active:
            draw.line((x, 33, x + 86, 33), fill="#2f8dff", width=2)
    draw.line((10, 33, 1640, 33), fill="#d9dee8", width=1)


def draw_toolbar(draw: ImageDraw.ImageDraw, y=54):
    text(draw, (10, y + 17), "节点类型:", image_font(14), "#303744", "lm")
    input_box(draw, (78, y, 257, y + 34), "请选择节点类型", arrow=True)
    text(draw, (276, y + 17), "主域名:", image_font(14), "#303744", "lm")
    input_box(draw, (327, y, 507, y + 34), "请输入主域名")
    button(draw, (523, y, 585, y + 34), "搜索", True)
    button(draw, (601, y, 663, y + 34), "重置")
    button(draw, (1576, y, 1638, y + 34), "新增", True)


def draw_status_cell(draw: ImageDraw.ImageDraw, x, y, width):
    f = image_font(13)
    text(draw, (x + 10, y + 15), "jocelyn.ns.cloudflare.com", f, "#172033", "lm")
    copy_icon(draw, x + 180, y + 9)
    text(draw, (x + 10, y + 35), "zeus.ns.cloudflare.com", f, "#172033", "lm")
    copy_icon(draw, x + 170, y + 29)
    text(draw, (x + width / 2, y + 55), "验证通过", image_font(12), "#23a812", "mm")


def draw_domain_cell(draw: ImageDraw.ImageDraw, x, y, domain):
    text(draw, (x + 10, y + 34), domain, image_font(13), "#172033", "lm")
    copy_icon(draw, x + 10 + min(152, int(len(domain) * 6.4)), y + 27)


def draw_footer(draw: ImageDraw.ImageDraw, count):
    y = 648
    text(draw, (1124, y + 17), "|‹", image_font(18), "#b9c0cc", "lm")
    text(draw, (1158, y + 17), "‹‹", image_font(18), "#b9c0cc", "lm")
    text(draw, (1197, y + 17), "‹", image_font(18), "#b9c0cc", "lm")
    input_box(draw, (1222, y, 1278, y + 34), "1")
    text(draw, (1288, y + 17), "/ 1", image_font(16), "#172033", "lm")
    text(draw, (1323, y + 17), "›", image_font(18), "#b9c0cc", "lm")
    text(draw, (1362, y + 17), "››", image_font(18), "#b9c0cc", "lm")
    text(draw, (1397, y + 17), "›|", image_font(18), "#b9c0cc", "lm")
    input_box(draw, (1425, y, 1551, y + 34), "10条/页", arrow=True)
    text(draw, (1560, y + 17), f"共 {count} 条记录", image_font(14), "#172033", "lm")


def draw_table(draw: ImageDraw.ImageDraw, active: str, headers, widths, rows):
    x0, y0 = 10, 104
    header_h = 48
    row_h = 68 if active == "frontend" else 70
    total_w = sum(widths)
    x = x0
    for header, width in zip(headers, widths):
        rect(draw, (x, y0, x + width, y0 + header_h), "#fafbfc", "#e3e8f0")
        text(draw, (x + width / 2, y0 + header_h / 2), header, image_font(13, True), "#303744", "mm")
        x += width
    for r_index, row in enumerate(rows):
        y = y0 + header_h + r_index * row_h
        x = x0
        for c_index, width in enumerate(widths):
            rect(draw, (x, y, x + width, y + row_h), "#ffffff", "#e3e8f0")
            header = headers[c_index]
            if header == "CDN节点名称":
                text(draw, (x + width / 2, y + row_h / 2), "CF", image_font(13), "#172033", "mm")
            elif header == "域名":
                draw_domain_cell(draw, x, y, row["domain"])
            elif header == "主域名是否已激活":
                draw_status_cell(draw, x, y, width)
            elif header == "站点":
                if row.get("site"):
                    text(draw, (x + width / 2 - 8, y + row_h / 2), row["site"], image_font(13), "#172033", "mm")
                    text(draw, (x + width / 2 + 58, y + row_h / 2), "修改", image_font(12), "#1677ff", "lm")
                else:
                    text(draw, (x + width / 2, y + row_h / 2), "绑定站点", image_font(12), "#1677ff", "mm")
            elif header == "落地页模板":
                text(draw, (x + width / 2 - 18, y + row_h / 2), row.get("template", ""), image_font(13), "#172033", "mm")
                text(draw, (x + width / 2 + 30, y + row_h / 2), "修改", image_font(12), "#1677ff", "lm")
            elif header == "域名状态":
                color = "#172033"
                text(draw, (x + width / 2, y + row_h / 2), row["status"], image_font(13), color, "mm")
            elif header == "操作人":
                text(draw, (x + width / 2, y + row_h / 2), "white", image_font(13), "#172033", "mm")
            elif header == "操作时间":
                text(draw, (x + width / 2, y + row_h / 2), row["time"], image_font(12), "#172033", "mm")
            elif header == "操作":
                action = "停用" if row["enabled"] else "启用"
                text(draw, (x + width / 2, y + row_h / 2 - 10), action, image_font(12), "#1677ff", "mm")
                text(draw, (x + width / 2, y + row_h / 2 + 12), "删除", image_font(12), "#ff4d4f", "mm")
            x += width
    rect(draw, (x0, y0, x0 + total_w, 625), fill=None, outline="#e3e8f0")
    draw_footer(draw, len(rows))


def save_tab_image(active: str, filename: str, rows):
    img = Image.new("RGB", (1645, 700), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw_tabs(draw, active)
    draw_toolbar(draw)
    if active == "promotion":
        headers = ["CDN节点名称", "域名", "主域名是否已激活", "站点", "落地页模板", "域名状态", "过期时间", "操作人", "操作时间", "操作"]
        widths = [104, 194, 290, 193, 160, 140, 120, 150, 150, 130]
    else:
        headers = ["CDN节点名称", "域名", "主域名是否已激活", "站点", "域名状态", "过期时间", "操作人", "操作时间", "操作"]
        widths = [112, 208, 348, 210, 172, 146, 130, 162, 138]
    draw_table(draw, active, headers, widths, rows)
    img.save(ASSET_DIR / filename)


def save_modal_image():
    img = Image.new("RGB", (629, 392), "#ffffff")
    draw = ImageDraw.Draw(img)
    rect(draw, (0, 0, 628, 391), "#ffffff", "#cfd4dc")
    rect(draw, (0, 0, 628, 40), "#ffffff", "#dcdfe6")
    text(draw, (10, 20), "新增前台域名", image_font(14, True), "#172033", "lm")
    draw.line((599, 15, 611, 27), fill="#606266", width=2)
    draw.line((611, 15, 599, 27), fill="#606266", width=2)
    text(draw, (31, 84), "*", image_font(14), "#f56c6c", "rm")
    text(draw, (37, 84), "CDN节\n点:", image_font(14), "#303744", "lm")
    draw.ellipse((100, 78, 114, 92), outline="#1677ff", width=2)
    draw.ellipse((104, 82, 110, 88), fill="#1677ff")
    text(draw, (123, 85), "CF", image_font(14), "#172033", "lm")
    text(draw, (31, 151), "*", image_font(14), "#f56c6c", "rm")
    text(draw, (37, 151), "主域名:", image_font(14), "#303744", "lm")
    rounded(draw, (95, 136, 611, 256), fill="#ffffff", outline="#cfd8e6", radius=3)
    text(draw, (110, 157), "支持批量添加，最多20个，多个域名请换行", image_font(14), "#647083", "lm")
    text(draw, (110, 276), "温馨提示：支持顶级域名及子域名。", image_font(12), "#ff7a00", "lm")
    button(draw, (483, 345, 545, 379), "取消")
    button(draw, (553, 345, 615, 379), "提交", True)
    img.save(ASSET_DIR / "add-domain-modal.png")


def generate_images():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    frontend_rows = [
        {"domain": "test3.gbsoft8686.com", "site": "", "status": "已停用", "enabled": False, "time": "2026-05-13T13:39:48..."},
        {"domain": "test2.gbsoft8686.com", "site": "", "status": "已停用", "enabled": False, "time": "2026-05-13T13:39:48..."},
        {"domain": "test1.gbsoft8686.com", "site": "(102)测试站点2", "status": "正常使用", "enabled": True, "time": "2026-05-13T18:36:32..."},
        {"domain": "cloud222w.jyowhite.cc", "site": "", "status": "已停用", "enabled": False, "time": "2026-05-13T13:32:50..."},
        {"domain": "cloud22w.jyowhite.cc", "site": "", "status": "已停用", "enabled": False, "time": "2026-05-13T13:32:29..."},
        {"domain": "www.jyowhite.cc", "site": "(101)测试站点", "status": "正常使用", "enabled": True, "time": "2026-05-13T13:56:59..."},
        {"domain": "jyowhite.cc", "site": "(101)测试站点", "status": "正常使用", "enabled": True, "time": "2026-02-24T16:00:24..."},
    ]
    backend_rows = [
        {"domain": "admin.jyowhite.cc", "site": "(101)测试站点", "status": "正常使用", "enabled": True, "time": "2026-05-13T14:08:21..."},
        {"domain": "bo-test.gbsoft8686.com", "site": "", "status": "已停用", "enabled": False, "time": "2026-05-13T13:47:06..."},
        {"domain": "manage.jyowhite.cc", "site": "(102)测试站点2", "status": "正常使用", "enabled": True, "time": "2026-04-28T19:22:10..."},
    ]
    promotion_rows = [
        {"domain": "uiguaranwda.jyowhite.cc", "site": "(101)测试站点", "template": "1123", "status": "正常使用", "enabled": True, "time": "2026-05-13T15:11:1..."},
    ]
    save_tab_image("frontend", "frontend-domain-list.png", frontend_rows)
    save_tab_image("backend", "backend-domain-list.png", backend_rows)
    save_tab_image("promotion", "promotion-domain-list.png", promotion_rows)
    save_modal_image()


def set_run(run, *, size=10, bold=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)


def insert_after(paragraph, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = paragraph._parent.add_paragraph()
    inserted._p = new_p
    if style:
        inserted.style = style
    return inserted


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def reset_section(document: Document):
    start = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == PAGE_TITLE:
            start = index
            break
    if start is None:
        raise RuntimeError("未找到“站点域名”章节。")
    end = len(document.paragraphs)
    for index in range(start + 1, len(document.paragraphs)):
        if document.paragraphs[index].style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            end = index
            break
    for paragraph in list(document.paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)
    heading = document.paragraphs[start]
    for run in heading.runs:
        set_run(run, size=12.5, bold=True, color=BLUE)
    return heading


def add_body(anchor, value):
    paragraph = insert_after(anchor, "Normal")
    run = paragraph.add_run(value)
    set_run(run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_heading(anchor, value):
    paragraph = insert_after(anchor, "Heading 4")
    run = paragraph.add_run(value)
    set_run(run, size=11, bold=True, color=BLUE)
    paragraph.paragraph_format.space_after = Pt(5)
    return paragraph


def add_bullet(anchor, label, value):
    paragraph = insert_after(anchor, "List Bullet")
    label_run = paragraph.add_run(label)
    set_run(label_run, size=10, bold=True, color=BLUE)
    body_run = paragraph.add_run("：" + value)
    set_run(body_run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.12
    return paragraph


def add_picture(anchor, filename, caption, width):
    paragraph = insert_after(anchor, "Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(ASSET_DIR / filename), width=Inches(width))
    paragraph.paragraph_format.space_after = Pt(2)
    cap = insert_after(paragraph, "Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run(run, size=9, color=CAPTION_GRAY)
    cap.paragraph_format.space_after = Pt(6)
    return cap


def update_document():
    document = Document(DOCX)
    last = reset_section(document)
    last = add_body(
        last,
        "站点域名用于维护不同业务入口的域名配置，页面通过顶部页签区分前台域名管理、后台域名管理和推广域名。操作时先根据蓝色选中态识别当前页面，再执行查询、新增、绑定站点、修改落地页模板、启用、停用或删除等操作。",
    )

    last = add_heading(last, "前台域名管理")
    last = add_picture(last, "frontend-domain-list.png", "图18：站点域名-前台域名管理，用于维护用户访问前台站点使用的域名。", 6.5)
    for label, value in [
        ("选中状态", "顶部“前台域名管理”为蓝色文字并显示下划线时，代表当前列表展示前台域名数据。"),
        ("筛选条件", "支持按节点类型和主域名查询；节点类型示例为 CF，主域名可输入完整或部分域名后点击搜索，点击重置清空条件。"),
        ("列表字段", "表格展示 CDN 节点名称、域名、主域名是否已激活、站点、域名状态、过期时间、操作人、操作时间和操作列；主域名激活区域会展示 NS 地址和验证结果。"),
        ("站点绑定", "未绑定站点时显示“绑定站点”，已绑定时展示站点编号和名称，并提供“修改”入口；域名启用中时应先停用再修改绑定关系。"),
        ("状态操作", "操作列支持启用、停用和删除。启用前需确认域名解析、主域名验证和站点绑定均正确，删除前应确认该域名不再承担线上访问入口。"),
    ]:
        last = add_bullet(last, label, value)

    last = add_heading(last, "后台域名管理")
    last = add_picture(last, "backend-domain-list.png", "图19：站点域名-后台域名管理，用于维护后台管理端访问域名。", 6.5)
    for label, value in [
        ("选中状态", "顶部“后台域名管理”为蓝色文字并显示下划线时，代表当前列表展示后台域名数据。"),
        ("适用范围", "后台域名用于运营、客服、财务、风控和管理人员访问后台系统，配置前应确认该域名只面向授权后台入口使用。"),
        ("操作要点", "新增或启用后台域名前，应检查 CDN 节点、NS 验证、站点绑定和访问权限策略，避免后台入口误暴露或与前台域名混用。"),
        ("停用删除", "停用会影响后台访问入口，应先确认替代后台域名可用；删除前需确认无账号、公告、白名单或运维文档仍指向该域名。"),
    ]:
        last = add_bullet(last, label, value)

    last = add_heading(last, "推广域名")
    last = add_picture(last, "promotion-domain-list.png", "图20：站点域名-推广域名，用于维护推广入口域名及落地页模板。", 6.5)
    for label, value in [
        ("选中状态", "顶部“推广域名”为蓝色文字并显示下划线时，代表当前列表展示推广域名数据。"),
        ("差异字段", "推广域名列表比前台和后台域名多出“落地页模板”列，可查看当前模板编号，并通过“修改”调整绑定模板。"),
        ("使用场景", "推广域名通常用于投放渠道、活动入口或拉新页面，配置后应同步检查落地页模板、活动内容、渠道参数和前台跳转链路。"),
        ("上线复核", "启用推广域名前，应确认域名解析和主域名验证通过，落地页模板已配置正确，且推广链接不会跳转到停用站点或过期活动。"),
    ]:
        last = add_bullet(last, label, value)

    last = add_heading(last, "新增域名弹窗")
    last = add_picture(last, "add-domain-modal.png", "图21：站点域名-新增域名弹窗，三类域名新增表单字段一致。", 4.9)
    for label, value in [
        ("入口识别", "点击右上角“新增”打开弹窗，弹窗标题根据当前选中的页签自动识别为“新增前台域名”“新增后台域名”或“新增推广域名”。"),
        ("CDN节点", "当前示例默认选择 CF；如后续支持更多 CDN 节点，应按实际解析服务选择，避免域名添加到错误节点。"),
        ("主域名", "支持批量添加，最多 20 个，多个域名需换行填写；支持顶级域名及子域名，提交前应检查拼写、后缀和重复项。"),
        ("提交结果", "点击提交后按当前页签类型新增到对应域名列表；取消或右上角关闭不会保存输入内容。新增后继续完成解析验证、站点绑定和启用操作。"),
    ]:
        last = add_bullet(last, label, value)

    last = add_heading(last, "常见问题")
    for label, value in [
        ("为什么新增弹窗标题不同但字段一样", "页面通过当前选中页签判断新增类型，因此只需要在对应页签下点击新增；表单字段保持一致，提交后的归属由当前页签决定。"),
        ("绑定站点无法修改怎么办", "先检查域名是否处于正常使用状态。启用中的域名通常需要先停用，再修改绑定站点，避免线上访问入口在运行中被直接切换。"),
        ("主域名验证未通过怎么办", "核对 NS 地址是否已在域名服务商处配置，等待 DNS 生效后重新检查；仍未通过时确认域名是否填错、节点是否选错或解析记录被其他服务覆盖。"),
    ]:
        last = add_bullet(last, label, value)

    document.save(DOCX)


def main():
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    generate_images()
    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{DOCX.stem}.site-domain-{stamp}.bak.docx"
    shutil.copy2(DOCX, backup)
    update_document()
    print(f"backup={backup}")
    print(f"assets={ASSET_DIR}")


if __name__ == "__main__":
    main()
