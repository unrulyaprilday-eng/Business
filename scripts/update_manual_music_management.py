from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"


TITLE = "音乐管理"


SECTIONS = [
    (
        "页面说明",
        {
            "path": ROOT / "custom" / "assets" / "manual" / "music-management" / "music-management-list-user.png",
            "caption": "图：音乐管理列表，用于维护音乐开关、曲目状态和列表操作。",
            "width": 6.5,
        },
        [
            "音乐管理用于维护客户端背景音乐或侧滑菜单音乐播放器中可播放的曲目。页面顶部提供“音乐开关”，用于统一控制客户端音乐播放功能；下方列表展示已上传的音乐文件，支持新增、修改、删除和单曲启用或停用。",
            "当音乐开关关闭时，客户端音乐播放功能关闭，侧滑菜单中的音乐播放器不再展示。关闭前应确认当前站点是否仍需要背景音乐、活动音乐或用户侧音乐入口，避免影响前台体验。",
        ],
    ),
    (
        "列表字段",
        None,
        [
            "序号：显示当前音乐在列表中的顺序编号，用于快速核对曲目数量和位置。",
            "歌曲名称：展示音乐名称，例如 BGM1、BGM2。名称建议简短清晰，便于运营人员识别用途或播放场景。",
            "大小/MB：展示源文件大小，便于判断文件是否符合上传限制以及是否可能影响前台加载速度。",
            "源文件：展示音乐文件地址。列表中较长地址会自动省略，核对文件时可结合修改入口或源文件管理记录确认完整链接。",
            "停/启用：控制单首音乐是否参与前台展示或播放。启用后该曲目可被客户端使用；停用后该曲目不应继续出现在用户侧音乐播放列表中。",
            "操作：提供“修改”和“删除”。修改用于调整歌曲名称、排序或重新上传文件；删除用于移除不再使用的曲目，删除前应确认该音乐没有被活动、页面或运营配置依赖。",
        ],
    ),
    (
        "新增音乐",
        {
            "path": ROOT / "custom" / "assets" / "manual" / "music-management" / "music-management-add-modal-user.png",
            "caption": "图：新增音乐弹窗，用于填写歌曲名称、排序并上传 MP3 文件。",
            "width": 4.7,
        },
        [
            "点击右上角“新增”打开“新增音乐”弹窗，按页面提示填写歌曲名称、排序并上传音乐文件。带星号的歌曲名称和音乐文件为必填项，提交前应确认名称不重复、排序值符合展示顺序、文件格式和大小符合限制。",
            "歌曲名称：填写后台列表和前台可能展示的曲目名称。建议使用统一命名规则，例如 BGM1、活动主题曲、首页背景音乐等。",
            "排序：填写数字排序值，用于控制音乐在列表或前台播放器中的排列位置。数字规则以页面实际排序口径为准，保存后应回到列表确认顺序是否符合预期。",
            "音乐文件：点击“上传歌曲”选择本地 MP3 文件。页面提示文件需为 MP3 格式且小于 10240KB；超过限制或格式不符时应重新压缩或转换后再上传。",
            "提交与取消：点击“提交”保存新增音乐并回到列表；点击“取消”或右上角关闭不会保存本次填写内容。提交后应检查列表中歌曲名称、大小、源文件地址和启用状态是否正常。",
        ],
    ),
    (
        "修改、停用与删除",
        None,
        [
            "修改音乐：在列表操作列点击“修改”，进入编辑弹窗后可调整歌曲名称、排序或重新上传源文件。替换文件前应确认新文件试听正常、大小符合限制，并避免误覆盖线上正在使用的曲目。",
            "启用或停用单曲：通过列表“停/启用”开关控制单首音乐状态。批量调整多首音乐前，应先确认全局音乐开关处于预期状态，否则单曲启用后用户侧也可能无法播放。",
            "删除音乐：点击“删除”移除曲目。删除属于高影响操作，执行前应确认该歌曲不再用于前台播放、活动配置或运营素材；误删后通常需要重新上传文件并重新配置排序。",
        ],
    ),
    (
        "注意事项",
        None,
        [
            "音乐文件会影响用户侧加载和播放体验，建议控制文件大小、音量和时长，避免上传过大的 MP3 或音质异常文件。",
            "新增或修改后应在前台对应入口做一次播放验证，重点检查播放器是否展示、歌曲是否能正常加载、切换曲目是否顺畅，以及关闭音乐开关后用户侧入口是否按预期隐藏。",
            "若前台没有音乐入口，优先检查全局音乐开关是否开启、单曲是否启用、源文件地址是否可访问、文件格式是否为 MP3，以及当前站点模板是否支持音乐播放器展示。",
        ],
    ),
]


def set_field_run_style(run):
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(31, 78, 121)


def set_body_run_style(run):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10)


def set_heading_style(paragraph):
    for run in paragraph.runs:
        run.bold = True
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor(31, 78, 121)
        run.font.size = Pt(10.5)


def set_caption_style(paragraph):
    paragraph.alignment = 1
    for run in paragraph.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(127, 127, 127)


def add_label_text(paragraph, text):
    label, sep, body = text.partition("：")
    if sep:
        label_run = paragraph.add_run(label + sep)
        set_field_run_style(label_run)
        body_run = paragraph.add_run(body)
        set_body_run_style(body_run)
    else:
        run = paragraph.add_run(text)
        set_body_run_style(run)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_music_range(document):
    start = None
    for idx, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name.startswith("Heading") and paragraph.text.strip() == TITLE:
            start = idx
            break
    if start is None:
        raise RuntimeError("未找到音乐管理章节")

    end = len(document.paragraphs)
    for idx in range(start + 1, len(document.paragraphs)):
        paragraph = document.paragraphs[idx]
        if paragraph.style.name in ("Heading 1", "Heading 2", "Heading 3") and paragraph.text.strip():
            end = idx
            break
    return start, end


def update_music_section():
    BACKUP_DIR.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = BACKUP_DIR / f"B端后台操作手册.music-management-{timestamp}.bak.docx"
    shutil.copy2(DOCX_PATH, backup_path)

    document = Document(DOCX_PATH)
    start, end = find_music_range(document)

    anchor = document.paragraphs[start]
    for paragraph in list(document.paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)

    current = anchor
    for heading, image, bullets in SECTIONS:
        current = insert_paragraph_after(current, style="Heading 4")
        current.add_run(heading)
        set_heading_style(current)

        if image:
            if not image["path"].exists():
                raise FileNotFoundError(image["path"])
            current = insert_paragraph_after(current, style="Normal")
            current.alignment = 1
            current.add_run().add_picture(str(image["path"]), width=Inches(image["width"]))
            current = insert_paragraph_after(current, image["caption"], style="Caption")
            set_caption_style(current)

        for bullet in bullets:
            current = insert_paragraph_after(current, style="List Bullet")
            add_label_text(current, bullet)

    try:
        document.save(DOCX_PATH)
    except PermissionError:
        print(f"PERMISSION_ERROR backup={backup_path}")
        raise

    return backup_path


if __name__ == "__main__":
    backup = update_music_section()
    print(f"updated music management; backup={backup}")
