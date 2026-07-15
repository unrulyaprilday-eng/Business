from __future__ import annotations

import argparse
from dataclasses import dataclass
from io import BytesIO
import json
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

import extract_chat_images_from_codex_log as chat_images


ROOT = Path(__file__).resolve().parents[1]
BACKUP_DIR = ROOT / "backups"

FONT_NAME = "Microsoft YaHei"
BLUE = RGBColor(47, 85, 151)
CAPTION_GRAY = RGBColor(112, 126, 146)
HEADING_SIZES = {
    "Heading 1": 14,
    "Heading 2": 12.5,
    "Heading 3": 10.5,
    "Heading 4": 10.5,
}
HEADING_RANKS = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
}


class DocumentLockedError(RuntimeError):
    pass


@dataclass(frozen=True)
class SectionLocation:
    start: int
    end: int
    matched_style: str


@dataclass(frozen=True)
class EmbeddedImage:
    data: bytes
    caption: str | None = None


@dataclass
class BlockContext:
    config_dir: Path
    embedded_images: list[EmbeddedImage]
    chat_images: dict[str, EmbeddedImage]
    image_ordinal: int = 0


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Generic manual updater for B端后台操作手册.docx. "
            "Uses a small JSON config to update one or more Heading 3 sections, "
            "optionally preserve existing media, and optionally extract chat images."
        )
    )
    parser.add_argument("--config", required=True, help="JSON config path")
    parser.add_argument("--docx", help="Optional docx path override")
    parser.add_argument("--dry-run", action="store_true", help="Validate config and section lookup without saving")
    return parser.parse_args()


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def resolve_path(value: str | None, base_dir: Path) -> Path | None:
    if value is None:
        return None
    path = Path(value)
    if not path.is_absolute():
        path = (base_dir / path).resolve()
    return path


def find_default_docx(root: Path = ROOT) -> Path:
    candidates = sorted(
        path for path in root.glob("*.docx") if not path.name.startswith("~$")
    )
    if not candidates:
        raise FileNotFoundError(f"no docx found in {root}")
    if len(candidates) > 1:
        names = ", ".join(path.name for path in candidates)
        raise RuntimeError(f"multiple docx files found; set config.docx or --docx: {names}")
    return candidates[0]


def ensure_document_writable(docx_path: Path) -> None:
    try:
        with docx_path.open("r+b"):
            pass
    except PermissionError as exc:
        raise DocumentLockedError(
            f"document is in use; close it and rerun the same config: {docx_path}"
        ) from exc


def set_run(run, *, size: float = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def insert_after(paragraph, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    return new_para


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//w:drawing"))


def paragraph_is_media(paragraph) -> bool:
    if paragraph_has_drawing(paragraph):
        return True
    if paragraph.style and paragraph.style.name == "Caption":
        return True
    text = paragraph.text.strip()
    return text.startswith(("图：", "圖："))


def collect_section_images(
    document: Document, location: SectionLocation
) -> list[EmbeddedImage]:
    paragraphs = document.paragraphs
    images: list[EmbeddedImage] = []
    for index in range(location.start + 1, location.end):
        paragraph = paragraphs[index]
        relationship_ids = paragraph._element.xpath(".//a:blip/@r:embed")
        if not relationship_ids:
            continue
        caption = None
        if index + 1 < location.end:
            next_paragraph = paragraphs[index + 1]
            if paragraph_is_media(next_paragraph) and not paragraph_has_drawing(next_paragraph):
                caption = next_paragraph.text.strip() or None
        for relationship_id in relationship_ids:
            image_part = document.part.related_parts.get(relationship_id)
            if image_part is None or not hasattr(image_part, "blob"):
                raise RuntimeError(
                    f"unable to read embedded image relationship: {relationship_id}"
                )
            images.append(EmbeddedImage(data=image_part.blob, caption=caption))
    return images


def default_stop_styles(heading_style: str) -> set[str]:
    rank = HEADING_RANKS.get(heading_style)
    if rank is None:
        return {"Heading 1", "Heading 2", "Heading 3"}
    return {style for style, style_rank in HEADING_RANKS.items() if style_rank <= rank}


def paragraph_is_within_parent(
    paragraphs: list[Paragraph],
    index: int,
    parent_title: str | None,
    parent_heading_style: str,
) -> bool:
    if not parent_title:
        return True
    for parent_index in range(index - 1, -1, -1):
        paragraph = paragraphs[parent_index]
        if paragraph.style.name == parent_heading_style:
            return paragraph.text.strip() == parent_title
    return False


def find_section_range(
    document: Document,
    title: str,
    heading_style: str = "Heading 3",
    stop_styles: list[str] | None = None,
    parent_title: str | None = None,
    parent_heading_style: str = "Heading 2",
    match_index: int | None = None,
    allow_style_fallback: bool = True,
) -> SectionLocation:
    paragraphs = document.paragraphs
    matches = [
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph.style.name == heading_style
        and paragraph.text.strip() == title
        and paragraph_is_within_parent(
            paragraphs, index, parent_title, parent_heading_style
        )
    ]

    if not matches and allow_style_fallback:
        matches = [
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.style.name in HEADING_RANKS
            and paragraph.text.strip() == title
            and paragraph_is_within_parent(
                paragraphs, index, parent_title, parent_heading_style
            )
        ]
        if len(matches) == 1:
            actual_style = paragraphs[matches[0]].style.name
            print(
                f"section_style_fallback title={title} "
                f"requested={heading_style} actual={actual_style}"
            )

    if match_index is not None:
        if match_index < 1 or match_index > len(matches):
            raise RuntimeError(
                f"section match_index out of range: {title} index={match_index} matches={len(matches)}"
            )
        start = matches[match_index - 1]
    elif len(matches) == 1:
        start = matches[0]
    elif not matches:
        raise RuntimeError(f"section not found: {title} style={heading_style}")
    else:
        raise RuntimeError(
            f"section is ambiguous: {title} style={heading_style} matches={len(matches)}; "
            "set parent_title or match_index"
        )

    matched_style = paragraphs[start].style.name
    stop_style_set = set(stop_styles) if stop_styles else default_stop_styles(matched_style)
    end = len(paragraphs)
    for index in range(start + 1, len(paragraphs)):
        if paragraphs[index].style.name in stop_style_set:
            end = index
            break
    return SectionLocation(start=start, end=end, matched_style=matched_style)


def style_size(style: str, default: float = 10) -> float:
    return float(HEADING_SIZES.get(style, default))


def format_heading_paragraph(paragraph, style: str) -> None:
    for run in paragraph.runs:
        set_run(run, size=style_size(style), bold=True, color=BLUE)


def reset_section(
    document: Document,
    title: str,
    keep_existing_media: bool,
    heading_style: str = "Heading 3",
    result_heading_style: str | None = None,
    stop_styles: list[str] | None = None,
    parent_title: str | None = None,
    parent_heading_style: str = "Heading 2",
    match_index: int | None = None,
):
    location = find_section_range(
        document,
        title=title,
        heading_style=heading_style,
        stop_styles=stop_styles,
        parent_title=parent_title,
        parent_heading_style=parent_heading_style,
        match_index=match_index,
    )
    for paragraph in list(document.paragraphs[location.start + 1 : location.end]):
        if keep_existing_media and paragraph_is_media(paragraph):
            continue
        delete_paragraph(paragraph)

    anchor = document.paragraphs[location.start]
    output_style = result_heading_style or location.matched_style
    anchor.style = output_style
    format_heading_paragraph(anchor, output_style)
    return anchor


def resolve_keep_existing_media(config: dict[str, Any], section: dict[str, Any]) -> bool:
    if "keep_existing_media" in section:
        return bool(section["keep_existing_media"])
    if "keep_existing_media" in config:
        return bool(config["keep_existing_media"])
    return False


def add_paragraph_after(anchor, text: str, style: str = "Normal"):
    paragraph = insert_after(anchor, style=style)
    run = paragraph.add_run(text)
    set_run(run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_submodule_title_after(anchor, text: str):
    paragraph = insert_after(anchor, style="Heading 4")
    run = paragraph.add_run(text)
    set_run(run, size=10.5, bold=True, color=BLUE)
    run.font.italic = True
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_heading_after(anchor, text: str, style: str):
    paragraph = insert_after(anchor, style=style)
    run = paragraph.add_run(text)
    set_run(run, size=style_size(style), bold=True, color=BLUE)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullet_after(anchor, label: str | None, text: str):
    paragraph = insert_after(anchor, style="List Bullet")
    if label:
        label_run = paragraph.add_run(label)
        set_run(label_run, size=10, bold=True, color=BLUE)
        body_run = paragraph.add_run("：" + text)
        set_run(body_run, size=10)
    else:
        body_run = paragraph.add_run(text)
        set_run(body_run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_picture_after(anchor, image_source: Path | bytes, width: float = 6.4):
    paragraph = insert_after(anchor, style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if isinstance(image_source, bytes):
        picture_source = BytesIO(image_source)
    else:
        picture_source = str(image_source)
    paragraph.add_run().add_picture(picture_source, width=Inches(width))
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def add_caption_after(anchor, text: str):
    paragraph = insert_after(anchor, style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run(run, size=9, color=CAPTION_GRAY)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def load_chat_images(task: dict[str, Any], config_dir: Path) -> dict[str, EmbeddedImage]:
    if task.get("thread_read_file"):
        _, message, images = chat_images.extract_from_thread_read_file(
            resolve_path(task["thread_read_file"], config_dir),
            message_text=task.get("message_text"),
            message_id=task.get("message_id"),
        )
        source = "thread_read_file"
    else:
        thread_id = task.get("thread_id")
        if not thread_id:
            raise RuntimeError("chat_image task requires thread_read_file or thread_id")
        _, _, message, images = chat_images.locate_from_logs(
            thread_id=thread_id,
            message_text=task.get("message_text"),
            message_id=task.get("message_id"),
            date=task.get("date"),
        )
        source = "local_codex_logs"

    chat_images.validate_count(len(images), task.get("expected_count"))
    names = task.get("names") or task.get("rename_to") or []
    if names and len(names) != len(images):
        raise RuntimeError("chat image names count does not match extracted image count")
    if not names:
        names = [f"chat-{index}.{extension}" for index, (extension, _) in enumerate(images, 1)]

    loaded: dict[str, EmbeddedImage] = {}
    task_id = task.get("id")
    for name, (_, data) in zip(names, images):
        key = str(name)
        aliases = [key]
        if task_id:
            aliases.append(f"{task_id}:{key}")
        for alias in aliases:
            if alias in loaded:
                raise RuntimeError(f"duplicate chat image name: {alias}")
            loaded[alias] = EmbeddedImage(data=data)

    if task.get("output_dir"):
        print("chat_images in_memory output_dir_ignored=true")
    print(
        f"chat_images source={source} message_id={message.get('id', '')} "
        f"count={len(images)} mode=in_memory"
    )
    return loaded


def load_section_chat_images(
    section: dict[str, Any], config_dir: Path
) -> dict[str, EmbeddedImage]:
    loaded: dict[str, EmbeddedImage] = {}
    for task in section.get("chat_image_tasks", []):
        for name, image in load_chat_images(task, config_dir=config_dir).items():
            if name in loaded:
                raise RuntimeError(f"duplicate chat image name across tasks: {name}")
            loaded[name] = image
    return loaded


def image_from_embedded_source(
    block: dict[str, Any], context: BlockContext, ordinal: int
) -> bytes:
    if "source_index" in block:
        source_index = int(block["source_index"])
        if source_index < 1 or source_index > len(context.embedded_images):
            raise RuntimeError(
                f"embedded image source_index out of range: {source_index} "
                f"count={len(context.embedded_images)}"
            )
        return context.embedded_images[source_index - 1].data

    if "source_caption" in block:
        source_caption = str(block["source_caption"])
        matches = [
            image
            for image in context.embedded_images
            if image.caption == source_caption
        ]
        if len(matches) != 1:
            raise RuntimeError(
                f"embedded image caption must match exactly once: {source_caption} "
                f"matches={len(matches)}"
            )
        return matches[0].data

    if ordinal < len(context.embedded_images):
        return context.embedded_images[ordinal].data
    raise RuntimeError(
        f"no embedded image available for block position {ordinal + 1}; "
        "set path, chat_image, source_index, or source_caption"
    )


def resolve_image_source(
    block: dict[str, Any], context: BlockContext
) -> Path | bytes:
    ordinal = context.image_ordinal
    context.image_ordinal += 1

    chat_image_name = block.get("chat_image")
    if chat_image_name:
        image = context.chat_images.get(str(chat_image_name))
        if image is None:
            raise RuntimeError(f"chat image not found: {chat_image_name}")
        return image.data

    if "source_index" in block or "source_caption" in block:
        return image_from_embedded_source(block, context, ordinal)

    path_value = block.get("path")
    if path_value:
        image_path = resolve_path(path_value, context.config_dir)
        if image_path is not None and image_path.exists():
            return image_path
        if ordinal < len(context.embedded_images):
            print(
                f"image_path_missing fallback=embedded source_index={ordinal + 1} "
                f"path={image_path}"
            )
            return context.embedded_images[ordinal].data
        raise FileNotFoundError(image_path)

    return image_from_embedded_source(block, context, ordinal)


def apply_blocks(anchor, blocks: list[dict[str, Any]], context: BlockContext):
    last = anchor
    for block in blocks:
        block_type = block["type"]
        if block_type == "paragraph":
            last = add_paragraph_after(last, block["text"], style=block.get("style", "Normal"))
            continue
        if block_type == "submodule_title":
            last = add_submodule_title_after(last, block["text"])
            continue
        if block_type == "heading":
            last = add_heading_after(last, block["text"], style=block.get("style", "Heading 3"))
            continue
        if block_type == "bullet":
            last = add_bullet_after(last, block.get("label"), block["text"])
            continue
        if block_type in {"image", "document_image", "chat_image"}:
            image_source = resolve_image_source(block, context)
            last = add_picture_after(
                last, image_source=image_source, width=float(block.get("width", 6.4))
            )
            caption = block.get("caption")
            if caption:
                last = add_caption_after(last, caption)
            continue
        if block_type == "module":
            last = add_submodule_title_after(last, block["title"])
            intro = block.get("intro")
            if intro:
                last = add_paragraph_after(last, intro, style=block.get("intro_style", "Normal"))
            if block.get("pre_blocks"):
                last = apply_blocks(last, block.get("pre_blocks", []), context=context)
            image_source = resolve_image_source(block, context)
            last = add_picture_after(
                last, image_source=image_source, width=float(block.get("width", 6.4))
            )
            caption = block.get("caption")
            if caption:
                last = add_caption_after(last, caption)
            last = apply_blocks(last, block.get("blocks", []), context=context)
            continue
        raise RuntimeError(f"unsupported block type: {block_type}")
    return last


def backup_docx(docx_path: Path, tag: str) -> Path:
    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{docx_path.stem}.{tag}-{stamp}.bak{docx_path.suffix}"
    shutil.copy2(docx_path, backup)
    return backup


def section_lookup_kwargs(section: dict[str, Any]) -> dict[str, Any]:
    return {
        "title": section["title"],
        "heading_style": section.get("heading_style", "Heading 3"),
        "stop_styles": section.get("stop_styles"),
        "parent_title": section.get("parent_title"),
        "parent_heading_style": section.get("parent_heading_style", "Heading 2"),
        "match_index": section.get("match_index"),
    }


def validate_blocks(blocks: list[dict[str, Any]]) -> None:
    supported = {
        "paragraph",
        "submodule_title",
        "heading",
        "bullet",
        "image",
        "document_image",
        "chat_image",
        "module",
    }
    for block in blocks:
        block_type = block.get("type")
        if block_type not in supported:
            raise RuntimeError(f"unsupported block type: {block_type}")
        if block_type == "module":
            if not block.get("title"):
                raise RuntimeError("module block requires title")
            validate_blocks(block.get("pre_blocks", []))
            validate_blocks(block.get("blocks", []))


def locate_sections(
    document: Document, sections: list[dict[str, Any]]
) -> list[SectionLocation]:
    locations = [
        find_section_range(document, **section_lookup_kwargs(section))
        for section in sections
    ]
    ordered = sorted(
        ((location.start, location.end, sections[index]["title"]) for index, location in enumerate(locations)),
        key=lambda item: item[0],
    )
    for previous, current in zip(ordered, ordered[1:]):
        if current[0] < previous[1]:
            raise RuntimeError(
                f"overlapping section updates are not supported: {previous[2]} / {current[2]}"
            )
    return locations


def validate_document_sections(
    document: Document, sections: list[dict[str, Any]]
) -> None:
    for section in sections:
        location = find_section_range(document, **section_lookup_kwargs(section))
        section_paragraphs = document.paragraphs[location.start : location.end]
        section_text = "\n".join(paragraph.text for paragraph in section_paragraphs)
        validation = section.get("validate", {})
        for required_text in validation.get("required_text", []):
            if required_text not in section_text:
                raise RuntimeError(
                    f"required text validation failed: {section['title']} / {required_text}"
                )
        if "expected_images" in validation:
            actual_images = sum(
                len(paragraph._element.xpath(".//a:blip/@r:embed"))
                for paragraph in section_paragraphs
            )
            expected_images = int(validation["expected_images"])
            if actual_images != expected_images:
                raise RuntimeError(
                    f"image validation failed: {section['title']} "
                    f"expected={expected_images} actual={actual_images}"
                )


def main() -> None:
    args = parse_args()
    config_path = Path(args.config).resolve()
    config_dir = config_path.parent
    config = load_json(config_path)

    docx_value = args.docx or config.get("docx")
    docx_path = resolve_path(docx_value, config_dir) if docx_value else find_default_docx()
    assert docx_path is not None
    if not docx_path.exists():
        raise FileNotFoundError(docx_path)

    sections = config.get("sections", [])
    if not sections:
        raise RuntimeError("config.sections is empty")
    for section in sections:
        validate_blocks(section.get("blocks", []))

    document = Document(docx_path)
    initial_locations = locate_sections(document, sections)
    embedded_images = [
        collect_section_images(document, location) for location in initial_locations
    ]

    if not args.dry_run:
        ensure_document_writable(docx_path)

    section_chat_images = [
        load_section_chat_images(section, config_dir=config_dir) for section in sections
    ]

    for section_index, section in enumerate(sections):
        anchor = reset_section(
            document,
            title=section["title"],
            keep_existing_media=resolve_keep_existing_media(config, section),
            heading_style=section.get("heading_style", "Heading 3"),
            result_heading_style=section.get("result_heading_style"),
            stop_styles=section.get("stop_styles"),
            parent_title=section.get("parent_title"),
            parent_heading_style=section.get("parent_heading_style", "Heading 2"),
            match_index=section.get("match_index"),
        )
        context = BlockContext(
            config_dir=config_dir,
            embedded_images=embedded_images[section_index],
            chat_images=section_chat_images[section_index],
        )
        apply_blocks(anchor, section.get("blocks", []), context=context)

    validate_document_sections(document, sections)

    if args.dry_run:
        print(f"dry_run ok docx={docx_path}")
        return

    backup = backup_docx(docx_path, tag=config.get("backup_tag", "manual-sections"))
    try:
        document.save(docx_path)
    except PermissionError as exc:
        raise DocumentLockedError(
            f"document became locked during save; close it and rerun the same config: {docx_path}"
        ) from exc

    saved_document = Document(docx_path)
    validate_document_sections(saved_document, sections)

    print(f"updated manual docx={docx_path}")
    print(f"backup={backup}")


if __name__ == "__main__":
    try:
        main()
    except DocumentLockedError as exc:
        print(f"DOCX_LOCKED {exc}")
        raise SystemExit(2) from None
