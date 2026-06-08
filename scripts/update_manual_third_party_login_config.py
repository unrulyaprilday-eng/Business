from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"
ASSET_DIR = ROOT / "custom" / "assets" / "manual-third-party-login-config"
OVERVIEW_IMAGE = ASSET_DIR / "third-party-login-config-overview.png"

FONT_NAME = "微软雅黑"
BLUE = RGBColor(47, 85, 151)
CAPTION_GRAY = RGBColor(112, 126, 146)


CONTENT = [
    (
        "页面介绍：第三方登录配置用于维护站点登录相关的外部服务参数，包括自动注册开关，以及邮箱、Facebook、Whatsapp、Telegram、短信和 Google 等登录或验证服务。页面以配置卡片展示每个服务的启用状态、配置状态和关键参数，运营或技术人员可在上线前逐项核对。",
        "List Bullet",
    ),
    ("IMAGE", str(OVERVIEW_IMAGE), "图：第三方登录配置页面概览。"),
    (
        "自动注册：控制用户通过登录入口进入时，若账号不存在是否自动创建账户。开启前需确认当前站点允许第三方登录自动开户，并与风控、渠道归因、用户协议和注册校验策略保持一致；关闭后，未注册用户应按站点既有注册流程处理。",
        "List Bullet",
    ),
    (
        "配置卡片状态：每张卡片右上角提供启用/关闭开关和“编辑”入口，标题下方显示“已配置”或“未配置”。已配置代表至少存在有效参数；启用状态代表当前服务是否允许在前台登录或验证流程中使用。上线前不要只看“已配置”，还要核对开关是否处于预期状态。",
        "List Bullet",
    ),
    (
        "邮箱配置：用于维护邮箱登录或邮件验证能力，字段包括 SMTP服务器、SMTP端口、SMTP账号、SMTP密码和发件人名称。保存前需确认服务器地址、端口、账号权限、密码或授权码、发件人名称与邮件服务商要求一致。",
        "List Bullet",
    ),
    (
        "Facebook配置：用于维护 Facebook 第三方登录参数，字段包括 AppId 和密钥。修改后应同步确认前台登录页展示、回调域名、应用状态和密钥是否匹配，避免用户点击 Facebook 登录后无法完成授权。",
        "List Bullet",
    ),
    (
        "Whatsapp配置：用于维护 Whatsapp 相关登录或消息验证参数，字段包括 ApiKey、PhoneId 和 BusinessId。该类参数通常与第三方 Business 后台绑定，保存后需用测试账号验证发送链路和号码归属是否正常。",
        "List Bullet",
    ),
    (
        "Telegram配置：用于维护 Telegram 登录或机器人验证参数，字段包括 BotToken 和 Bot名称。BotToken 属于敏感信息，页面展示时会做脱敏处理；更新后需确认机器人名称、Token、前台入口和回调配置一致。",
        "List Bullet",
    ),
    (
        "短信配置：用于维护短信服务参数，字段包括短信提供商（twilio/mock）、AccountSID、AuthToken 和发送号码。生产站点应使用正式服务商参数；mock 仅适合联调或演示环境，切换前需确认不会影响真实用户验证码接收。",
        "List Bullet",
    ),
    (
        "Google配置：用于维护 Google 第三方登录参数，字段包括 ClientId 和密钥。保存后需检查前台域名、OAuth 应用配置、回调地址和启用状态，确保用户能够正常跳转授权并返回站点。",
        "List Bullet",
    ),
    (
        "编辑与保存：点击某张卡片的“编辑”后，该卡片进入编辑态，可修改输入框和开关；点击“取消”放弃本次调整并恢复原展示值；点击“保存”将当前输入同步到展示态。涉及密钥、Token、AuthToken 等敏感字段时，保存前应由技术或运维人员复核来源，避免粘贴测试环境参数。",
        "List Bullet",
    ),
    (
        "注意事项：第三方登录配置会直接影响前台用户登录、注册和验证链路。修改前应确认是否处于低峰期或有回滚方案；修改后至少验证前台入口是否展示、授权跳转是否成功、自动注册是否符合预期、未配置或关闭的服务是否不会被前台继续调用。",
        "List Bullet",
    ),
    (
        "常见问题：若前台看不到某个第三方登录入口，优先检查该服务是否已配置、开关是否开启、登录模板是否展示对应入口；若点击后授权失败，优先检查 AppId、ClientId、Token、密钥、回调域名、第三方应用状态和当前站点域名是否一致。",
        "List Bullet",
    ),
]


def image_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    for item in candidates:
        path = Path(item)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def draw_text(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, font, fill=(31, 41, 55)) -> None:
    draw.text(xy, text, font=font, fill=fill)


def text_center(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill=(31, 41, 55)) -> None:
    left, top, right, bottom = box
    bbox = draw.textbbox((0, 0), text, font=font)
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    draw.text((left + (right - left - width) // 2, top + (bottom - top - height) // 2 - 1), text, font=font, fill=fill)


def rounded(draw: ImageDraw.ImageDraw, box, radius=6, fill=(255, 255, 255), outline=(220, 226, 235), width=1) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def switch(draw: ImageDraw.ImageDraw, x: int, y: int, on: bool) -> None:
    fill = (103, 180, 255) if on else (204, 209, 218)
    draw.rounded_rectangle((x, y, x + 42, y + 22), radius=11, fill=fill)
    if on:
        draw_text(draw, (x + 7, y + 5), "开启", image_font(10, True), (255, 255, 255))
        knob_x = x + 23
    else:
        draw_text(draw, (x + 18, y + 5), "关闭", image_font(10, True), (255, 255, 255))
        knob_x = x + 3
    draw.ellipse((knob_x, y + 3, knob_x + 16, y + 19), fill=(255, 255, 255))


def service_icon(draw: ImageDraw.ImageDraw, x: int, y: int, label: str) -> None:
    colors = {
        "邮箱": (239, 243, 255),
        "Facebook": (235, 242, 255),
        "Whatsapp": (239, 243, 255),
        "Telegram": (239, 243, 255),
        "短信": (239, 243, 255),
        "Google": (239, 243, 255),
    }
    glyphs = {
        "邮箱": "✉",
        "Facebook": "▯",
        "Whatsapp": "☵",
        "Telegram": "✈",
        "短信": "",
        "Google": "⌕",
    }
    rounded(draw, (x, y, x + 40, y + 40), radius=8, fill=colors.get(label, (239, 243, 255)), outline=(239, 243, 255))
    if glyphs[label]:
        text_center(draw, (x, y, x + 40, y + 40), glyphs[label], image_font(22, True), (15, 23, 42))


def draw_card(draw: ImageDraw.ImageDraw, x: int, y: int, title: str, fields: list[tuple[str, str]], configured: bool, on: bool) -> None:
    width = 410
    height = 158
    rounded(draw, (x, y, x + width, y + height), radius=8, fill=(255, 255, 255), outline=(221, 228, 238))
    draw.rounded_rectangle((x, y, x + width, y + 58), radius=8, fill=(251, 252, 255), outline=(221, 228, 238))
    draw.rectangle((x + 1, y + 42, x + width - 1, y + 58), fill=(251, 252, 255))
    service_name = title.replace("配置", "")
    service_icon(draw, x + 18, y + 14, service_name)
    draw_text(draw, (x + 66, y + 18), title, image_font(14, True), (6, 18, 49))
    dot_color = (105, 214, 81) if configured else (215, 220, 229)
    draw.ellipse((x + 66, y + 39, x + 73, y + 46), fill=dot_color)
    draw_text(draw, (x + 79, y + 35), "已配置" if configured else "未配置", image_font(11), (111, 126, 148))
    switch(draw, x + 312, y + 22, on)
    draw_text(draw, (x + 366, y + 25), "↗ 编辑", image_font(11), (38, 135, 255))

    field_x = x + 18
    field_y = y + 82
    col_w = 200
    for index, (label, value) in enumerate(fields):
        fx = field_x + (index % 2) * col_w
        fy = field_y + (index // 2) * 42
        draw_text(draw, (fx, fy), label.upper(), image_font(9), (169, 181, 199))
        draw_text(draw, (fx, fy + 17), value, image_font(12), (61, 75, 96))


def create_overview_image() -> None:
    width, height = 1320, 520
    img = Image.new("RGB", (width, height), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    border = (224, 230, 240)
    muted = (113, 127, 149)

    draw_text(draw, (22, 22), "⚙ 第三方配置", image_font(19, True), (6, 18, 49))
    draw_text(draw, (22, 52), "管理平台第三方服务集成配置，点击编辑按钮修改对应服务的配置信息", image_font(12), muted)
    rounded(draw, (22, 88, width - 22, 142), radius=8, fill=(249, 251, 255), outline=border)
    draw_text(draw, (42, 104), "自动注册", image_font(14, True), (6, 18, 49))
    draw_text(draw, (42, 124), "登录界面，若用户不存在则自动创建账户", image_font(11), muted)
    switch(draw, width - 82, 104, True)

    cards = [
        ("邮箱配置", [("SMTP服务器", "-"), ("SMTP端口", "-"), ("SMTP账号", "-"), ("SMTP密码", "-")], True, False),
        ("Facebook配置", [("AppId", "-"), ("密钥", "-")], True, True),
        ("Whatsapp配置", [("ApiKey", "gf••••••"), ("PhoneId", "df••••••"), ("BusinessId", "••••")], True, False),
        ("Telegram配置", [("BotToken", "76••••••"), ("Bot名称", "jy••••••")], True, True),
        ("短信配置", [("短信提供商", "-"), ("AccountSID", "-"), ("AuthToken", "-"), ("发送号码", "-")], True, True),
        ("Google配置", [("ClientId", "-"), ("密钥", "-")], False, False),
    ]
    positions = [(22, 164), (455, 164), (888, 164), (22, 340), (455, 340), (888, 340)]
    for pos, card in zip(positions, cards):
        draw_card(draw, pos[0], pos[1], *card)

    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    img.save(OVERVIEW_IMAGE)


def set_run_font(run, *, bold: bool = False, color: RGBColor | None = None, size: float = 10) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def style_paragraph(paragraph) -> None:
    text = paragraph.text
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    if "：" in text:
        label, body = text.split("：", 1)
        label_run = paragraph.add_run(label + "：")
        set_run_font(label_run, bold=True, color=BLUE)
        body_run = paragraph.add_run(body)
        set_run_font(body_run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15


def insert_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(anchor, text: str, style: str):
    paragraph = insert_after(anchor, text, style)
    style_paragraph(paragraph)
    return paragraph


def insert_image_after(anchor, image_path: str, caption: str):
    image_paragraph = insert_after(anchor, style="Normal")
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(image_path, width=Inches(6.35))
    image_paragraph.paragraph_format.space_after = Pt(2)

    caption_paragraph = insert_after(image_paragraph, style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, color=CAPTION_GRAY, size=9)
    caption_paragraph.paragraph_format.space_after = Pt(6)
    return caption_paragraph


def find_heading_index(document: Document, title: str) -> int:
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name.startswith("Heading") and paragraph.text.strip() == title:
            return index
    raise ValueError(f"未找到章节：{title}")


def reset_section(document: Document, title: str):
    start = find_heading_index(document, title)
    paragraphs = document.paragraphs
    end = len(paragraphs)
    for index in range(start + 1, len(paragraphs)):
        if paragraphs[index].style.name.startswith("Heading"):
            end = index
            break

    for paragraph in list(paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)

    heading = document.paragraphs[start]
    for run in heading.runs:
        set_run_font(run, bold=True, color=BLUE, size=12.5)
    return heading


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    create_overview_image()
    BACKUP_DIR.mkdir(exist_ok=True)
    backup = BACKUP_DIR / f"B端后台操作手册.third-party-login-config-{datetime.now():%Y%m%d_%H%M%S}.bak.docx"
    shutil.copy2(DOCX, backup)

    document = Document(DOCX)
    cursor = reset_section(document, "第三方登录配置")
    for item in CONTENT:
        if item[0] == "IMAGE":
            cursor = insert_image_after(cursor, item[1], item[2])
        else:
            text, style = item
            cursor = insert_paragraph_after(cursor, text, style)

    document.save(DOCX)
    print(f"backup={backup}")
    print(f"image={OVERVIEW_IMAGE}")
    print("updated=第三方登录配置")


if __name__ == "__main__":
    main()
