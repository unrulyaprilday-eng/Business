from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"
ASSET_DIR = ROOT / "custom" / "assets" / "manual-promotion-activity-list"

LIST_IMAGE = ASSET_DIR / "page-overview.png"
CREATE_IMAGE = ASSET_DIR / "add-activity-modal.png"

FONT_NAME = "微软雅黑"
BLUE = RGBColor(47, 85, 151)
CAPTION_GRAY = RGBColor(112, 126, 146)

LIST_TITLE = "优惠活动列表"
CREATE_TITLE = "新增活动"


def set_run(run, *, size: float = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def insert_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_body_after(anchor: Paragraph, text: str) -> Paragraph:
    paragraph = insert_after(anchor, style="Normal")
    run = paragraph.add_run(text)
    set_run(run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullet_after(anchor: Paragraph, label: str, text: str) -> Paragraph:
    paragraph = insert_after(anchor, style="List Bullet")
    label_run = paragraph.add_run(label)
    set_run(label_run, size=10, bold=True, color=BLUE)
    body_run = paragraph.add_run("：" + text)
    set_run(body_run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_picture_after(anchor: Paragraph, image_path: Path, width: float) -> Paragraph:
    paragraph = insert_after(anchor, style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def add_caption_after(anchor: Paragraph, text: str) -> Paragraph:
    paragraph = insert_after(anchor, style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run(run, size=9, color=CAPTION_GRAY)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def find_heading(document: Document, title: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == title:
            return paragraph
    raise RuntimeError(f"未找到“{title}”章节。")


def clear_section(document: Document, title: str) -> Paragraph:
    paragraphs = document.paragraphs
    start = None
    for index, paragraph in enumerate(paragraphs):
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == title:
            start = index
            break
    if start is None:
        raise RuntimeError(f"未找到“{title}”章节。")

    end = len(paragraphs)
    for index in range(start + 1, len(paragraphs)):
        if paragraphs[index].style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            end = index
            break

    for paragraph in list(paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)

    return document.paragraphs[start]


def remove_heading_section_if_exists(document: Document, title: str) -> None:
    paragraphs = document.paragraphs
    start = None
    for index, paragraph in enumerate(paragraphs):
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == title:
            start = index
            break
    if start is None:
        return

    end = len(paragraphs)
    for index in range(start + 1, len(paragraphs)):
        if paragraphs[index].style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            end = index
            break

    for paragraph in list(paragraphs[start:end]):
        delete_paragraph(paragraph)


def update_list_section(document: Document) -> None:
    anchor = clear_section(document, LIST_TITLE)
    last = add_body_after(
        anchor,
        "优惠活动列表用于统一查看站点已配置的优惠活动，并完成筛选查询、状态维护和新增入口跳转，适合运营人员按活动类型、时间和展示状态跟进活动上线情况。",
    )
    last = add_picture_after(last, LIST_IMAGE, 6.45)
    last = add_caption_after(last, "图：优惠活动列表页面，用于筛选活动记录、维护启用状态并进入新增或编辑操作。")

    bullets = [
        ("活动类型筛选", "页面左上方提供活动类型下拉框，可按拼多多、打码返水、充值送金、游戏排行榜、签到奖励等模板快速缩小查询范围。"),
        ("搜索与重置", "选择筛选条件后点击“搜索”刷新列表；点击“重置”可清空当前条件并恢复查看全部活动。"),
        ("新增活动", "点击“新增活动”进入活动创建页面，继续填写活动模板、活动时间、奖励规则和展示内容。"),
        ("列表字段", "列表集中展示活动名称、游戏类型、活动标题、活动类型、开始时间、结束时间、打码倍数、是否可见、启用状态、排序、操作人和操作，便于一次性核对活动基础信息与当前状态。"),
        ("是否可见", "“是否可见”开关用于控制活动是否在前台展示；下线展示入口前，建议先确认活动奖励是否已发放完成，避免玩家仍可参与但找不到入口。"),
        ("启用状态", "“启用状态”开关用于控制活动是否生效。停用前应先核对活动周期、已参与玩家和关联奖励，避免活动中途关闭造成客诉。"),
        ("排序", "排序值用于调整活动展示顺序，通常数值越小越靠前。修改排序后建议同步检查前台活动位和专题页展示顺序。"),
        ("查看 / 编辑 / 删除", "每条记录右侧提供查看、编辑和删除操作。查看用于核对活动详情，编辑用于调整活动信息，删除前需确认该活动不再使用且不会影响历史记录查询。"),
        ("使用要点", "日常维护时建议重点核对活动时间、活动类型、游戏适用范围、打码倍数以及开关状态，确保后台配置与前台宣传和结算规则一致。"),
    ]

    for label, text in bullets:
        last = add_bullet_after(last, label, text)

    last = add_picture_after(last, CREATE_IMAGE, 5.7)
    last = add_caption_after(last, "图：点击“新增活动”后打开的新增活动弹窗，用于填写活动基础信息、奖励档位、Banner 和活动描述。")

    create_bullets = [
        ("新增活动弹窗", "在优惠活动列表点击“新增活动”后打开新增活动弹窗，用于录入活动模板、标题、时间、奖励规则和展示内容。"),
        ("活动模板", "先选择活动模板，再继续填写对应规则。当前弹窗示例为“首充奖励”，不同模板会带出各自适用的规则区和字段。"),
        ("活动名称与活动标题", "活动名称用于后台识别和列表检索，活动标题用于前台展示或活动页标题，录入时应保持命名清晰并与宣传文案一致。"),
        ("活动时间", "活动时间需同时设置开始时间和结束时间，用于控制活动生效周期。提交前应确认时间覆盖宣传排期、报名期和奖励领取期。"),
        ("奖励类型", "奖励类型支持固定金额和百分比金额。选择固定金额时按档位配置具体奖励；选择百分比金额时应同步核对计算基数和上限规则。"),
        ("派发方式", "派发方式支持立即派发和玩家自领。立即派发适合自动到账场景；玩家自领适合需要玩家主动领取奖励的活动。"),
        ("档位规则", "档位区域用于维护充值金额门槛与奖励金额，支持逐档新增。配置多档奖励时，应保证门槛和奖励递进关系清晰，避免档位重叠或金额倒挂。"),
        ("活动Banner", "活动 Banner 用于前台展示活动入口，上传前应核对尺寸、比例和活动主题文案，避免图片变形或展示内容与活动规则不一致。"),
        ("打码倍数与排序", "打码倍数用于约束奖励相关流水要求，排序用于控制活动展示顺序。两个字段都会直接影响活动参与体验，提交前需与活动规则表一致。"),
        ("外部链接", "外部链接为选填项，可配置活动跳转地址。填写时应使用完整可访问链接，并确认跳转页面与活动内容一致。"),
        ("活动描述", "弹窗底部富文本编辑区用于补充活动规则、参与条件、奖励说明和注意事项。文案应写清领取条件、流水要求、有效期和失效规则，便于前台直接展示。"),
        ("保存与取消", "填写完成后点击“确定”提交新活动；点击“取消”或关闭窗口会放弃当前未保存内容。正式提交前建议逐项复核必填字段、奖励档位和展示素材。"),
    ]

    for label, text in create_bullets:
        last = add_bullet_after(last, label, text)

def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    if not LIST_IMAGE.exists():
        raise FileNotFoundError(LIST_IMAGE)
    if not CREATE_IMAGE.exists():
        raise FileNotFoundError(CREATE_IMAGE)

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{DOCX.stem}.promotion-activity-list-{stamp}.bak.docx"
    shutil.copy2(DOCX, backup)

    document = Document(DOCX)
    remove_heading_section_if_exists(document, CREATE_TITLE)
    update_list_section(document)

    try:
        document.save(DOCX)
    except PermissionError:
        print(f"PERMISSION_ERROR backup={backup}")
        raise

    print(f"updated promotion activity list; backup={backup}; list_image={LIST_IMAGE}; create_image={CREATE_IMAGE}")


if __name__ == "__main__":
    main()
