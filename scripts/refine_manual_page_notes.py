from __future__ import annotations

import html
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "B端后台操作手册.docx"
SOURCE_PATH = ROOT / "backups" / "B端后台操作手册.simplify-page-notes-20260608_153804.bak.docx"


ALIASES = {
    "站点应用配置": "站点配置.html",
    "游戏人数据表": "游戏人数报表.html",
    "游戏厂商统计": "游戏厂商报表.html",
    "充值留存报表": "付费留存率报表.html",
    "票券使用记录": "优惠活动-票券使用记录.html",
    "代理配置": "代理中心_代理配置.html",
    "日志管理": "管理员操作日志.html",
}

FIELD_HINTS = [
    "编号",
    "名称",
    "标题",
    "类型",
    "状态",
    "时间",
    "日期",
    "金额",
    "账号",
    "会员",
    "玩家",
    "代理",
    "商户",
    "站点",
    "渠道",
    "域名",
    "币种",
    "排序",
    "开关",
    "备注",
    "操作人",
    "充值",
    "提现",
    "投注",
    "输赢",
    "返水",
    "活动",
    "游戏",
    "角色",
    "权限",
    "IP",
    "ID",
]

ACTION_WORDS = [
    "搜索",
    "查询",
    "重置",
    "新增",
    "编辑",
    "修改",
    "删除",
    "查看",
    "详情",
    "导出",
    "保存",
    "提交",
    "取消",
    "启用",
    "停用",
    "审核",
    "通过",
    "驳回",
    "绑定",
    "上传",
    "配置",
    "预览",
]

DROP_WORDS = (
    "后续补充",
    "后续",
    "补充",
    "具体说明可在",
    "页面截图",
)

GENERIC_SENTENCES = (
    "进入页面后先通过筛选条件定位数据或配置项，再执行页面支持的新增、编辑、启用、停用、审核、导出或查看详情等操作。",
    "按页面截图补齐关键字段、按钮说明、弹窗流程、保存校验和常见问题。",
)

BAD_TOKENS = {
    "function",
    "return",
    "resources",
    "transparent",
    "reload",
    "Other",
    "html",
    "gif",
    "axure",
    "utils",
}


def has_drawing(paragraph: Paragraph) -> bool:
    return any(run._element.xpath(".//w:drawing") for run in paragraph.runs)


def is_caption(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    return paragraph.style.name == "Caption" or bool(re.match(r"^图\d*[：:]", text))


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_after(anchor: Paragraph, text: str, style: str = "List Bullet") -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    paragraph.style = style
    run = paragraph.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10)
    return paragraph


def set_fonts(document: Document) -> None:
    for style_name in ("Normal", "List Bullet", "List Number"):
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = "微软雅黑"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            style.font.size = Pt(10)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = "微软雅黑"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            style.font.color.rgb = RGBColor(31, 78, 121)


def normalize_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text.strip())
    text = re.sub(r"^[-•\d.、\s]+", "", text)
    for prefix in ("页面介绍", "页面说明", "模块定位", "字段说明", "操作流程", "注意事项", "常见问题", "常用操作", "使用要点", "上线检查"):
        text = re.sub(rf"^{re.escape(prefix)}[：:]\s*", "", text)
    for bad in GENERIC_SENTENCES:
        text = text.replace(bad, "")
    text = text.replace("具体说明可在后续补充截图时继续细化。", "")
    text = text.replace("按页面截图补齐关键字段、按钮说明、弹窗流程、保存校验和常见问题。", "")
    text = text.replace("后续补充截图时，可在对应组件下继续追加字段截图、弹窗说明或配置示例。", "")
    text = re.sub(r"用于处理“([^”]+)”相关后台业务[，,。]*", r"用于维护\1相关数据。", text)
    text = re.sub(r"^(筛选条件|列表字段|操作方式|按钮说明|入口|用途|字段|查看详情|新增|编辑|修改|删除)[：:]\s*", "", text)
    text = re.sub(r"[，,；;：:]+。", "。", text)
    return text.strip()


def useful_line(text: str) -> bool:
    if not text:
        return False
    return not any(word in text for word in DROP_WORDS)


def split_sentences(lines: list[str]) -> list[str]:
    sentences: list[str] = []
    for line in lines:
        line = normalize_text(line)
        if not useful_line(line):
            continue
        for sentence in re.split(r"(?<=[。！？])", line):
            sentence = normalize_text(sentence)
            if useful_line(sentence) and len(sentence) >= 8 and sentence not in sentences:
                if not sentence.endswith(("。", "！", "？")):
                    sentence += "。"
                sentences.append(sentence)
    return sentences


def limit_sentence(text: str, limit: int = 120) -> str:
    text = normalize_text(text)
    if len(text) > limit:
        text = text[: limit - 1].rstrip("，,；;、 ") + "。"
    if text and not text.endswith(("。", "！", "？")):
        text += "。"
    return text


def section_bounds(document: Document, start: int) -> int:
    style = document.paragraphs[start].style.name
    if style == "Heading 3":
        stop_styles = {"Heading 1", "Heading 2", "Heading 3"}
    elif style == "Heading 2":
        stop_styles = {"Heading 1", "Heading 2"}
    else:
        stop_styles = {"Heading 1"}
    for index in range(start + 1, len(document.paragraphs)):
        paragraph = document.paragraphs[index]
        if paragraph.style.name in stop_styles and paragraph.text.strip():
            return index
    return len(document.paragraphs)


def page_starts(document: Document) -> list[tuple[int, str]]:
    starts: list[tuple[int, str]] = []
    in_page_notes = False
    for index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        style = paragraph.style.name
        if style == "Heading 1":
            in_page_notes = False
        elif style == "Heading 2" and text == "页面说明":
            in_page_notes = True
        elif style == "Heading 2" and text == "默认看板":
            starts.append((index, text))
        elif style == "Heading 3" and in_page_notes:
            starts.append((index, text))
    return starts


def collect_source_sections() -> dict[str, list[str]]:
    source = Document(SOURCE_PATH)
    sections: dict[str, list[str]] = {}
    for start, title in page_starts(source):
        end = section_bounds(source, start)
        lines: list[str] = []
        for paragraph in source.paragraphs[start + 1 : end]:
            text = paragraph.text.strip()
            if text and not is_caption(paragraph) and not paragraph.style.name.startswith("Heading"):
                lines.append(text)
        sections[title] = lines
    return sections


def prototype_path(title: str) -> Path | None:
    filename = ALIASES.get(title, f"{title}.html")
    path = ROOT / filename
    return path if path.exists() else None


def prototype_tokens(title: str) -> list[str]:
    path = prototype_path(title)
    if not path:
        return []
    source = path.read_text(encoding="utf-8", errors="ignore")
    source = re.sub(r"<script[\s\S]*?</script>", " ", source, flags=re.I)
    source = re.sub(r"<style[\s\S]*?</style>", " ", source, flags=re.I)
    source = re.sub(r"<[^>]+>", " ", source)
    source = html.unescape(source)
    source = re.sub(r"\s+", " ", source)
    tokens: list[str] = []
    for token in re.split(r"\s+|[|/\\]+", source):
        token = token.strip(" ：:，,。.；;（）()[]【】<>\"'“”")
        if not token or token == title or len(token) > 24:
            continue
        if len(token) == 1 and token not in ACTION_WORDS:
            continue
        if re.fullmatch(r"[\d:./_-]+", token):
            continue
        if any(bad.lower() in token.lower() for bad in BAD_TOKENS):
            continue
        if token not in tokens:
            tokens.append(token)
        if len(tokens) >= 160:
            break
    return tokens


def dedupe(values: list[str], limit: int) -> list[str]:
    result: list[str] = []
    for value in values:
        value = normalize_text(value).strip(" ：:，,。.；;（）()[]【】")
        if not value or value in result or len(value) > 22:
            continue
        if any(bad.lower() in value.lower() for bad in BAD_TOKENS):
            continue
        result.append(value)
        if len(result) >= limit:
            break
    return result


def labels_from_lines(lines: list[str]) -> list[str]:
    labels: list[str] = []
    for line in lines:
        line = normalize_text(line)
        match = re.match(r"^([^：:]{2,18})[：:]", line)
        if match:
            label = match.group(1)
            if label not in {"页面说明", "字段说明", "操作流程", "注意事项", "常见问题", "操作方式"}:
                labels.append(label)
    return dedupe(labels, 10)


def choose_purpose(title: str, sentences: list[str]) -> str:
    for sentence in sentences:
        if "用于" in sentence or "支持" in sentence or title in sentence:
            return limit_sentence(sentence, 130)
    return f"{title}用于维护后台中与{title}相关的配置、记录或数据查询。"


def choose_fields(title: str, lines: list[str]) -> list[str]:
    labels = labels_from_lines(lines)
    tokens = prototype_tokens(title)
    html_fields = dedupe([t for t in tokens if any(h in t for h in FIELD_HINTS)], 10)
    fields = labels or html_fields
    return fields[:8]


def choose_actions(title: str, lines: list[str]) -> list[str]:
    text = "\n".join(lines)
    tokens = prototype_tokens(title)
    from_text = [word for word in ACTION_WORDS if word in text]
    from_html = [word for word in ACTION_WORDS if word in tokens]
    return dedupe(from_text + from_html, 8)


def choose_flow(actions: list[str], sentences: list[str], purpose: str) -> str:
    flow_candidates = [
        sentence
        for sentence in sentences
        if any(word in sentence for word in ("点击", "进入", "选择", "填写", "保存", "提交", "查看", "筛选"))
        and "后续" not in sentence
        and sentence != purpose
        and not (sentence.startswith("用于") or "用于维护" in sentence[:20])
    ]
    if flow_candidates:
        return limit_sentence(flow_candidates[0], 135)
    if actions:
        return f"进入页面后先按业务条件筛选数据，再根据需要执行{'、'.join(actions[:6])}等操作，保存或提交后回到列表核对结果。"
    return "进入页面后先确认当前站点和筛选条件，再查看列表、配置项或详情内容，完成调整后核对页面展示结果。"


def choose_note(sentences: list[str]) -> str:
    for sentence in sentences:
        if any(word in sentence for word in ("确认", "检查", "避免", "影响", "生效", "上线", "复核", "权限", "口径")):
            return limit_sentence(sentence, 135)
    return "修改前应确认影响范围、当前站点和账号权限；涉及线上入口、资金、订单或风控结果的内容，保存后应回到前台或列表做一次复核。"


def build_notes(title: str, source_lines: list[str]) -> list[str]:
    lines = [normalize_text(line) for line in source_lines if useful_line(normalize_text(line))]
    sentences = split_sentences(lines)
    fields = choose_fields(title, lines)
    actions = choose_actions(title, lines)

    purpose = choose_purpose(title, sentences)
    notes = [f"页面用途：{purpose}"]
    if fields:
        notes.append(f"关键字段：重点查看或维护{'、'.join(fields)}等信息，具体必填项以页面表单和列表展示为准。")
    else:
        notes.append("关键字段：重点查看页面筛选条件、列表字段、状态标识和操作列，必要时进入详情或编辑弹窗核对完整信息。")
    notes.append(f"操作流程：{choose_flow(actions, sentences, purpose)}")
    note = choose_note([sentence for sentence in sentences if sentence != purpose])
    if normalize_text(note) == normalize_text(purpose):
        note = "修改前应确认影响范围、当前站点和账号权限；涉及线上入口、资金、订单或风控结果的内容，保存后应回到前台或列表做一次复核。"
    notes.append(f"使用要点：{note}")
    return notes


def remove_chapter_from(document: Document, heading_prefix: str) -> None:
    start = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "Heading 1" and paragraph.text.strip().startswith(heading_prefix):
            start = index
            break
    if start is None:
        return
    for paragraph in reversed(document.paragraphs[start:]):
        delete_paragraph(paragraph)


def rewrite_section(document: Document, start: int, notes: list[str]) -> None:
    end = section_bounds(document, start)
    for paragraph in reversed(document.paragraphs[start + 1 : end]):
        if has_drawing(paragraph) or is_caption(paragraph):
            continue
        delete_paragraph(paragraph)
    anchor = document.paragraphs[start]
    for note in reversed(notes):
        insert_after(anchor, note, "List Bullet")


def cleanup_sections(document: Document) -> None:
    starts = page_starts(document)
    for start, _title in reversed(starts):
        end = section_bounds(document, start)
        for paragraph in reversed(document.paragraphs[start + 1 : end]):
            text = paragraph.text.strip()
            if has_drawing(paragraph) or is_caption(paragraph):
                continue
            if text.startswith(("页面用途：", "关键字段：", "操作流程：", "使用要点：")):
                continue
            delete_paragraph(paragraph)


def main() -> None:
    source_sections = collect_source_sections()
    document = Document(DOCX_PATH)
    set_fonts(document)
    starts = page_starts(document)
    for start, title in reversed(starts):
        notes = build_notes(title, source_sections.get(title, []))
        rewrite_section(document, start, notes)
    cleanup_sections(document)
    remove_chapter_from(document, "十六、典型业务流程")
    document.save(DOCX_PATH)
    print(f"refined_sections={len(starts)}")


if __name__ == "__main__":
    main()
