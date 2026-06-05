# -*- coding: utf-8 -*-
from datetime import datetime
from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


DOCX_PATH = Path("B端后台操作手册.docx")
FONT_NAME = "微软雅黑"
DEEP_BLUE = RGBColor(37, 64, 97)
MUTED = RGBColor(96, 108, 128)

RIGHT_NAV_ITEMS = [
    "点击主题开关，可切换页面显示模式。",
    "点击声音图标，可开启或关闭系统提示音。",
    "点击背景样式图标，可切换页面背景显示效果。",
    "点击界面尺寸选项，可在默认、中、小、迷你之间切换页面显示尺寸。",
    "点击语言入口，可切换系统显示语言。",
    "点击通知图标，可查看系统通知或消息提醒。",
    "点击站点下拉框，可切换当前操作和查看的站点。",
    "点击头像图标，可退出当前登录账号。",
]

COMMON_ITEMS = [
    "筛选：输入时间、账号、状态、渠道、类型等条件后点击查询；结果不符合预期时先重置再重新查询。",
    "表格：表格通常承载列表数据，支持查看状态、金额、时间、操作人和行内操作。",
    "分页：列表数据较多时通过分页切换；排查问题时注意当前页码和每页条数。",
    "弹窗：新增、编辑、审核、确认等操作通常在弹窗中完成，关闭前需确认是否保存。",
    "状态开关：启用后配置生效，停用后配置隐藏或失效；停用前应确认是否影响线上业务。",
]


def set_run_font(run, size=10, bold=False, color=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def clear_paragraph(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def replace_text(paragraph, text, size=10, color=None):
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color)


def add_after(anchor, text="", style=None):
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = anchor._parent.add_paragraph()
    paragraph._p = new_p
    if style:
        paragraph.style = style
    if text:
        run = paragraph.add_run(text)
        set_run_font(run)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def move_after(paragraph, anchor):
    element = paragraph._element
    element.getparent().remove(element)
    anchor._element.addnext(element)
    return paragraph


def main():
    backup = DOCX_PATH.with_suffix(f".docx.nav-{datetime.now().strftime('%Y%m%d-%H%M%S')}.bak")
    shutil.copy2(DOCX_PATH, backup)

    doc = Document(str(DOCX_PATH))

    caption2_index = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("图2："))
    caption2 = doc.paragraphs[caption2_index]
    image2_index = caption2_index - 1
    image2 = doc.paragraphs[image2_index]

    replace_text(caption2, "图2：顶部导航栏右侧区域。", size=9, color=MUTED)

    menu_heading = next(p for p in doc.paragraphs if p.text.strip() == "菜单导航")
    common_heading = next(p for p in doc.paragraphs if p.text.strip().startswith("筛选、表格"))
    common_heading.style = "Heading 2"

    # Move image and caption to the end of menu navigation.
    caption2_index = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("图2："))
    caption2 = doc.paragraphs[caption2_index]
    image2 = doc.paragraphs[caption2_index - 1]
    common_heading = next(p for p in doc.paragraphs if p.text.strip().startswith("筛选、表格"))
    common_heading_index = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("筛选、表格"))
    left_last = doc.paragraphs[common_heading_index - 1]
    move_after(image2, left_last)
    move_after(caption2, image2)

    last = caption2
    for item in RIGHT_NAV_ITEMS:
        last = add_after(last, item, style="List Bullet")

    # Remove existing common-operation bullets between common heading and next chapter.
    common_heading_index = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("筛选、表格"))
    start = common_heading_index + 1
    end = next(i for i, p in enumerate(doc.paragraphs[start:], start) if p.style.name == "Heading 1")
    for p in list(doc.paragraphs[start:end]):
        delete_paragraph(p)

    # Rebuild common-operation bullets after the common heading.
    common_heading = next(p for p in doc.paragraphs if p.text.strip().startswith("筛选、表格"))
    last = common_heading
    for item in COMMON_ITEMS:
        last = add_after(last, item, style="List Bullet")

    try:
        doc.save(str(DOCX_PATH))
    except PermissionError:
        print("blocked=word_file_open")
        print("请先关闭 B端后台操作手册.docx，然后回复我继续。")
        return

    print(f"backup={backup}")
    print("updated=nav_section")


if __name__ == "__main__":
    main()
