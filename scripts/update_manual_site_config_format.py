from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX_PATH = "B端后台操作手册.docx"
FONT_NAME = "微软雅黑"


def set_font(run, size=None, bold=None):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


doc = Document(DOCX_PATH)

start = next(
    i for i, paragraph in enumerate(doc.paragraphs)
    if paragraph.text.strip() == "站点配置" and i > 60
)
end = next(
    i for i, paragraph in enumerate(doc.paragraphs[start + 1:], start + 1)
    if paragraph.text.strip() == "运营中心"
)

for i in range(start, end):
    paragraph = doc.paragraphs[i]
    text = paragraph.text.strip()

    for run in paragraph.runs:
        set_font(run)

    if i == start:
        for run in paragraph.runs:
            set_font(run, size=12.5, bold=True)
        continue

    if paragraph.style.name.startswith("Heading") and text:
        for run in paragraph.runs:
            set_font(run, size=10.5, bold=True)

doc.save(DOCX_PATH)
print(f"updated paragraphs {start}-{end - 1}")
