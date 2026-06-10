from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"
ASSET_DIR = ROOT / "custom" / "assets" / "manual-task-center-original"

FONT_NAME = "微软雅黑"
BLUE = RGBColor(47, 85, 151)
CAPTION_GRAY = RGBColor(112, 126, 146)

IMAGE_SPECS = [
    ("task-center-newbie-list.png", "图：任务中心页面的新人福利页签，用于维护新手任务奖励、领取方式和启停状态。", 6.45),
    ("task-center-daily-edit.png", "图：每日任务页签，支持按累计充值、累计打码和单局打码维护阶梯奖励。", 6.45),
    ("task-center-weekly-edit.png", "图：每周任务页签，按周配置奖励门槛、奖励金额、奖励活跃度和任务说明。", 6.45),
    ("task-center-rule-modal.png", "图：规则设置弹窗，用于配置打码倍数和任务启用状态。", 3.6),
    ("task-center-chest-list.png", "图：活跃度宝箱页签，用于查看宝箱门槛、奖励类型和金额范围。", 6.45),
    ("task-center-chest-edit.png", "图：宝箱编辑弹窗，用于维护宝箱名称、所需活跃度和奖励金额。", 4.2),
]


def set_run(run, *, size: float = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def insert_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_heading_after(anchor, text: str, level: int):
    paragraph = insert_after(anchor, style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_run(run, size=10.5 if level >= 4 else 12.5, bold=True, color=BLUE)
    paragraph.paragraph_format.space_after = Pt(5)
    return paragraph


def add_body_after(anchor, text: str):
    paragraph = insert_after(anchor, style="Normal")
    run = paragraph.add_run(text)
    set_run(run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullet_after(anchor, label: str, text: str):
    paragraph = insert_after(anchor, style="List Bullet")
    label_run = paragraph.add_run(label)
    set_run(label_run, size=10, bold=True, color=BLUE)
    body_run = paragraph.add_run("：" + text)
    set_run(body_run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_picture_after(anchor, image_path: Path, width: float):
    paragraph = insert_after(anchor, style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def add_caption_after(anchor, text: str):
    paragraph = insert_after(anchor, style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run(run, size=9, color=CAPTION_GRAY)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def find_task_center_range(document: Document) -> tuple[int, int]:
    start = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == "任务中心":
            start = index
            break
    if start is None:
        raise RuntimeError("未找到“任务中心”章节。")

    end = len(document.paragraphs)
    for index in range(start + 1, len(document.paragraphs)):
        if document.paragraphs[index].style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            end = index
            break
    return start, end


def reset_task_center_section(document: Document):
    start, end = find_task_center_range(document)
    for paragraph in list(document.paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)
    anchor = document.paragraphs[start]
    for run in anchor.runs:
        set_run(run, size=12.5, bold=True, color=BLUE)
    return anchor


def maybe_add_original_image(anchor, filename: str, caption: str, width: float):
    image_path = ASSET_DIR / filename
    if not image_path.exists():
        return anchor
    last = add_picture_after(anchor, image_path, width)
    return add_caption_after(last, caption)


def fill_task_center_section(document: Document) -> None:
    anchor = reset_task_center_section(document)

    last = add_body_after(
        anchor,
        "任务中心用于集中维护新人福利、每日任务、每周任务和活跃度宝箱。运营人员可在不同页签下配置任务门槛、奖励金额、奖励活跃度、任务说明、领取方式和启停状态，并通过规则设置统一维护打码倍数等公共参数。",
    )

    last = maybe_add_original_image(last, *IMAGE_SPECS[0])

    last = add_heading_after(last, "新人福利", 4)
    for label, text in [
        ("页签用途", "新人福利页签按任务条件展示新手阶段可领取的任务，如注册账号、首笔充值、绑定邮箱、绑定手机号和绑定社交账号等，适合统一维护新玩家成长奖励。"),
        ("列表字段", "列表通常包含任务条件、图标、奖励金额、奖励活跃度、任务介绍、领取方式、是否开启、提示气泡和操作。领取方式用于区分自动派发和手动领取；启用和提示气泡开关用于控制前台是否展示及是否提示。"),
        ("修改任务", "在目标任务行点击“修改”后，更新宣传图、奖励金额、奖励活跃度、打码倍数和任务介绍，并按需要调整是否开启和提示气泡。保存前应确认图片规格、奖励数值和介绍文案已与当前活动口径一致。"),
        ("使用要点", "奖励金额与奖励活跃度属于直接发放内容，修改后要同步核对预算、站内展示和奖励发放规则；对暂未上线的任务，应先关闭开关，避免玩家提前看到入口。"),
    ]:
        last = add_bullet_after(last, label, text)

    last = add_heading_after(last, "每日任务", 4)
    last = maybe_add_original_image(last, *IMAGE_SPECS[1])
    for label, text in [
        ("奖励维度", "每日任务支持在累计充值、累计打码和单局打码之间切换，分别维护不同的任务口径和阶梯配置。切换维度后，应分别检查每一档门槛和奖励是否完整。"),
        ("编辑方式", "点击页面底部“修改”进入编辑态后，可逐行维护累计金额、奖励金额、奖励活跃度和任务介绍。任务介绍建议直接写明达成条件，便于前台玩家理解领取要求。"),
        ("保存规则", "每一行至少要配置奖励金额或奖励活跃度中的一种；两项都为空的数据在保存时会被过滤。保存前要检查阶梯是否按从低到高递增，避免门槛交叉或奖励倒挂。"),
    ]:
        last = add_bullet_after(last, label, text)

    last = add_heading_after(last, "每周任务", 4)
    last = maybe_add_original_image(last, *IMAGE_SPECS[2])
    for label, text in [
        ("配置逻辑", "每周任务的维护方式与每日任务一致，同样支持累计充值、累计打码和单局打码三类任务口径，但奖励门槛通常按周目标设置，数值会高于每日任务。"),
        ("核对重点", "调整每周阶梯时，应同时核对奖励金额、奖励活跃度和说明文案，确保周任务目标与站点活动周期一致；如果周任务与每日任务同时存在，需避免同一行为重复激励过高。"),
    ]:
        last = add_bullet_after(last, label, text)

    last = add_heading_after(last, "规则设置", 4)
    last = maybe_add_original_image(last, *IMAGE_SPECS[3])
    for label, text in [
        ("打码倍数", "规则设置弹窗用于维护任务对应的打码倍数。修改前应确认平台当前流水要求，避免奖励到账后无法满足提款或活动结算条件。"),
        ("启用状态", "是否启用用于控制当前任务规则是否生效。停用前要确认前台是否仍有展示入口，以及是否存在未完成但仍需结算的历史任务。"),
    ]:
        last = add_bullet_after(last, label, text)

    last = add_heading_after(last, "活跃度宝箱", 4)
    last = maybe_add_original_image(last, *IMAGE_SPECS[4])
    for label, text in [
        ("列表管理", "活跃度宝箱页签用于维护按活跃度领取的奖励宝箱，列表展示宝箱名称、所需活跃度、奖励类型、奖励金额和操作。可通过“新增宝箱”补充新档位，也可对现有档位执行修改或删除。"),
        ("奖励类型", "奖励类型分为随机和固定。随机奖励通常需要维护起止金额区间；固定奖励则直接配置一个确定金额。配置前应确认是否与活动预算、前台文案和派奖逻辑一致。"),
    ]:
        last = add_bullet_after(last, label, text)

    last = maybe_add_original_image(last, *IMAGE_SPECS[5])
    for label, text in [
        ("弹窗字段", "宝箱编辑弹窗通常需要填写宝箱名称、所需活跃度、奖励类型和奖励金额。宝箱名称建议使用便于识别的档位名称，所需活跃度需与前台活跃度产出规则保持一致。"),
        ("提交流程", "新增或修改后点击“提交”保存；如果本次只是查看配置，可点击“取消”关闭。保存前要检查名称长度、活跃度门槛和奖励区间上下限，避免出现空值、倒置或重复档位。"),
    ]:
        last = add_bullet_after(last, label, text)

    add_bullet_after(
        last,
        "注意事项",
        "任务中心同时影响前台任务展示和奖励发放，调整任何门槛、奖励或启停状态前，都应先确认活动预算、任务周期、前台文案和结算规则，必要时与运营、风控和财务口径同步后再保存。",
    )


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = BACKUP_DIR / f"{DOCX.stem}.task-center-original-only-{stamp}.bak.docx"
    shutil.copy2(DOCX, backup)

    document = Document(DOCX)
    fill_task_center_section(document)

    try:
        document.save(DOCX)
    except PermissionError:
        print(f"PERMISSION_ERROR backup={backup}")
        raise

    existing = sum(1 for name, _, _ in IMAGE_SPECS if (ASSET_DIR / name).exists())
    print(f"updated task center original-only; backup={backup}; existing_original_images={existing}")


if __name__ == "__main__":
    main()
