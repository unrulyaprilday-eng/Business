# -*- coding: utf-8 -*-
from datetime import datetime
from pathlib import Path
import shutil

from docx import Document


DOCX_PATH = Path("B端后台操作手册.docx")

REPLACEMENTS = {
    "图2：后台通用工具栏、站点切换与账号入口。": "图2：顶部导航栏右侧区域，用于说明显示模式、声音、背景样式、界面尺寸、语言、通知、站点切换和账号入口。",
    "站点设置页面补充说明": "站点设置",
    "首页与模板管理补充说明": "首页与模板管理",
}


def replace_paragraph_text(paragraph, new_text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def main():
    backup = DOCX_PATH.with_suffix(f".docx.labels-{datetime.now().strftime('%Y%m%d-%H%M%S')}.bak")
    shutil.copy2(DOCX_PATH, backup)

    doc = Document(str(DOCX_PATH))
    changed = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text in REPLACEMENTS:
            replace_paragraph_text(paragraph, REPLACEMENTS[text])
            changed.append((index, text, REPLACEMENTS[text]))

    doc.save(str(DOCX_PATH))
    print(f"backup={backup}")
    for index, old, new in changed:
        print(f"{index}: {old} -> {new}")
    print(f"changed={len(changed)}")


if __name__ == "__main__":
    main()
