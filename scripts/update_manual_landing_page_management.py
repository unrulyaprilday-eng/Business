from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "B端后台操作手册.docx"
BACKUP_DIR = ROOT / "backups"
ASSET_DIR = ROOT / "custom" / "assets" / "manual-landing-page-management-original"

TITLE = "落地页管理"
FONT_NAME = "微软雅黑"
BLUE = RGBColor(47, 85, 151)
CAPTION_GRAY = RGBColor(112, 126, 146)

IMAGE_SPECS = [
    ("list-table.png", "图：落地页管理列表页面，用于查看模板信息、下载地址、评分、截图和行内操作。", 6.45),
    ("add-modal.png", "图：新增落地页弹窗，用于维护模板样式、图标、下载地址、评分、截图和应用介绍。", 4.9),
]


def set_run(run, *, size: float = 10, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def insert_after(paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_body_after(anchor: Paragraph, text: str) -> Paragraph:
    paragraph = insert_after(anchor, style="Normal")
    run = paragraph.add_run(text)
    set_run(run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_bullet_after(anchor: Paragraph, label: str, text: str) -> Paragraph:
    paragraph = insert_after(anchor, style="List Bullet")
    label_run = paragraph.add_run(label)
    set_run(label_run, size=10, bold=True, color=BLUE)
    body_run = paragraph.add_run("：" + text)
    set_run(body_run, size=10)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    return paragraph


def add_picture_after(anchor: Paragraph, image_path: Path, width: float) -> Paragraph:
    paragraph = insert_after(anchor, style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def add_caption_after(anchor: Paragraph, text: str) -> Paragraph:
    paragraph = insert_after(anchor, style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run(run, size=9, color=CAPTION_GRAY)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def maybe_add_original_image(anchor: Paragraph, filename: str, caption: str, width: float) -> Paragraph:
    image_path = ASSET_DIR / filename
    if not image_path.exists():
        return anchor
    last = add_picture_after(anchor, image_path, width)
    return add_caption_after(last, caption)


def reset_section(document: Document) -> Paragraph:
    start = None
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name == "Heading 3" and paragraph.text.strip() == TITLE:
            start = index
            break
    if start is None:
        raise RuntimeError(f"未找到《{TITLE}》章节。")

    end = len(document.paragraphs)
    for index in range(start + 1, len(document.paragraphs)):
        if document.paragraphs[index].style.name in {"Heading 1", "Heading 2", "Heading 3"}:
            end = index
            break

    for paragraph in list(document.paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)

    anchor = document.paragraphs[start]
    for run in anchor.runs:
        set_run(run, size=10.5, bold=True, color=BLUE)
    return anchor


def fill_section(document: Document) -> None:
    anchor = reset_section(document)

    last = add_body_after(
        anchor,
        "落地页管理用于维护推广落地页的模板样式、应用名称、下载地址、评分、截图和详情介绍，是运营维护投放页素材和跳转链路的统一入口。新增、编辑或删除落地页后，会直接影响渠道投放页展示内容，提交前应同步复核模板名称、应用图标、下载地址和截图素材是否与当前投放计划一致。",
    )

    last = maybe_add_original_image(last, *IMAGE_SPECS[0])

    for label, body in [
        ("筛选查询", "页面顶部支持按名称输入关键字后执行“搜索”，也可通过“重置”快速清空条件并恢复全部数据，适合在模板数量较多时先定位目标落地页。"),
        ("新增入口", "点击“新增”可打开新增落地页弹窗，录入模板样式、模板名称、应用图标、下载地址、评分、截图和应用详细介绍。新增完成后，新记录会出现在列表中，便于继续核对素材和链接。"),
        ("列表字段", "列表展示 ID、模板名称、样式、图标、APP 名字、安卓下载地址、评分和截图缩略图，便于运营快速确认当前投放页使用的模板和素材版本。"),
        ("下载地址", "安卓下载地址字段用于展示当前落地页跳转链接。维护时应填写完整可访问地址，并在保存前检查渠道参数、页面路径和打开结果是否正确。"),
        ("截图预览", "截图列用于展示当前落地页的页面预览图，适合快速核对展示风格、页面主题和素材是否与投放内容一致。上传或替换截图后，应确认缩略图顺序与实际展示顺序一致。"),
        ("编辑与删除", "点击“编辑”可修改当前落地页的模板信息和素材；点击“删除”会移除当前记录。删除前应确认该模板已停用或有替代页面，避免影响正在使用的推广链接。"),
    ]:
        last = add_bullet_after(last, label, body)

    last = maybe_add_original_image(last, *IMAGE_SPECS[1])

    for label, body in [
        ("模板样式", "模板样式用于区分当前落地页使用的展示风格，应先选择对应模板，再继续填写名称和素材，避免后续页面风格与投放渠道不匹配。"),
        ("模板名称", "模板名称建议填写便于识别的投放页名称或活动名称，方便后续通过名称筛选快速找到目标模板。"),
        ("APP 图标", "APP 图标要求尺寸为 64×64，支持等比例缩放。上传前应确认图标清晰、品牌版本正确，避免出现旧图标或模糊素材。"),
        ("下载地址", "安卓下载地址为必填项，iOS 下载地址可按投放需要补充。保存前应分别检查链接是否可访问、是否带齐渠道参数，以及是否跳转到正确的下载页面。"),
        ("APP 评分", "APP 评分用于展示落地页应用评分，录入时应控制在 0 到 5 分范围内，并与当前投放素材口径保持一致。"),
        ("截图上传", "截图建议尺寸为 330×587，最多可上传 5 张。截图应覆盖落地页核心展示内容，顺序建议与页面实际浏览顺序一致，方便前台按预期展示。"),
        ("详细介绍", "APP 详细介绍用于补充应用卖点或功能说明，建议直接描述核心内容，避免文案过长或与页面素材不一致。"),
        ("保存检查", "点击“确定”后保存新增内容；点击“取消”或关闭弹窗会放弃本次未保存修改。提交前应重点复核模板样式、图标、下载地址、评分、截图数量和详细介绍，确保投放页素材完整可用。"),
    ]:
        last = add_bullet_after(last, label, body)


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    BACKUP_DIR.mkdir(exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = BACKUP_DIR / f"{DOCX_PATH.stem}.landing-page-management-original-{stamp}.bak.docx"
    shutil.copy2(DOCX_PATH, backup_path)

    document = Document(DOCX_PATH)
    fill_section(document)

    try:
        document.save(DOCX_PATH)
    except PermissionError:
        print(f"PERMISSION_ERROR backup={backup_path}")
        raise

    existing = sum(1 for name, _, _ in IMAGE_SPECS if (ASSET_DIR / name).exists())
    print(f"updated landing page management original-only; backup={backup_path}; existing_original_images={existing}")


if __name__ == "__main__":
    main()
